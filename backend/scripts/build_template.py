"""
build_template.py  —  Generate a clean OziCyber pentest report template.

Uses <<placeholder>> markers for scalars and <<f.field>> / <<r.field>> /
<<v.field>> / <<s.field>> markers in table rows for loops.  No Jinja2 —
works with the python-docx generator in reports/generator.py.

Run inside the backend container:
    docker exec ozireport-backend python /app/scripts/build_template.py

Output:  /app/report_fixtures/template.docx
"""

import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = '/app/report_fixtures/template.docx'

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY   = '1B3A6B'
TEAL   = '006D75'
LGREY  = 'F0F3F8'
MGREY  = 'D0D8E8'
DGREY  = '3C3C3C'
WHITE  = 'FFFFFF'
C_CRIT = 'AE0000'
C_HIGH = 'C75000'
C_MED  = 'A07000'
C_LOW  = '2E6B2E'
C_INFO = '175FA0'
FONT   = 'Calibri'


# ── XML helpers ───────────────────────────────────────────────────────────────

def _rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def shade(cell, hexcol):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hexcol.lstrip('#'))
    tcPr.append(shd)


def table_borders(table, size=4, color=MGREY):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    brd = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color.lstrip('#'))
        brd.append(el)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(brd)


def col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def cp(cell, text, bold=False, italic=False, size=Pt(10), color=DGREY,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=Pt(3), sa=Pt(3)):
    """Write text into the first paragraph of a cell."""
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = sb
    p.paragraph_format.space_after  = sa
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = size; run.font.bold = bold
    run.font.italic = italic; run.font.color.rgb = _rgb(color)
    return p


def bp(doc, text='', bold=False, italic=False, size=Pt(10), color=DGREY,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=Pt(0), sa=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = sb
    p.paragraph_format.space_after  = sa
    if text:
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = size; r.font.bold = bold
        r.font.italic = italic; r.font.color.rgb = _rgb(color)
    return p


def hp(doc, text, level=1, sb=None):
    p = doc.add_paragraph()
    sizes   = {1: Pt(16), 2: Pt(13), 3: Pt(11)}
    colors  = {1: NAVY,   2: TEAL,   3: NAVY}
    p.paragraph_format.space_before = sb or (Pt(18) if level == 1 else Pt(12))
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name = FONT; r.font.bold = True
    r.font.size = sizes.get(level, Pt(10))
    r.font.color.rgb = _rgb(colors.get(level, DGREY))
    return p


def pgbreak(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p.add_run()._r.append(br)


# ── Build ─────────────────────────────────────────────────────────────────────

doc = Document()

sec = doc.sections[0]
sec.page_width    = Inches(8.27)
sec.page_height   = Inches(11.69)
sec.left_margin   = Inches(1.18)
sec.right_margin  = Inches(1.18)
sec.top_margin    = Inches(0.98)
sec.bottom_margin = Inches(0.98)

doc.styles['Normal'].font.name = FONT
doc.styles['Normal'].font.size = Pt(10)


# ═══ COVER PAGE ═══════════════════════════════════════════════════════════════

banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.LEFT
bc = banner.cell(0, 0)
shade(bc, NAVY)
p = bc.paragraphs[0]
p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after  = Pt(14)
r = p.add_run('OziCyber Security')
r.font.name = FONT; r.font.size = Pt(24); r.font.bold = True
r.font.color.rgb = _rgb(WHITE)

bp(doc, sb=Pt(36))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run('<<client_name>>')
r.font.name = FONT; r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = _rgb(NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(32)
r = p.add_run('<<report_title>>')
r.font.name = FONT; r.font.size = Pt(20); r.font.bold = False
r.font.color.rgb = _rgb(TEAL)

meta = doc.add_table(rows=7, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
table_borders(meta)
col_widths(meta, [4.5, 9.5])
for i, (label, val) in enumerate([
    ('Engagement Reference', '<<project_id>>'),
    ('Engagement Type',      '<<engagement_type>>'),
    ('Report Version',       '<<report_version>>'),
    ('Classification',       'CONFIDENTIAL'),
    ('Status',               '<<is_draft>>'),
    ('Date',                 '<<today>>'),
    ('Prepared By',          '<<author_name>>'),
]):
    shade(meta.rows[i].cells[0], LGREY)
    cp(meta.rows[i].cells[0], label, bold=True, color=NAVY)
    cp(meta.rows[i].cells[1], val,   color=DGREY)

bp(doc, sb=Pt(36))

conf = doc.add_table(rows=1, cols=1)
cc = conf.cell(0, 0)
shade(cc, LGREY)
p = cc.paragraphs[0]
p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)
r = p.add_run(
    'CONFIDENTIAL — This document contains sensitive security information. '
    'Distribution is restricted to authorised recipients only.'
)
r.font.name = FONT; r.font.size = Pt(9); r.font.italic = True
r.font.color.rgb = _rgb('606060')

pgbreak(doc)


# ═══ DOCUMENT CONTROL ═════════════════════════════════════════════════════════

hp(doc, 'Document Control')
bp(doc, 'This document is version controlled. The table below records all revisions.',
   color=DGREY)

ver = doc.add_table(rows=2, cols=4)
table_borders(ver)
col_widths(ver, [2.5, 3.5, 4.5, 6.0])

for i, lbl in enumerate(['Version', 'Date', 'Author', 'Description']):
    shade(ver.rows[0].cells[i], NAVY)
    cp(ver.rows[0].cells[i], lbl, bold=True, color=WHITE)

# Template row — generator clones this for each version entry
cp(ver.rows[1].cells[0], '<<v.version>>', color=DGREY)
cp(ver.rows[1].cells[1], '<<v.date>>',    color=DGREY)
cp(ver.rows[1].cells[2], '<<v.author>>',  color=DGREY)
cp(ver.rows[1].cells[3], '<<v.comment>>', color=DGREY)

pgbreak(doc)


# ═══ DISCLAIMER ═══════════════════════════════════════════════════════════════

hp(doc, 'Disclaimer')
for para_text in [
    'This report has been prepared by OziCyber Security solely for the benefit of '
    '<<client_name>> (the "Client"). The information contained in this report is '
    'confidential and is intended only for the Client. OziCyber Security accepts no '
    'liability to any third party who may obtain access to this report.',

    "The findings described in this report reflect the state of the Client's environment "
    'at the time of testing only. The security posture of any environment may change '
    'over time; accordingly, this report should not be relied upon as a definitive '
    'assessment of security at any time other than the period of testing specified herein.',

    'All testing was conducted within the agreed scope and in accordance with the rules '
    'of engagement defined prior to commencement. OziCyber Security shall not be held '
    'liable for any damage resulting from actions taken outside the agreed scope.',
]:
    bp(doc, para_text, color=DGREY, sa=Pt(10))

pgbreak(doc)


# ═══ 1. EXECUTIVE SUMMARY ═════════════════════════════════════════════════════

hp(doc, '1. Executive Summary')

risk = doc.add_table(rows=1, cols=2)
table_borders(risk)
col_widths(risk, [6.5, 10.5])
shade(risk.cell(0, 0), NAVY)
cp(risk.cell(0, 0), 'Overall Risk Rating', bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
cp(risk.cell(0, 1), '<<overall_risk>>', bold=True, color=DGREY,
   align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))

bp(doc, sb=Pt(10))

# Chart placeholder — generator replaces this paragraph with the bar chart image
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run('<<SEVERITY_CHART>>')
r.font.name = FONT; r.font.size = Pt(10); r.font.color.rgb = _rgb(LGREY)

bp(doc, sb=Pt(6))

# Severity count table
scnt = doc.add_table(rows=2, cols=5)
table_borders(scnt)
col_widths(scnt, [3.4, 3.4, 3.4, 3.4, 3.4])
for i, (lbl, col) in enumerate([
    ('Critical', C_CRIT), ('High', C_HIGH), ('Medium', C_MED),
    ('Low', C_LOW), ('Informational', C_INFO),
]):
    shade(scnt.rows[0].cells[i], col)
    cp(scnt.rows[0].cells[i], lbl, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, val in enumerate(['<<critical_count>>', '<<high_count>>', '<<medium_count>>',
                          '<<low_count>>', '<<info_count>>']):
    shade(scnt.rows[1].cells[i], LGREY)
    cp(scnt.rows[1].cells[i], val, bold=True, color=NAVY,
       align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(16))

bp(doc, sb=Pt(10))
bp(doc, '<<executive_summary>>', color=DGREY, sa=Pt(10))

pgbreak(doc)


# ═══ 2. ENGAGEMENT OVERVIEW ════════════════════════════════════════════════════

hp(doc, '2. Engagement Overview')

eng = doc.add_table(rows=10, cols=2)
table_borders(eng)
col_widths(eng, [6.0, 11.0])
for i, (lbl, val) in enumerate([
    ('Engagement Reference', '<<project_id>>'),
    ('Client Organisation',  '<<client_name>>'),
    ('Contact Name',         '<<client_contact_name>>'),
    ('Contact Email',        '<<client_contact_email>>'),
    ('Contact Phone',        '<<client_contact_phone>>'),
    ('Engagement Type',      '<<engagement_type>>'),
    ('Start Date',           '<<engagement_start>>'),
    ('End Date',             '<<engagement_end>>'),
    ('Report Due',           '<<report_due_date>>'),
    ('Status',               '<<engagement_status>>'),
]):
    shade(eng.rows[i].cells[0], LGREY)
    cp(eng.rows[i].cells[0], lbl, bold=True, color=NAVY)
    shade(eng.rows[i].cells[1], WHITE if i % 2 else LGREY)
    cp(eng.rows[i].cells[1], val, color=DGREY)

bp(doc, sb=Pt(12))
hp(doc, '2.1  Testing Team', level=2)

team = doc.add_table(rows=3, cols=2)
table_borders(team)
col_widths(team, [6.0, 11.0])
for i, (lbl, val) in enumerate([
    ('Lead Pentester',  '<<lead_tester>> (<<lead_tester_email>>)'),
    ('Project Manager', '<<project_manager>> (<<project_manager_email>>)'),
    ('Report Author',   '<<author_name>>'),
]):
    shade(team.rows[i].cells[0], LGREY)
    cp(team.rows[i].cells[0], lbl, bold=True, color=NAVY)
    shade(team.rows[i].cells[1], WHITE if i % 2 else LGREY)
    cp(team.rows[i].cells[1], val, color=DGREY)

bp(doc, sb=Pt(12))
hp(doc, '2.2  Objectives', level=2)
bp(doc, '<<objectives>>', color=DGREY)

pgbreak(doc)


# ═══ 3. SCOPE ═════════════════════════════════════════════════════════════════

hp(doc, '3. Scope of Assessment')
hp(doc, '3.1  In-Scope Targets', level=2)
bp(doc, 'The following assets and systems were included in the scope of this assessment:',
   color=DGREY)

sc = doc.add_table(rows=2, cols=2)
table_borders(sc)
col_widths(sc, [1.5, 15.5])
shade(sc.rows[0].cells[0], NAVY); cp(sc.rows[0].cells[0], '#', bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
shade(sc.rows[0].cells[1], NAVY); cp(sc.rows[0].cells[1], 'Target / Asset', bold=True, color=WHITE)
# Template row — generator finds <<s. prefix, clones for each scope item
shade(sc.rows[1].cells[0], LGREY); cp(sc.rows[1].cells[0], '•', color=DGREY, align=WD_ALIGN_PARAGRAPH.CENTER)
cp(sc.rows[1].cells[1], '<<s.item>>', color=DGREY)

bp(doc, sb=Pt(12))
hp(doc, '3.2  Out-of-Scope Items', level=2)
bp(doc, 'The following were explicitly excluded from the scope of this engagement:', color=DGREY)
bp(doc, '<<out_of_scope_list>>', color=DGREY)

pgbreak(doc)


# ═══ 4. METHODOLOGY ═══════════════════════════════════════════════════════════

hp(doc, '4. Methodology')
bp(doc, '<<methodology>>', color=DGREY, sa=Pt(10))

pgbreak(doc)


# ═══ 5. FINDINGS SUMMARY ══════════════════════════════════════════════════════

hp(doc, '5. Findings Summary')
bp(doc,
   'A total of <<finding_count>> finding(s) were identified during this engagement. '
   'The table below provides an overview of all findings by severity.',
   color=DGREY)

vl = doc.add_table(rows=2, cols=4)
table_borders(vl)
col_widths(vl, [2.5, 2.5, 3.0, 9.0])
for i, lbl in enumerate(['#', 'Reference', 'Severity', 'Finding Title']):
    shade(vl.rows[0].cells[i], NAVY)
    cp(vl.rows[0].cells[i], lbl, bold=True, color=WHITE)
# Template row — generator finds <<f. prefix, clones for each finding
cp(vl.rows[1].cells[0], '<<f.number>>',          color=DGREY, align=WD_ALIGN_PARAGRAPH.CENTER)
cp(vl.rows[1].cells[1], '<<f.code>>',             color=NAVY,  bold=True)
cp(vl.rows[1].cells[2], '<<f.severity_display>>', color=DGREY)
cp(vl.rows[1].cells[3], '<<f.title>>',            color=DGREY)

pgbreak(doc)


# ═══ 6. FINDINGS DETAIL ════════════════════════════════════════════════════════

hp(doc, '6. Findings Detail')
bp(doc,
   'Detailed information for each finding is provided below, including description, '
   'impact, supporting evidence, and remediation guidance.',
   color=DGREY, sa=Pt(14))

# One 10-row template table — generator clones rows 1..9 for each finding,
# replacing the <<f.field>> markers. Row 0 is the styled header row.
# Because all 10 rows share the same <<f. markers, _process_loop_tables will
# find row 0 (which has <<f.code>>) and clone the entire table isn't right —
# we need a different approach for multi-row per-finding blocks.
#
# Solution: use a single-row table per-finding-field isn't practical.
# Instead we use a PARAGRAPH-based approach: one styled block paragraph per
# finding using <<f. markers — the generator detects these aren't in tables
# and handles them as inline replacements after the loop table.
#
# For this template we put a single compact table with all fields in 2 rows
# per finding to make the cloning work: the template row IS the entire block.
# We use a 1-row (per finding) wide table with merged cells and line breaks.

ftbl = doc.add_table(rows=2, cols=4)
table_borders(ftbl)
col_widths(ftbl, [4.5, 4.5, 4.5, 3.5])

# ── Row 0: header row (stays, not cloned) ──
shade(ftbl.rows[0].cells[0], NAVY)
cp(ftbl.rows[0].cells[0], 'Reference', bold=True, color=WHITE)
shade(ftbl.rows[0].cells[1], NAVY)
cp(ftbl.rows[0].cells[1], 'Finding Title', bold=True, color=WHITE)
ftbl.rows[0].cells[1].merge(ftbl.rows[0].cells[2])
shade(ftbl.rows[0].cells[3], NAVY)
cp(ftbl.rows[0].cells[3], 'Severity', bold=True, color=WHITE)

# ── Row 1: template row — ONE row per finding, multi-line via \n in text ──
# Generator clones this row for each finding and fills <<f.field>> markers.
shade(ftbl.rows[1].cells[0], LGREY)
cp(ftbl.rows[1].cells[0], '<<f.code>>', bold=True, color=NAVY, size=Pt(11))
ftbl.rows[1].cells[1].merge(ftbl.rows[1].cells[2])
# Pack multiple fields into one cell using labelled lines
p = ftbl.rows[1].cells[1].paragraphs[0]
p.clear()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after  = Pt(3)

def _labelled_run(para, label, marker, last=False):
    rb = para.add_run(label)
    rb.font.name = FONT; rb.font.size = Pt(9); rb.font.bold = True
    rb.font.color.rgb = _rgb(TEAL)
    rv = para.add_run(marker + ('' if last else '\n'))
    rv.font.name = FONT; rv.font.size = Pt(9); rv.font.bold = False
    rv.font.color.rgb = _rgb(DGREY)

_labelled_run(p, 'Title: ',       '<<f.title>>')
_labelled_run(p, 'Asset: ',       '<<f.affected_asset>>')
_labelled_run(p, 'Type: ',        '<<f.pentest_type>>')
_labelled_run(p, 'Status: ',      '<<f.status_display>>')
_labelled_run(p, 'CVSS: ',        '<<f.cvss_score>>')
_labelled_run(p, 'CWE: ',         '<<f.cwe_id>>')
_labelled_run(p, 'CVE: ',         '<<f.cve_id>>')
_labelled_run(p, 'Description:\n','<<f.description>>')
_labelled_run(p, '\nImpact:\n',   '<<f.impact>>')
_labelled_run(p, '\nEvidence:\n', '<<f.supporting_evidence>>')
_labelled_run(p, '\nRecommendations:\n', '<<f.recommendations>>', last=True)

shade(ftbl.rows[1].cells[3], LGREY)
p3 = ftbl.rows[1].cells[3].paragraphs[0]
p3.clear()
p3.paragraph_format.space_before = Pt(3)
p3.paragraph_format.space_after  = Pt(3)
_labelled_run(p3, 'Severity:\n',    '<<f.severity_display>>')
_labelled_run(p3, '\nConsequence:\n','<<f.consequence>>')
_labelled_run(p3, '\nLikelihood:\n', '<<f.likelihood_label>>')
_labelled_run(p3, '\nRefs:\n',       '<<f.references>>', last=True)

pgbreak(doc)


# ═══ 7. REMEDIATION TRACKER ════════════════════════════════════════════════════

hp(doc, '7. Remediation Tracker')
bp(doc,
   'The table below tracks the remediation status of all findings identified during '
   'this engagement.',
   color=DGREY)

rem = doc.add_table(rows=2, cols=5)
table_borders(rem)
col_widths(rem, [2.5, 2.5, 7.0, 2.5, 2.5])
for i, lbl in enumerate(['Reference', 'Severity', 'Finding Title', 'Remediated', 'Date']):
    shade(rem.rows[0].cells[i], NAVY)
    cp(rem.rows[0].cells[i], lbl, bold=True, color=WHITE)
# Template row
cp(rem.rows[1].cells[0], '<<r.code>>',       bold=True, color=NAVY)
cp(rem.rows[1].cells[1], '<<r.severity>>',   color=DGREY)
cp(rem.rows[1].cells[2], '<<r.title>>',      color=DGREY)
cp(rem.rows[1].cells[3], '<<r.remediated>>', color=DGREY, align=WD_ALIGN_PARAGRAPH.CENTER)
cp(rem.rows[1].cells[4], '<<r.date>>',       color=DGREY, align=WD_ALIGN_PARAGRAPH.CENTER)

pgbreak(doc)


# ═══ 8. CONCLUSION ════════════════════════════════════════════════════════════

hp(doc, '8. Conclusion')
bp(doc, '<<conclusion>>', color=DGREY, sa=Pt(10))
bp(doc, '<<client_notes>>', color=DGREY, italic=True, sa=Pt(10))

pgbreak(doc)


# ═══ APPENDIX A — CVSS ════════════════════════════════════════════════════════

hp(doc, 'Appendix A — CVSS Score Reference')
bp(doc,
   'The following table lists the CVSS v3.1 base metric values for each finding.',
   color=DGREY)

cvss = doc.add_table(rows=2, cols=10)
table_borders(cvss)
col_widths(cvss, [2.0, 1.8, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5])
for i, lbl in enumerate(['Finding', 'Score', 'AV', 'AC', 'PR', 'UI', 'S', 'C', 'I', 'A']):
    shade(cvss.rows[0].cells[i], NAVY)
    cp(cvss.rows[0].cells[i], lbl, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, val in enumerate([
    '<<f.code>>', '<<f.cvss_score>>', '<<f.cvss_av>>', '<<f.cvss_ac>>',
    '<<f.cvss_pr>>', '<<f.cvss_ui>>', '<<f.cvss_s>>', '<<f.cvss_c>>',
    '<<f.cvss_i>>', '<<f.cvss_a>>',
]):
    shade(cvss.rows[1].cells[i], LGREY if i % 2 == 0 else WHITE)
    cp(cvss.rows[1].cells[i], val, color=DGREY, align=WD_ALIGN_PARAGRAPH.CENTER)

bp(doc, sb=Pt(14))
hp(doc, 'CVSS Metric Legend', level=2)

leg = doc.add_table(rows=9, cols=3)
table_borders(leg)
col_widths(leg, [3.5, 5.5, 8.0])
for i, lbl in enumerate(['Metric', 'Abbreviation', 'Values']):
    shade(leg.rows[0].cells[i], NAVY)
    cp(leg.rows[0].cells[i], lbl, bold=True, color=WHITE)
for i, (metric, abbr, vals) in enumerate([
    ('Attack Vector',          'AV', 'Network / Adjacent / Local / Physical'),
    ('Attack Complexity',      'AC', 'Low / High'),
    ('Privileges Required',    'PR', 'None / Low / High'),
    ('User Interaction',       'UI', 'None / Required'),
    ('Scope',                  'S',  'Unchanged / Changed'),
    ('Confidentiality Impact', 'C',  'None / Low / High'),
    ('Integrity Impact',       'I',  'None / Low / High'),
    ('Availability Impact',    'A',  'None / Low / High'),
]):
    bg = LGREY if i % 2 == 0 else WHITE
    for j, txt in enumerate([metric, abbr, vals]):
        shade(leg.rows[i + 1].cells[j], bg)
        cp(leg.rows[i + 1].cells[j], txt, color=DGREY)


# ── Save ──────────────────────────────────────────────────────────────────────

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f'Template saved → {OUT_PATH}')
print()
print('Placeholder syntax:')
print('  Scalars  →  <<placeholder_name>>  (e.g. <<report_title>>, <<client_name>>)')
print('  Findings →  <<f.field>>  in a table row  (generator clones the row per finding)')
print('  Remediation  →  <<r.field>>')
print('  Version log  →  <<v.field>>')
print('  Scope items  →  <<s.item>>')
print('  Chart        →  <<SEVERITY_CHART>>  paragraph (replaced with bar chart PNG)')
