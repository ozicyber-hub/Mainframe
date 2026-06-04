"""
Transform OziCyber template.docx into a docxtpl-compatible template.

Run inside the backend container:
  docker exec ozireport-backend python /app/scripts/make_template.py

Output: /app/report_fixtures/template_docxtpl.docx
Upload this file via Repository → Report Templates, then set it as default.

Placeholders added
==================
Cover/header
  {{ client_name }}         Organisation / client name
  {{ report_title }}        Report title
  {{ project_id }}          e.g. ENG-0001
  {{ author_name }}         Lead tester full name
  {{ today }}               Generation date

Executive summary
  {{ executive_summary }}   Plain-text exec summary

Scope table (one row per item)
  {{ item }}                Each scope item

Vuln list table (one row per finding)
  {{ f.severity }}
  {{ f.code }}              e.g. WEB-01
  {{ f.title }}

Finding detail table (8 rows × 4 cols, repeated per finding)
  {{ f.code }}
  {{ f.title }}
  {{ f.severity }}
  {{ f.affected_asset }}
  {{ f.consequence }}       e.g. Moderate
  {{ f.likelihood_label }}  e.g. Likely
  {{ f.description }}
  {{ f.impact }}
  {{ f.supporting_evidence }}
  {{ f.recommendations }}

Remediation tracker (one row per finding)
  {{ r.code }}
  {{ r.title }}
  {{ r.remediated }}        Yes / No
  {{ r.date }}

Version control (one row per entry)
  {{ v.version }}
  {{ v.date }}
  {{ v.author }}
  {{ v.comment }}
"""

import os
from docx import Document
from docx.oxml import OxmlElement

TEMPLATE_IN  = '/app/report_fixtures/template.docx'
TEMPLATE_OUT = '/app/report_fixtures/template_docxtpl.docx'

doc = Document(TEMPLATE_IN)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def para_replace(para, old, new):
    """Replace text in a paragraph, handling runs split across word boundaries."""
    if old not in para.text:
        return
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    # Split-run fallback: collapse all runs into the first
    replacement = para.text.replace(old, new)
    for i, run in enumerate(para.runs):
        run.text = replacement if i == 0 else ''

def replace_everywhere(old, new):
    for p in doc.paragraphs:
        para_replace(p, old, new)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    para_replace(p, old, new)

def cell_set(cell, text):
    """Set a cell's text (first paragraph, first run). Preserves run formatting."""
    paras = cell.paragraphs
    if not paras:
        return
    p = paras[0]
    # Remove extra paragraphs beyond the first
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)

def cell_append_tag(cell, tag):
    """Add a docxtpl control tag as a new paragraph at the end of a cell."""
    p = cell.add_paragraph()
    p.add_run(tag)

def remove_rows_from(table, keep_header=1):
    """Remove all data rows, keeping the first `keep_header` rows."""
    for row in list(table.rows)[keep_header:]:
        row._element.getparent().remove(row._element)


# ─── 1. Cover page ────────────────────────────────────────────────────────────

# Paragraph 0 is the Title: "CLIENTNAME\nReport Type"
p0 = doc.paragraphs[0]
if p0.runs:
    p0.runs[0].text = '{{ client_name }}\n{{ report_title }}'
    for r in p0.runs[1:]:
        r.text = ''

# Project ID and author (appear in body paragraphs and version table)
replace_everywhere('OCT-1002',      '{{ project_id }}')
replace_everywhere('By Alex Young', 'By {{ author_name }}')
replace_everywhere('Alexanda Young','{{ author_name }}')


# ─── 2. Executive summary ─────────────────────────────────────────────────────

# Find the body-text paragraphs immediately after the "Executive Summary" heading
# and replace the first content paragraph with {{ executive_summary }}
in_exec = False
replaced_exec = False
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and 'Executive Summary' in p.text:
        in_exec = True
        continue
    if in_exec and p.style.name.startswith('Heading'):
        in_exec = False  # reached next heading
    if in_exec and p.text.strip() and not replaced_exec:
        if p.runs:
            p.runs[0].text = '{{ executive_summary }}'
            for r in p.runs[1:]:
                r.text = ''
        replaced_exec = True
    elif in_exec and p.text.strip() and replaced_exec:
        # Blank out subsequent exec summary paragraphs
        for r in p.runs:
            r.text = ''


# ─── 3. Scope table (T0) – single loop row ────────────────────────────────────

scope_tbl = doc.tables[0]
remove_rows_from(scope_tbl, keep_header=1)
# Add one template row by cloning the header row structure
from copy import deepcopy
tmpl_row_elem = deepcopy(list(scope_tbl.rows)[0]._element)
scope_tbl._element.append(tmpl_row_elem)
# Now set that new row's content
new_row = scope_tbl.rows[-1]
cells = new_row.cells
cell_set(cells[0], '{%tr for item in scope_items %}{{ item }}')
if len(cells) > 1:
    cell_set(cells[1], '{%tr endfor %}')


# ─── 4. Vulnerability list table (T2) ─────────────────────────────────────────

vuln_tbl = doc.tables[2]
remove_rows_from(vuln_tbl, keep_header=1)
tmpl_row_elem = deepcopy(list(vuln_tbl.rows)[0]._element)
vuln_tbl._element.append(tmpl_row_elem)
new_row = vuln_tbl.rows[-1]
cells = new_row.cells
cell_set(cells[0], '{%tr for f in findings %}{{ f.severity }}')
if len(cells) > 1:
    cell_set(cells[1], '{{ f.code }}')
if len(cells) > 2:
    cell_set(cells[2], '{{ f.title }}{%tr endfor %}')


# ─── 5. Finding tables (T3–T7): keep T3, add loop tags, delete T4–T7 ─────────

ftbl = doc.tables[3]
rows = ftbl.rows

# Row 0: code | title | risk rating (cols 2+3 are merged in the original)
r0 = rows[0].cells
cell_set(r0[0], '{%tr for f in findings %}{{ f.code }}')
cell_set(r0[1], '{{ f.title }}')
cell_set(r0[2], 'RISK RATING: {{ f.severity }}')

# Rows 1-3: affected asset (cols 0+1 merged) | label (static) | value
for row_idx, value_tag in [
    (1, '{{ f.consequence }}'),
    (2, '{{ f.consequence }}'),
    (3, '{{ f.likelihood_label }}'),
]:
    r_cells = rows[row_idx].cells
    cell_set(r_cells[0], '{{ f.affected_asset }}')
    # r_cells[2] is the static label ("Level of Access:", "Consequence:", "Likelihood:") — leave it
    cell_set(r_cells[3], value_tag)

# Rows 4-7: fully merged narrative rows (all 4 cell references are the same element)
cell_set(rows[4].cells[0], 'Description\n{{ f.description }}')
cell_set(rows[5].cells[0], 'Consequence\n{{ f.impact }}')
cell_set(rows[6].cells[0], 'Supporting Evidence\n{{ f.supporting_evidence }}')
cell_set(rows[7].cells[0], 'Remediation\n{{ f.recommendations }}')
# Add endfor as a separate paragraph in row 7's cell so it's clearly separated
cell_append_tag(rows[7].cells[0], '{%tr endfor %}')

# Delete T4–T7 in reverse order (so indices stay valid)
for tbl_idx in range(7, 3, -1):
    tbl_elem = doc.tables[tbl_idx]._element
    tbl_elem.getparent().remove(tbl_elem)


# ─── 6. Remediation tracker table ─────────────────────────────────────────────

rem_tbl = None
for tbl in doc.tables:
    if tbl.rows and tbl.rows[0].cells[0].text.strip() == 'Finding Reference':
        rem_tbl = tbl
        break

if rem_tbl:
    remove_rows_from(rem_tbl, keep_header=1)
    tmpl_row_elem = deepcopy(list(rem_tbl.rows)[0]._element)
    rem_tbl._element.append(tmpl_row_elem)
    new_row = rem_tbl.rows[-1]
    c = new_row.cells
    cell_set(c[0], '{%tr for r in remediation_items %}{{ r.code }}')
    if len(c) > 1: cell_set(c[1], '{{ r.title }}')
    if len(c) > 2: cell_set(c[2], '{{ r.remediated }}')
    if len(c) > 3: cell_set(c[3], '{{ r.date }}')
    if len(c) > 4: cell_set(c[4], '{%tr endfor %}')


# ─── 7. Version control table ─────────────────────────────────────────────────

ver_tbl = None
for tbl in doc.tables:
    if tbl.rows:
        hdrs = [c.text.strip() for c in tbl.rows[0].cells]
        if 'Version' in hdrs and 'Date' in hdrs and 'Name' in hdrs:
            ver_tbl = tbl
            break

if ver_tbl:
    # Keep header row only, add a loop row for generated entries
    remove_rows_from(ver_tbl, keep_header=1)
    tmpl_row_elem = deepcopy(list(ver_tbl.rows)[0]._element)
    ver_tbl._element.append(tmpl_row_elem)
    new_row = ver_tbl.rows[-1]
    c = new_row.cells
    cell_set(c[0], '{%tr for v in version_rows %}{{ v.version }}')
    if len(c) > 1: cell_set(c[1], '{{ v.date }}')
    if len(c) > 2: cell_set(c[2], '{{ v.author }}')
    if len(c) > 3: cell_set(c[3], '{{ v.comment }}{%tr endfor %}')


# ─── 8. Appendix I scoped targets table (T9 → now T5) ────────────────────────

app1_tbl = None
for tbl in doc.tables:
    if tbl.rows and tbl.rows[0].cells[0].text.strip() == 'Domain':
        app1_tbl = tbl
        break

if app1_tbl:
    remove_rows_from(app1_tbl, keep_header=1)
    tmpl_row_elem = deepcopy(list(app1_tbl.rows)[0]._element)
    app1_tbl._element.append(tmpl_row_elem)
    new_row = app1_tbl.rows[-1]
    c = new_row.cells
    cell_set(c[0], '{%tr for item in scope_items %}{{ item }}{%tr endfor %}')


# ─── 9. Save ──────────────────────────────────────────────────────────────────

os.makedirs(os.path.dirname(TEMPLATE_OUT), exist_ok=True)
doc.save(TEMPLATE_OUT)
print(f'Saved: {TEMPLATE_OUT}')
print()
print('Next steps:')
print('  1. Download the file from the backend container or media folder')
print('  2. Open in Word to review / tweak the placeholders')
print('  3. Upload via Repository → Report Templates → Upload Template')
print('  4. Set as Default (star icon)')
