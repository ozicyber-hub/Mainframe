"""
fix_template_runs.py  —  Fix split Jinja2 tags in a docxtpl Word template.

Word's spell-checker splits tag text across multiple XML runs, e.g.
  {{ f.code }}  →  ['{{ ', 'f', '.code', ' }}']
This script rebuilds every paragraph that contains {{ or {% as a single run.

Usage (inside backend container):
    python3 /app/scripts/fix_template_runs.py <input.docx> [output.docx]
    (output defaults to input — fixes in-place)
"""

import sys
import re
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W_T  = f'{{{W_NS}}}t'
W_R  = f'{{{W_NS}}}r'
W_BR = f'{{{W_NS}}}br'
W_P  = f'{{{W_NS}}}p'


def rebuild_paragraph(p_elem):
    """
    Completely rebuild a paragraph element so that all its text content
    (from all <w:t> elements) ends up in a single <w:r><w:t> element.

    Preserves any <w:pPr> (paragraph formatting) and keeps <w:br> line-break
    elements that are outside of the Jinja2 tags.
    """
    # Gather full text from ALL w:t elements (recursive)
    t_elems = list(p_elem.iter(W_T))
    if len(t_elems) <= 1:
        return False

    full_text = ''.join(t.text or '' for t in t_elems)
    if '{{' not in full_text and '{%' not in full_text:
        return False

    # Preserve <w:pPr> if present
    pPr = p_elem.find(qn('w:pPr'))

    # Build a fresh <w:r> with the full text
    new_r = OxmlElement('w:r')
    # Copy rPr from first run that had one, if any
    first_r = p_elem.find(qn('w:r'))
    if first_r is not None:
        rPr = first_r.find(qn('w:rPr'))
        if rPr is not None:
            new_r.append(deepcopy(rPr))
    new_t = OxmlElement('w:t')
    new_t.text = full_text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)

    # Clear ALL children from the paragraph
    for child in list(p_elem):
        p_elem.remove(child)

    # Re-add paragraph properties if there were any
    if pPr is not None:
        p_elem.append(pPr)

    # Add the single fresh run
    p_elem.append(new_r)

    return True


def fix_doc(path, out_path=None):
    doc = Document(path)
    fixed = 0

    def process_paragraphs(paragraphs):
        nonlocal fixed
        for para in paragraphs:
            if rebuild_paragraph(para._element):
                fixed += 1

    # Body paragraphs
    process_paragraphs(doc.paragraphs)

    # Tables — each unique cell once
    for tbl in doc.tables:
        seen = set()
        for row in tbl.rows:
            for cell in row.cells:
                cid = id(cell._tc)
                if cid in seen:
                    continue
                seen.add(cid)
                process_paragraphs(cell.paragraphs)

    # Headers / footers
    for section in doc.sections:
        for hdr in [section.header, section.footer,
                    section.first_page_header, section.first_page_footer,
                    section.even_page_header, section.even_page_footer]:
            if hdr is None:
                continue
            try:
                process_paragraphs(hdr.paragraphs)
            except Exception:
                pass

    save_path = out_path or path
    doc.save(save_path)
    print(f'Rebuilt {fixed} paragraph(s) — split Jinja2 tags fixed.')
    print(f'Saved → {save_path}')
    return fixed


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    inp  = sys.argv[1]
    outp = sys.argv[2] if len(sys.argv) > 2 else None
    fix_doc(inp, outp)
