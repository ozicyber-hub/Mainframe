from django.http import HttpResponse
from django.utils import timezone
from rest_framework import viewsets, status, generics
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated

from .models import (
    AssessmentTemplate, Assessment, AssessmentResponse, AssessmentEvidence,
    SectionEvidenceRequirement, EvidenceSubmission,
)
from .serializers import (
    AssessmentTemplateSerializer, AssessmentTemplateListSerializer,
    AssessmentSerializer, AssessmentListSerializer,
    AssessmentResponseSerializer, AssessmentEvidenceSerializer,
    SectionEvidenceRequirementSerializer, EvidenceSubmissionSerializer,
)


def get_org(request):
    return getattr(request.user, 'organization', None)


class AssessmentTemplateViewSet(viewsets.ReadOnlyModelViewSet):
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        if self.action == 'retrieve':
            return AssessmentTemplateSerializer
        return AssessmentTemplateListSerializer

    def get_queryset(self):
        return AssessmentTemplate.objects.filter(is_active=True).prefetch_related(
            'sections__questions', 'sections__evidence_requirements',
        )

    @action(detail=True, methods=['post'])
    def suggest_requirements(self, request, pk=None):
        """
        AI-generate evidence requirements for a section from its controls.
        Idempotent — returns existing requirements if already generated.
        """
        import json, re as _re

        template   = self.get_object()
        section_id = request.data.get('section_id')
        if not section_id:
            return Response({'error': 'section_id required.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            section = template.sections.prefetch_related('questions').get(id=section_id)
        except Exception:
            return Response({'error': 'Section not found.'}, status=status.HTTP_404_NOT_FOUND)

        # Idempotent — return existing if already created
        existing = SectionEvidenceRequirement.objects.filter(section=section)
        if existing.exists():
            return Response(SectionEvidenceRequirementSerializer(existing, many=True).data)

        questions_text = '\n'.join(
            f'- {q.text}' for q in section.questions.all()[:12]
        ) or '(no questions defined)'

        valid_types = 'POLICY|PROCEDURE|PLAN|LOG|REPORT|CERTIFICATION|SCREENSHOT|CONTRACT|TRAINING|OTHER'

        prompt = f"""You are a GRC compliance expert building a client-facing compliance portal like Vanta or Scytale.

Based on the following compliance section from a {template.get_framework_display()} assessment, determine exactly what evidence documents a client organisation must submit to demonstrate compliance.

SECTION: {section.name}
DESCRIPTION: {section.description or '(not specified)'}
CONTROLS:
{questions_text}

Return a JSON array of 3–5 evidence requirements. Be specific about what each document must contain. These will be shown to the client as their compliance checklist.

Return ONLY a valid JSON array:
[
  {{
    "title": "<specific document name, e.g. Incident Response Plan>",
    "description": "<1-2 sentences: exactly what this document must contain to satisfy the controls above>",
    "document_type": "<one of: {valid_types}>",
    "required": true,
    "validation_prompt": "<comma-separated specific checks: e.g. 'check for escalation matrix, verify RTO is defined, confirm annual review date present'>"
  }}
]"""

        try:
            raw   = _call_ai(prompt)
            match = _re.search(r'\[[\s\S]*\]', raw)
            if not match:
                return Response({'error': 'AI could not generate requirements.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            items = json.loads(match.group())
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        valid_doc_types = {c[0] for c in SectionEvidenceRequirement.DOC_TYPE_CHOICES}
        created = []
        for i, item in enumerate(items[:5]):
            doc_type = item.get('document_type', 'OTHER')
            if doc_type not in valid_doc_types:
                doc_type = 'OTHER'
            req = SectionEvidenceRequirement.objects.create(
                section           = section,
                title             = item.get('title', 'Untitled'),
                description       = item.get('description', ''),
                document_type     = doc_type,
                required          = bool(item.get('required', True)),
                validation_prompt = item.get('validation_prompt', ''),
                order             = i,
            )
            created.append(req)

        return Response(
            SectionEvidenceRequirementSerializer(created, many=True).data,
            status=status.HTTP_201_CREATED,
        )

    @action(detail=True, methods=['get', 'post'])
    def requirements(self, request, pk=None):
        """List or create evidence requirements for template sections."""
        template  = self.get_object()
        section_id = request.query_params.get('section')

        if request.method == 'GET':
            qs = SectionEvidenceRequirement.objects.filter(section__template=template)
            if section_id:
                qs = qs.filter(section_id=section_id)
            return Response(SectionEvidenceRequirementSerializer(qs, many=True).data)

        # POST — create
        data = request.data.copy()
        sid  = data.get('section')
        if not template.sections.filter(id=sid).exists():
            return Response({'error': 'Section does not belong to this template.'}, status=status.HTTP_400_BAD_REQUEST)
        serializer = SectionEvidenceRequirementSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['patch', 'delete'], url_path=r'requirements/(?P<req_id>[0-9]+)')
    def requirement_detail(self, request, pk=None, req_id=None):
        """Update or delete a single evidence requirement."""
        template = self.get_object()
        try:
            req = SectionEvidenceRequirement.objects.get(id=req_id, section__template=template)
        except SectionEvidenceRequirement.DoesNotExist:
            return Response({'error': 'Not found.'}, status=status.HTTP_404_NOT_FOUND)

        if request.method == 'DELETE':
            req.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)

        serializer = SectionEvidenceRequirementSerializer(req, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class AssessmentViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        if self.action in ('list',):
            return AssessmentListSerializer
        return AssessmentSerializer

    def get_queryset(self):
        from django.db.models import Q
        user = self.request.user
        org  = get_org(self.request)

        if user.is_superuser:
            base = Assessment.objects.all()
        else:
            from engagements.models import Engagement
            # Materialise accessible engagement IDs to avoid complex subquery interactions
            eng_ids = list(
                Engagement.objects.filter(
                    Q(organization__users=user) |
                    Q(lead_pentester=user) |
                    Q(project_manager=user) |
                    Q(team_members=user)
                ).values_list('id', flat=True).distinct()
            )
            base = Assessment.objects.filter(
                Q(organization=org) | Q(engagement_id__in=eng_ids) |
                Q(grc_consultant=user) | Q(assessor=user)
            ).distinct()

        qs = base.select_related('template', 'engagement', 'assessor', 'created_by')

        # Only prefetch heavy nested data for detail/retrieve actions
        if self.action not in ('list',):
            qs = qs.prefetch_related('responses__evidence', 'responses__question')

        framework  = self.request.query_params.get('framework')
        engagement = self.request.query_params.get('engagement')
        status_    = self.request.query_params.get('status')

        if framework:
            qs = qs.filter(template__framework=framework)
        if engagement:
            qs = qs.filter(engagement=engagement)
        if status_:
            qs = qs.filter(status=status_)

        return qs

    def perform_create(self, serializer):
        user = self.request.user
        if not user.is_superuser and user.role not in ('SUPERADMIN', 'ADMIN', 'GRC_CONSULTANT', 'PROJECT_MANAGER'):
            from rest_framework.exceptions import PermissionDenied
            raise PermissionDenied('Your role does not allow creating assessments.')
        from engagements.models import Engagement
        org = get_org(self.request)
        engagement_id = self.request.data.get('engagement')
        if engagement_id:
            try:
                org = Engagement.objects.get(id=engagement_id).organization
            except Engagement.DoesNotExist:
                pass
        assessment = serializer.save(organization=org, created_by=self.request.user)

        # Notify assigned GRC consultant
        if assessment.grc_consultant and assessment.grc_consultant != self.request.user:
            try:
                from notifications.utils import notify_users
                notify_users(
                    [assessment.grc_consultant],
                    'ASSESSMENT_ASSIGNED',
                    f'Assessment assigned: {assessment.title}',
                    f'You have been assigned as GRC Consultant for "{assessment.title}". '
                    f'{"Scheduled: " + str(assessment.start_date) + " → " + str(assessment.end_date) + "." if assessment.start_date else ""}',
                    'assessment', assessment.id,
                )
            except Exception:
                pass

    @action(detail=True, methods=['post'])
    def complete(self, request, pk=None):
        assessment = self.get_object()
        assessment.score = assessment.calculate_score()
        assessment.status = 'COMPLETED'
        assessment.completed_at = timezone.now()
        assessment.save()
        return Response(AssessmentSerializer(assessment).data)

    @action(detail=True, methods=['post'])
    def set_baseline(self, request, pk=None):
        """Mark this assessment as the baseline for its org/template."""
        assessment = self.get_object()
        # Clear any existing baseline for same org + template
        Assessment.objects.filter(
            organization=assessment.organization,
            template=assessment.template,
            is_baseline=True,
        ).exclude(pk=assessment.pk).update(is_baseline=False)
        assessment.is_baseline = True
        assessment.save(update_fields=['is_baseline'])
        return Response(AssessmentSerializer(assessment).data)

    @action(detail=True, methods=['post'])
    def unset_baseline(self, request, pk=None):
        assessment = self.get_object()
        assessment.is_baseline = False
        assessment.save(update_fields=['is_baseline'])
        return Response(AssessmentSerializer(assessment).data)

    @action(detail=True, methods=['get'])
    def export_report(self, request, pk=None):
        """Generate and return a standalone HTML assessment report."""
        from .report_generator import generate_baseline_report, generate_comparison_report
        from .models import AssessmentTemplate as Tmpl
        assessment = self.get_object()
        template   = Tmpl.objects.prefetch_related('sections__questions').get(pk=assessment.template_id)

        # Reload with fresh prefetch for report generation
        assessment = Assessment.objects.prefetch_related(
            'responses__question', 'responses__evidence',
        ).select_related('template', 'engagement', 'assessor', 'baseline').get(pk=assessment.pk)

        compare = request.query_params.get('compare') == '1'
        if compare and assessment.baseline_id:
            baseline = Assessment.objects.prefetch_related('responses__question').get(pk=assessment.baseline_id)
            html = generate_comparison_report(assessment, baseline, template)
            fname = f'comparison_{assessment.id}.html'
        else:
            html = generate_baseline_report(assessment, template)
            fname = f'assessment_{assessment.id}.html'

        resp = HttpResponse(html, content_type='text/html; charset=utf-8')
        resp['Content-Disposition'] = f'attachment; filename="{fname}"'
        return resp

    @action(detail=True, methods=['get'])
    def export_pdf(self, request, pk=None):
        """Generate and return a PDF assessment report via WeasyPrint."""
        from .report_generator import generate_baseline_report, generate_comparison_report
        from .models import AssessmentTemplate as Tmpl
        assessment = self.get_object()
        template   = Tmpl.objects.prefetch_related('sections__questions').get(pk=assessment.template_id)

        assessment = Assessment.objects.prefetch_related(
            'responses__question', 'responses__evidence',
        ).select_related('template', 'engagement', 'assessor', 'baseline').get(pk=assessment.pk)

        compare = request.query_params.get('compare') == '1'
        if compare and assessment.baseline_id:
            baseline = Assessment.objects.prefetch_related('responses__question').get(pk=assessment.baseline_id)
            html  = generate_comparison_report(assessment, baseline, template)
            fname = f'comparison_{assessment.id}.pdf'
        else:
            html  = generate_baseline_report(assessment, template)
            fname = f'assessment_{assessment.id}.pdf'

        try:
            import weasyprint
            pdf_bytes = weasyprint.HTML(string=html).write_pdf()
        except ImportError:
            return Response({'error': 'WeasyPrint is not installed.'}, status=501)
        except Exception as e:
            import traceback
            return Response({'error': str(e), 'detail': traceback.format_exc()}, status=500)

        resp = HttpResponse(pdf_bytes, content_type='application/pdf')
        resp['Content-Disposition'] = f'attachment; filename="{fname}"'
        return resp

    @action(detail=True, methods=['post'])
    def save_responses(self, request, pk=None):
        """Bulk upsert responses for an assessment."""
        assessment = self.get_object()
        responses_data = request.data.get('responses', [])
        saved = []
        for item in responses_data:
            question_id = item.get('question')
            if not question_id:
                continue
            obj, _ = AssessmentResponse.objects.update_or_create(
                assessment=assessment,
                question_id=question_id,
                defaults={
                    'answer':           item.get('answer', ''),
                    'maturity_achieved':item.get('maturity_achieved', False),
                    'notes':            item.get('notes', ''),
                    # GRC control-level fields stored in the JSON notes column
                    # (gap_analysis/owner/due_date are encoded as JSON in notes if present)
                }
            )
            saved.append(obj)
        if assessment.status == 'DRAFT':
            assessment.status = 'IN_PROGRESS'
            assessment.save(update_fields=['status'])
        return Response(AssessmentResponseSerializer(saved, many=True).data)


    # ── GRC Evidence Submission Endpoints ────────────────────────────────────

    @action(detail=False, methods=['get'])
    def calendar_events(self, request):
        """Return assessments for calendar display.

        Returns two sets merged:
        - Dated assessments that overlap the requested date window (for grid rendering)
        - ALL unassigned assessments regardless of dates (for the pool panel)
        """
        import datetime
        from django.db.models import Q
        start_str = request.query_params.get('start')
        end_str   = request.query_params.get('end')

        base_qs = self.get_queryset().select_related('template', 'organization', 'grc_consultant')

        # Dated assessments overlapping the view window (shown in grid)
        dated_qs = base_qs.exclude(start_date__isnull=True).exclude(end_date__isnull=True)
        if start_str and end_str:
            try:
                start_d = datetime.date.fromisoformat(start_str)
                end_d   = datetime.date.fromisoformat(end_str)
                dated_qs = dated_qs.filter(
                    Q(start_date__lte=end_d) & Q(end_date__gte=start_d)
                )
            except ValueError:
                pass

        # ALL unassigned assessments (for pool display — no date filter)
        undated_unassigned_qs = base_qs.filter(grc_consultant__isnull=True)

        # Assessments where I'm the GRC consultant — ensures they always appear
        # even if dates fall outside the window or are missing
        my_assigned_qs = base_qs.filter(grc_consultant=request.user)

        seen_ids = set()
        data = []
        for a in list(dated_qs) + list(undated_unassigned_qs) + list(my_assigned_qs):
            if a.id in seen_ids:
                continue
            seen_ids.add(a.id)
            data.append({
                'id':                   a.id,
                'title':                a.title,
                'status':               a.status,
                'status_display':       a.get_status_display(),
                'start_date':           a.start_date.isoformat() if a.start_date else None,
                'end_date':             a.end_date.isoformat() if a.end_date else None,
                'framework':            a.template.framework,
                'framework_display':    a.template.get_framework_display(),
                'grc_consultant_id':    a.grc_consultant_id,
                'grc_consultant_name':  a.grc_consultant.get_full_name() if a.grc_consultant else None,
                'organization_name':    a.organization.name if a.organization else None,
            })
        return Response(data)

    @action(detail=True, methods=['get'])
    def submissions(self, request, pk=None):
        """List all evidence submissions for this assessment."""
        assessment = self.get_object()
        section_id = request.query_params.get('section')
        qs = assessment.evidence_submissions.select_related('requirement', 'submitted_by', 'reviewed_by')
        if section_id:
            qs = qs.filter(requirement__section_id=section_id)
        return Response(EvidenceSubmissionSerializer(qs, many=True, context={'request': request}).data)

    @action(detail=True, methods=['post'])
    def submit_evidence(self, request, pk=None):
        """Upload or replace a file for a specific evidence requirement."""
        from django.utils import timezone
        assessment = self.get_object()

        req_id = request.data.get('requirement')
        if not req_id:
            return Response({'error': 'requirement is required.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            requirement = SectionEvidenceRequirement.objects.get(id=req_id)
        except SectionEvidenceRequirement.DoesNotExist:
            return Response({'error': 'Requirement not found.'}, status=status.HTTP_404_NOT_FOUND)

        uploaded = request.FILES.get('file')
        sub, _ = EvidenceSubmission.objects.get_or_create(
            assessment=assessment, requirement=requirement,
            defaults={'status': 'not_started'}
        )

        if uploaded:
            if sub.file:
                try:
                    sub.file.delete(save=False)
                except Exception:
                    pass
            sub.file         = uploaded
            sub.filename     = uploaded.name
            sub.submitted_by = request.user
            sub.submitted_at = timezone.now()
            sub.status       = 'submitted'
            sub.ai_result    = None
            sub.ai_validated_at = None
            # Reset review when re-submitted
            sub.reviewer_notes = ''
            sub.reviewed_by    = None
            sub.reviewed_at    = None
        sub.save()

        # Notify GRC consultant + assessor when evidence is submitted
        if uploaded:
            try:
                from notifications.utils import notify_users
                recipients = []
                if assessment.grc_consultant:
                    recipients.append(assessment.grc_consultant)
                if assessment.assessor and assessment.assessor not in recipients:
                    recipients.append(assessment.assessor)
                if recipients:
                    submitter = request.user.get_full_name() or request.user.email
                    notify_users(
                        recipients,
                        'EVIDENCE_SUBMITTED',
                        f'Evidence submitted: {requirement.title}',
                        f'{submitter} submitted "{uploaded.name}" for the requirement '
                        f'"{requirement.title}" in assessment "{assessment.title}".',
                        'assessment', assessment.id,
                    )
            except Exception:
                pass

        return Response(EvidenceSubmissionSerializer(sub, context={'request': request}).data)

    @action(detail=True, methods=['post'], url_path=r'submissions/(?P<sub_id>[0-9]+)/review')
    def review_submission(self, request, pk=None, sub_id=None):
        """Auditor reviews a submission: accept / reject / na."""
        from django.utils import timezone
        assessment = self.get_object()
        try:
            sub = EvidenceSubmission.objects.get(id=sub_id, assessment=assessment)
        except EvidenceSubmission.DoesNotExist:
            return Response({'error': 'Submission not found.'}, status=status.HTTP_404_NOT_FOUND)

        new_status = request.data.get('status')
        valid      = {'accepted', 'rejected', 'na', 'submitted', 'ai_reviewed'}
        if new_status and new_status not in valid:
            return Response({'error': f'Invalid status. Valid: {", ".join(sorted(valid))}'}, status=status.HTTP_400_BAD_REQUEST)

        if new_status:
            sub.status = new_status
        if 'reviewer_notes' in request.data:
            sub.reviewer_notes = request.data['reviewer_notes']
        if new_status in ('accepted', 'rejected', 'na'):
            sub.reviewed_by = request.user
            sub.reviewed_at = timezone.now()

        sub.save()

        # Notify the submitter when their evidence is accepted or rejected
        if new_status in ('accepted', 'rejected') and sub.submitted_by:
            try:
                from notifications.utils import notify_users
                reviewer = request.user.get_full_name() or request.user.email
                notes_excerpt = (f' Note: "{sub.reviewer_notes[:120]}"') if sub.reviewer_notes else ''
                notify_users(
                    [sub.submitted_by],
                    'EVIDENCE_ACCEPTED' if new_status == 'accepted' else 'EVIDENCE_REJECTED',
                    f'Evidence {"accepted" if new_status == "accepted" else "rejected"}: {sub.requirement.title}',
                    f'Your evidence for "{sub.requirement.title}" in "{assessment.title}" has been '
                    f'{"accepted ✓" if new_status == "accepted" else "rejected — please resubmit"} by {reviewer}.{notes_excerpt}',
                    'assessment', assessment.id,
                )
            except Exception:
                pass

        return Response(EvidenceSubmissionSerializer(sub, context={'request': request}).data)

    @action(detail=True, methods=['post'], url_path=r'submissions/(?P<sub_id>[0-9]+)/ai_validate')
    def ai_validate_submission(self, request, pk=None, sub_id=None):
        """Run structured AI validation on a submitted document against its requirement."""
        import json, re as _re
        from django.utils import timezone

        assessment = self.get_object()
        try:
            sub = EvidenceSubmission.objects.get(id=sub_id, assessment=assessment)
        except EvidenceSubmission.DoesNotExist:
            return Response({'error': 'Submission not found.'}, status=status.HTTP_404_NOT_FOUND)

        if not sub.file:
            return Response({'error': 'No file uploaded for this submission.'}, status=status.HTTP_400_BAD_REQUEST)

        # Extract document text
        filename = (sub.filename or sub.file.name or '').lower()
        doc_text = ''
        try:
            with sub.file.open('rb') as f:
                if filename.endswith('.docx'):
                    from docx import Document
                    doc  = Document(f)
                    doc_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                elif filename.endswith('.pdf'):
                    import pypdf
                    reader   = pypdf.PdfReader(f)
                    doc_text = '\n'.join(page.extract_text() or '' for page in reader.pages)
                else:
                    doc_text = f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            doc_text = f'[Extraction failed: {e}]'

        req   = sub.requirement
        extra = f"\nAdditional validation criteria:\n{req.validation_prompt}" if req.validation_prompt else ''

        prompt = f"""You are a strict GRC compliance auditor validating a submitted document against a specific evidence requirement.

EVIDENCE REQUIREMENT:
  Title:       {req.title}
  Type:        {req.get_document_type_display()}
  Description: {req.description}{extra}

SUBMITTED DOCUMENT filename: {sub.filename or '(unknown)'}
SUBMITTED DOCUMENT content (first 7000 chars):
{doc_text[:7000]}

══════════════════════════════════════════════════
STEP 1 — DOCUMENT TYPE CHECK (do this first, it is mandatory)
══════════════════════════════════════════════════
Determine what type of document was actually submitted (e.g. Incident Response Plan, Security Awareness Training record, Network diagram, Invoice, etc.).

If the submitted document is fundamentally a DIFFERENT document type than what is required — for example:
  • An Incident Response Plan submitted for an Information Security Policy requirement
  • A training completion record submitted for a risk assessment requirement
  • A network diagram submitted for a policy document requirement
  • Any document that is clearly NOT the type specified in the requirement title/description

…then you MUST immediately return:
  "is_correct_document": false
  "status": "Wrong Document"
  "coverage_score": 0
  "gaps": ["The submitted document is a <detected type>, not a <required type>. Please upload the correct document."]
  "strengths": []

Do NOT award partial credit for the wrong document type, even if it contains related content.

══════════════════════════════════════════════════
STEP 2 — CONTENT ANALYSIS (only if document type is correct)
══════════════════════════════════════════════════
If the document IS the correct type, analyse it for completeness:
1. Does it contain the required content, policies, and procedures described above?
2. Is it sufficiently detailed and actionable?
3. Are there critical gaps?

Respond ONLY with a single valid JSON object — no markdown fences, no prose outside the JSON:
{{
  "is_correct_document": <true|false>,
  "document_type_detected": "<what this document actually is>",
  "coverage_score": <integer 0-100; must be 0 if wrong document type>,
  "status": "<Compliant|Partially Compliant|Non-Compliant|Wrong Document>",
  "summary": "<2-3 sentence assessment>",
  "gaps": ["<specific gap 1>", "<specific gap 2>"],
  "strengths": ["<confirmed strength 1>", "<confirmed strength 2>"],
  "recommendation": "<specific, actionable recommendation>"
}}"""

        try:
            raw   = _call_ai(prompt)
            match = _re.search(r'\{[\s\S]*\}', raw)
            if match:
                ai_result = json.loads(match.group())
            else:
                ai_result = {'error': 'Could not parse AI response', 'raw': raw[:500]}
        except Exception as e:
            return Response({'error': f'AI validation failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        sub.ai_result       = ai_result
        sub.ai_validated_at = timezone.now()
        # Wrong document type — keep as submitted so auditor action is still required
        if not ai_result.get('is_correct_document', True) or ai_result.get('status') == 'Wrong Document':
            sub.status = 'submitted'
        else:
            sub.status = 'ai_reviewed'
        sub.save()

        return Response(EvidenceSubmissionSerializer(sub, context={'request': request}).data)

    @action(detail=True, methods=['post'])
    def ai_audit(self, request, pk=None):
        """
        AI-powered document audit.
        Accepts multipart/form-data:
          file     — optional uploaded document (PDF / DOCX / TXT)
          doc_text — optional pre-extracted text (client cache across messages)
          message  — user chat message
          mode     — 'scan' (full structured audit) | 'chat' (conversational)
        """
        assessment = self.get_object()
        mode       = request.data.get('mode', 'chat')
        message    = request.data.get('message', '')
        doc_text   = request.data.get('doc_text', '')
        section_id = request.data.get('section_id')

        uploaded = request.FILES.get('file')
        if uploaded:
            doc_text = _extract_document_text(uploaded)

        sections = list(
            assessment.template.sections
            .prefetch_related('questions')
            .order_by('order')
        )

        # Scope analysis to a single section when section_id is provided
        if section_id:
            sections = [s for s in sections if str(s.id) == str(section_id)]

        if mode == 'scan':
            result = _ai_full_scan(sections, doc_text)
        elif message:
            result = _ai_chat(sections, doc_text, message)
        else:
            # No message — just returning extracted doc_text, no AI call needed
            result = {'type': 'chat', 'message': ''}

        # Always return the (possibly newly extracted) doc_text so the client
        # can cache it and avoid re-uploading on subsequent messages.
        result['doc_text'] = doc_text
        return Response(result)


# ── AI audit helpers ──────────────────────────────────────────────────────────

def _extract_document_text(file_obj):
    name = (file_obj.name or '').lower()
    try:
        if name.endswith('.docx'):
            from docx import Document
            doc = Document(file_obj)
            return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        elif name.endswith('.pdf'):
            import pypdf
            reader = pypdf.PdfReader(file_obj)
            return '\n'.join(page.extract_text() or '' for page in reader.pages)
        else:
            return file_obj.read().decode('utf-8', errors='ignore')
    except Exception as e:
        return f'[Document extraction failed: {e}]'


def _call_ai(prompt):
    from django.conf import settings
    provider = getattr(settings, 'AI_PROVIDER', 'ollama')
    if provider == 'gemini':
        from findings.ai_views import call_gemini
        return call_gemini(prompt)
    from findings.ai_views import call_ollama
    return call_ollama(prompt)


def _ai_full_scan(sections, doc_text):
    import json, re as _re

    controls_lines = []
    for s in sections:
        controls_lines.append(f'Section {s.id} — {s.name}:')
        for q in s.questions.all()[:6]:          # cap to keep prompt tight
            controls_lines.append(f'  • {q.text}')

    controls_str = '\n'.join(controls_lines)
    doc_preview  = doc_text[:7000] if doc_text else '(no document provided — assess based on absence of evidence)'

    prompt = f"""You are a professional security auditor conducting a compliance gap analysis.

ASSESSMENT CONTROLS:
{controls_str}

DOCUMENT UNDER REVIEW:
{doc_preview}

Review the document against each section above. For every section provide a score and finding.

Respond ONLY with a single valid JSON object — no markdown, no prose before or after:
{{
  "overall_score": <integer 0-100>,
  "summary": "<2-3 sentence executive summary of the overall compliance posture>",
  "sections": [
    {{
      "section_id": <integer>,
      "section_name": "<exact section name>",
      "score": <integer 0-100>,
      "status": "<Compliant|Partially Compliant|Non-Compliant|Not Applicable>",
      "finding": "<1-2 sentences: what evidence was or wasn't found in the document>",
      "recommendation": "<1-2 sentences: specific, actionable improvement>"
    }}
  ]
}}"""

    raw = _call_ai(prompt)

    try:
        match = _re.search(r'\{[\s\S]*\}', raw)
        if match:
            parsed = json.loads(match.group())
            return {'type': 'scan', 'result': parsed}
    except Exception:
        pass

    # Model didn't produce valid JSON — return as chat fallback
    return {'type': 'chat', 'message': raw}


def _ai_chat(sections, doc_text, message):
    controls_str = '\n'.join(f'• {s.name}' for s in sections)
    doc_context  = f'\n\nDOCUMENT CONTEXT:\n{doc_text[:6000]}' if doc_text else ''

    system = (
        'You are an expert security auditor assistant helping review a compliance assessment. '
        'Be concise, specific, and evidence-based. Do not add disclaimers.'
    )

    prompt = f"""{system}

ASSESSMENT SECTIONS:
{controls_str}
{doc_context}

USER: {message or 'Hello — how can you help me with this assessment?'}"""

    raw = _call_ai(prompt)
    return {'type': 'chat', 'message': raw}


class AssessmentResponseViewSet(viewsets.ModelViewSet):
    serializer_class   = AssessmentResponseSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        org = get_org(self.request)
        return AssessmentResponse.objects.filter(
            assessment__organization=org
        ).select_related('question', 'assessment').prefetch_related('evidence')

    def create(self, request, *args, **kwargs):
        assessment_id = request.data.get('assessment')
        question_id   = request.data.get('question')
        obj, created  = AssessmentResponse.objects.update_or_create(
            assessment_id=assessment_id,
            question_id=question_id,
            defaults={
                'answer':            request.data.get('answer', ''),
                'maturity_achieved': request.data.get('maturity_achieved', False),
                'notes':             request.data.get('notes', ''),
            }
        )
        code = status.HTTP_201_CREATED if created else status.HTTP_200_OK
        return Response(AssessmentResponseSerializer(obj).data, status=code)


class AssessmentEvidenceViewSet(viewsets.ModelViewSet):
    serializer_class   = AssessmentEvidenceSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        org = get_org(self.request)
        return AssessmentEvidence.objects.filter(
            response__assessment__organization=org
        ).select_related('uploaded_by')

    def perform_create(self, serializer):
        serializer.save(uploaded_by=self.request.user)
