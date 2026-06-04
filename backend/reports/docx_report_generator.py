"""
Word (.docx) report generator — OziCyber.
Builds a fully styled document from scratch using python-docx + raw XML injection.
"""
import io
import re
from datetime import date, datetime

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ── Brand colours ──────────────────────────────────────────────────────────────
GREEN       = '24483E'
GOLD        = 'C9A84C'
WHITE       = 'FFFFFF'
LIGHT_BG    = 'F2F5F3'
MID_BG      = 'E8EDEB'
DARK_TEXT   = '1A1A1A'
MID_TEXT    = '444444'
BORDER      = 'C8D4CF'

SEV_COLORS = {
    'CRITICAL':      ('7B0000', 'FFFFFF'),
    'HIGH':          ('B71C1C', 'FFFFFF'),
    'MEDIUM':        ('E65100', 'FFFFFF'),
    'LOW':           ('F9A825', '1A1A1A'),
    'INFORMATIONAL': ('1565C0', 'FFFFFF'),
}
SEV_ORDER = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFORMATIONAL']

CVSS_LABELS = {
    'av': {'N': 'Network', 'A': 'Adjacent', 'L': 'Local',  'P': 'Physical'},
    'ac': {'L': 'Low',     'H': 'High'},
    'pr': {'N': 'None',    'L': 'Low',     'H': 'High'},
    'ui': {'N': 'None',    'R': 'Required'},
    's':  {'U': 'Unchanged','C': 'Changed'},
    'c':  {'N': 'None',    'L': 'Low',     'H': 'High'},
    'i':  {'N': 'None',    'L': 'Low',     'H': 'High'},
    'a':  {'N': 'None',    'L': 'Low',     'H': 'High'},
}
CVSS_METRIC_NAMES = {
    'av': 'Attack Vector',
    'ac': 'Attack Complexity',
    'pr': 'Privileges Required',
    'ui': 'User Interaction',
    's':  'Scope',
    'c':  'Confidentiality',
    'i':  'Integrity',
    'a':  'Availability',
}

# Page dimensions (A4 with 2cm margins → 17cm usable width)
PAGE_W_DXA = int(17.0 * 567)   # twips (1 cm = 567 twips)
PAGE_W_EMU = int(17.0 * 914400 / 2.54)


# ══════════════════════════════════════════════════════════════════════════════
# XML / TABLE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _cell_shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.upper())
    tcPr.append(shd)


def _cell_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_dxa)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _cell_valign(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _tbl_width(tbl_elem, width_dxa):
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elem.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(width_dxa)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def _tbl_no_borders(tbl_elem):
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elem.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _tbl_cell_spacing(tbl_elem, dxa=0):
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elem.insert(0, tblPr)
    spacing = OxmlElement('w:tblCellSpacing')
    spacing.set(qn('w:w'),    str(dxa))
    spacing.set(qn('w:type'), 'dxa')
    tblPr.append(spacing)


def _set_cell_border(cell, **kwargs):
    """Set individual cell borders. kwargs: top/bottom/left/right → dict(val, sz, color)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, cfg in kwargs.items():
        if not cfg:
            continue
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   cfg.get('val',   'single'))
        el.set(qn('w:sz'),    cfg.get('sz',    '6'))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), cfg.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_para_left_border(paragraph, color_hex, sz='12', space='6'):
    """Add a coloured left border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    sz)
    left.set(qn('w:space'), space)
    left.set(qn('w:color'), color_hex.upper())
    pBdr.append(left)
    pPr.append(pBdr)


def _keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepNext')
    pPr.append(kwn)


def _add_page_number_field(paragraph):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run = paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_num_pages_field(paragraph):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' NUMPAGES '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run = paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ══════════════════════════════════════════════════════════════════════════════
# TEXT HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _plain_text(html_or_text):
    if not html_or_text:
        return ''
    s = str(html_or_text)
    if '<' in s:
        return BeautifulSoup(s, 'html.parser').get_text(separator='\n').strip()
    return s.strip()


def _fmt_date(d):
    if not d:
        return '—'
    if isinstance(d, (date, datetime)):
        return d.strftime('%d %B %Y')
    return str(d)


def _rgb(hex_color):
    h = hex_color.lstrip('#')
    return RGBColor(*bytes.fromhex(h))


# ══════════════════════════════════════════════════════════════════════════════
# LAYOUT HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)
    return p


def _section_heading(doc, text, add_spacer_before=True):
    """Full-width green band section heading with gold left accent."""
    if add_spacer_before:
        _spacer(doc, 8)

    tbl = doc.add_table(rows=1, cols=1)
    _tbl_no_borders(tbl._tbl)
    _tbl_width(tbl._tbl, PAGE_W_DXA)
    cell = tbl.rows[0].cells[0]
    _cell_shade(cell, GREEN)
    _cell_width(cell, PAGE_W_DXA)
    _cell_valign(cell, 'center')
    _set_cell_border(cell, left={'val': 'single', 'sz': '24', 'color': GOLD})

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.12)
    r = p.add_run(text.upper())
    r.bold           = True
    r.font.size      = Pt(10)
    r.font.color.rgb = _rgb(WHITE)

    _spacer(doc, 4)


def _content_section(doc, label, text, border_color=None, label_color=None):
    """Left-bordered content block: bold label line followed by body text."""
    if not text:
        return
    bc = (border_color or GREEN).lstrip('#')
    lc = (label_color  or GREEN).lstrip('#')

    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_before = Pt(5)
    label_p.paragraph_format.space_after  = Pt(1)
    label_p.paragraph_format.left_indent  = Inches(0.12)
    _add_para_left_border(label_p, bc, sz='18', space='8')
    lr = label_p.add_run(label.upper())
    lr.bold           = True
    lr.font.size      = Pt(8)
    lr.font.color.rgb = _rgb(lc)

    lines = [l for l in text.splitlines() if l.strip()]
    for line in lines:
        tp = doc.add_paragraph(line.strip())
        tp.paragraph_format.space_before = Pt(1)
        tp.paragraph_format.space_after  = Pt(1)
        tp.paragraph_format.left_indent  = Inches(0.18)
        _add_para_left_border(tp, bc, sz='6', space='8')
        for r in tp.runs:
            r.font.size = Pt(10)

    _spacer(doc, 2)


def _info_table(doc, rows_data, col_widths=None):
    """Two-column label/value table."""
    if not col_widths:
        col_widths = [int(PAGE_W_DXA * 0.32), int(PAGE_W_DXA * 0.68)]
    tbl = doc.add_table(rows=0, cols=2)
    _tbl_width(tbl._tbl, PAGE_W_DXA)
    _tbl_no_borders(tbl._tbl)
    for label, val in rows_data:
        row = tbl.add_row()
        lc, vc = row.cells[0], row.cells[1]
        _cell_shade(lc, MID_BG)
        _cell_shade(vc, LIGHT_BG)
        _cell_width(lc, col_widths[0])
        _cell_width(vc, col_widths[1])
        _set_cell_border(lc, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
        _set_cell_border(vc, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after  = Pt(3)
        lp.paragraph_format.left_indent  = Inches(0.08)
        lr = lp.add_run(label)
        lr.bold           = True
        lr.font.size      = Pt(9)
        lr.font.color.rgb = _rgb(GREEN)
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after  = Pt(3)
        vp.paragraph_format.left_indent  = Inches(0.08)
        vr = vp.add_run(str(val) if val else '—')
        vr.font.size = Pt(9)


def _sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(10)
    r.font.color.rgb = _rgb(GREEN)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# CHART HELPERS (matplotlib PNG — editable by deleting/replacing in Word)
# ══════════════════════════════════════════════════════════════════════════════

def _sev_bar_chart_png(sev_counts):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches

        labels = [s for s in SEV_ORDER if sev_counts.get(s, 0) > 0]
        if not labels:
            return None
        values = [sev_counts[s] for s in labels]
        bar_colors = {
            'CRITICAL': '#7B0000', 'HIGH': '#B71C1C', 'MEDIUM': '#E65100',
            'LOW': '#F9A825', 'INFORMATIONAL': '#1565C0',
        }
        colors = [bar_colors[s] for s in labels]

        fig, ax = plt.subplots(figsize=(5.5, max(1.8, len(labels) * 0.65)))
        fig.patch.set_facecolor('#F2F5F3')
        ax.set_facecolor('#F2F5F3')

        bars = ax.barh(labels, values, color=colors, height=0.55, edgecolor='white', linewidth=0.5)
        for bar, val in zip(bars, values):
            ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height() / 2,
                    str(val), va='center', fontsize=10, fontweight='bold', color='#1A1A1A')

        ax.set_xlim(0, max(values) * 1.35 if values else 5)
        ax.set_xlabel('Number of Findings', fontsize=9, color='#444444')
        ax.tick_params(colors='#444444', labelsize=9)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#C8D4CF')
        ax.spines['bottom'].set_color('#C8D4CF')
        ax.set_title('Findings by Severity', fontsize=10, fontweight='bold',
                     color='#24483E', pad=8)
        plt.tight_layout(pad=0.6)

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                    facecolor=fig.get_facecolor())
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _status_donut_chart_png(open_c, closed_c, in_prog_c):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        raw = [('Open', open_c, '#B71C1C'), ('In Progress', in_prog_c, '#E65100'), ('Resolved', closed_c, '#2E7D32')]
        data = [(l, v, c) for l, v, c in raw if v > 0]
        if not data:
            return None

        labels = [d[0] for d in data]
        values = [d[1] for d in data]
        colors = [d[2] for d in data]

        fig, ax = plt.subplots(figsize=(3.8, 3.2))
        fig.patch.set_facecolor('#F2F5F3')
        ax.set_facecolor('#F2F5F3')

        wedges, _ = ax.pie(values, colors=colors, wedgeprops={'width': 0.52, 'edgecolor': 'white', 'linewidth': 1.5},
                           startangle=90)
        total = sum(values)
        ax.text(0, 0, str(total), ha='center', va='center', fontsize=16, fontweight='bold', color='#24483E')
        ax.text(0, -0.22, 'Total', ha='center', va='center', fontsize=8, color='#444444')
        ax.legend(wedges, [f'{l}  ({v})' for l, v in zip(labels, values)],
                  loc='lower center', bbox_to_anchor=(0.5, -0.18),
                  ncol=len(data), fontsize=8, frameon=False)
        ax.set_title('Remediation Status', fontsize=10, fontweight='bold', color='#24483E', pad=4)
        plt.tight_layout(pad=0.4)

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                    facecolor=fig.get_facecolor())
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

def _build_cover_page(doc, report):
    eng      = report.engagement
    org      = getattr(eng, 'organization', None)
    org_name = getattr(org, 'name', 'Client Organisation') if org else 'Client Organisation'
    label    = 'DRAFT' if getattr(report, 'is_draft', True) else 'FINAL'

    # ── Outer green panel ──────────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=1)
    _tbl_no_borders(tbl._tbl)
    _tbl_width(tbl._tbl, PAGE_W_DXA)
    cell = tbl.rows[0].cells[0]
    _cell_shade(cell, GREEN)
    _cell_width(cell, PAGE_W_DXA)

    def _cp(text, size=11, bold=False, color=WHITE, space_before=8, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.2)
        r = p.add_run(text)
        r.bold           = bold
        r.italic         = italic
        r.font.size      = Pt(size)
        r.font.color.rgb = _rgb(color)
        return p

    # Gold top rule
    top_rule = cell.paragraphs[0]
    top_rule.paragraph_format.space_before = Pt(0)
    top_rule.paragraph_format.space_after  = Pt(0)
    _add_para_left_border(top_rule, GOLD, sz='0')
    tr = top_rule.add_run('')
    # Actually use a table border for the top gold line
    _set_cell_border(cell, top={'val': 'single', 'sz': '24', 'color': GOLD})

    _cp('OziCyber',                          size=13,  bold=True,  color=GOLD,  space_before=20)
    _cp('Security & Advisory',               size=10,  bold=False, color=GOLD,  space_before=2)

    # Gold divider line via paragraph
    div = cell.add_paragraph()
    div.paragraph_format.space_before = Pt(14)
    div.paragraph_format.space_after  = Pt(0)
    div.paragraph_format.left_indent  = Inches(0.2)
    _add_para_left_border(div, GOLD, sz='0')

    _cp(report.title or 'Penetration Test Report', size=26, bold=True,  color=WHITE, space_before=18)
    _cp(f'Penetration Testing Report  ·  {label}',    size=12, bold=False, color=GOLD,  space_before=8)

    _cp('',  size=6, space_before=22)  # spacer

    # Client details block
    for lbl, val in [
        ('Client',      org_name),
        ('Engagement',  getattr(eng, 'name', '—')),
        ('Version',     f'v{report.version}'),
        ('Date',        _fmt_date(getattr(report, 'created_at', date.today()))),
        ('Classification', 'CONFIDENTIAL'),
    ]:
        row_p = cell.add_paragraph()
        row_p.paragraph_format.space_before = Pt(3)
        row_p.paragraph_format.space_after  = Pt(0)
        row_p.paragraph_format.left_indent  = Inches(0.2)
        lb = row_p.add_run(f'{lbl}:  ')
        lb.bold           = True
        lb.font.size      = Pt(10)
        lb.font.color.rgb = _rgb(GOLD)
        vb = row_p.add_run(val)
        vb.font.size      = Pt(10)
        vb.font.color.rgb = _rgb(WHITE)

    _cp('', size=24, space_before=0)  # bottom padding

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTROL
# ══════════════════════════════════════════════════════════════════════════════

def _build_doc_control(doc, report):
    _section_heading(doc, '1.  Document Control', add_spacer_before=False)

    # Classification banner
    tbl = doc.add_table(rows=1, cols=1)
    _tbl_no_borders(tbl._tbl)
    _tbl_width(tbl._tbl, PAGE_W_DXA)
    cell = tbl.rows[0].cells[0]
    _cell_shade(cell, MID_BG)
    _cell_width(cell, PAGE_W_DXA)
    _set_cell_border(cell, left={'val': 'single', 'sz': '18', 'color': GOLD})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.12)
    r = p.add_run('CLASSIFICATION: CONFIDENTIAL — For Authorised Recipients Only')
    r.bold           = True
    r.font.size      = Pt(9)
    r.font.color.rgb = _rgb(GREEN)

    _spacer(doc, 8)

    _sub_heading(doc, 'Version History')

    hdrs   = ['Version', 'Date', 'Author', 'Change Description']
    widths = [int(PAGE_W_DXA * 0.12), int(PAGE_W_DXA * 0.20),
              int(PAGE_W_DXA * 0.28), int(PAGE_W_DXA * 0.40)]

    vtbl = doc.add_table(rows=1, cols=4)
    _tbl_width(vtbl._tbl, PAGE_W_DXA)
    _tbl_no_borders(vtbl._tbl)

    hrow = vtbl.rows[0]
    for i, (h, w) in enumerate(zip(hdrs, widths)):
        c = hrow.cells[i]
        _cell_shade(c, GREEN)
        _cell_width(c, w)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.06)
        r = p.add_run(h)
        r.bold           = True
        r.font.size      = Pt(9)
        r.font.color.rgb = _rgb(WHITE)

    author = getattr(report, 'created_by', None)
    a_name = ''
    if author:
        a_name = f'{getattr(author, "first_name", "")} {getattr(author, "last_name", "")}'.strip()
        a_name = a_name or getattr(author, 'email', 'OziCyber')
    a_name = a_name or 'OziCyber'

    drow = vtbl.add_row()
    row_vals = [f'v{report.version}',
                _fmt_date(getattr(report, 'created_at', date.today())),
                a_name,
                'DRAFT — Initial Release' if getattr(report, 'is_draft', True) else 'Final Release']
    for i, (val, w) in enumerate(zip(row_vals, widths)):
        c = drow.cells[i]
        _cell_shade(c, LIGHT_BG)
        _cell_width(c, w)
        _set_cell_border(c, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Inches(0.06)
        p.add_run(val).font.size = Pt(9)

    _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

def _build_exec_summary(doc, report, findings, sev_counts):
    _section_heading(doc, '2.  Executive Summary')

    summary_text = _plain_text(getattr(report, 'executive_summary', '') or '')
    if summary_text:
        p = doc.add_paragraph(summary_text)
        p.paragraph_format.space_after = Pt(10)
        for r in p.runs:
            r.font.size = Pt(10)

    # ── Severity tiles ─────────────────────────────────────────────────────────
    total     = len(findings)
    tile_data = [('TOTAL', str(total), GREEN, WHITE)] + [
        (s[:4] if s != 'INFORMATIONAL' else 'INFO',
         str(sev_counts.get(s, 0)),
         SEV_COLORS[s][0], SEV_COLORS[s][1])
        for s in SEV_ORDER
    ]
    tile_w = PAGE_W_DXA // 6

    ttbl = doc.add_table(rows=2, cols=6)
    _tbl_no_borders(ttbl._tbl)
    _tbl_width(ttbl._tbl, PAGE_W_DXA)
    _tbl_cell_spacing(ttbl._tbl, 20)

    for col_i, (label, count, bg, fg) in enumerate(tile_data):
        num_c   = ttbl.rows[0].cells[col_i]
        label_c = ttbl.rows[1].cells[col_i]
        for c in (num_c, label_c):
            _cell_shade(c, bg)
            _cell_width(c, tile_w)
            _cell_valign(c, 'center')

        np = num_c.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np.paragraph_format.space_before = Pt(8)
        np.paragraph_format.space_after  = Pt(0)
        nr = np.add_run(count)
        nr.bold           = True
        nr.font.size      = Pt(22)
        nr.font.color.rgb = _rgb(fg)

        lp = label_c.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(8)
        lr = lp.add_run(label)
        lr.bold           = True
        lr.font.size      = Pt(7)
        lr.font.color.rgb = _rgb(fg)

    _spacer(doc, 8)

    # ── Charts side by side ────────────────────────────────────────────────────
    open_c   = sum(1 for f in findings if getattr(f, 'status', '') in ('OPEN', 'DRAFT', 'IN_REVIEW'))
    closed_c = sum(1 for f in findings if getattr(f, 'status', '') in ('REMEDIATED', 'FALSE_POSITIVE', 'ACCEPTED_RISK'))
    in_prog  = max(0, total - open_c - closed_c)

    bar_png = _sev_bar_chart_png(sev_counts)
    pie_png = _status_donut_chart_png(open_c, closed_c, in_prog)

    if bar_png or pie_png:
        ctbl = doc.add_table(rows=1, cols=2)
        _tbl_no_borders(ctbl._tbl)
        _tbl_width(ctbl._tbl, PAGE_W_DXA)
        half = PAGE_W_DXA // 2
        _cell_width(ctbl.rows[0].cells[0], half)
        _cell_width(ctbl.rows[0].cells[1], half)

        if bar_png:
            cp = ctbl.rows[0].cells[0].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.add_run().add_picture(bar_png, width=Inches(3.2))

        if pie_png:
            pp = ctbl.rows[0].cells[1].paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.add_run().add_picture(pie_png, width=Inches(2.8))

    _spacer(doc, 8)

    # ── Key findings list ──────────────────────────────────────────────────────
    key_f = [f for f in findings if getattr(f, 'is_key_finding', False)]
    if not key_f:
        key_f = sorted(
            [f for f in findings if getattr(f, 'severity', '') in ('CRITICAL', 'HIGH')],
            key=lambda f: SEV_ORDER.index(f.severity)
        )[:5]

    if key_f:
        _sub_heading(doc, 'Key Findings')
        for f in key_f:
            sev    = getattr(f, 'severity', 'INFORMATIONAL')
            bg, fg = SEV_COLORS.get(sev, ('888888', 'FFFFFF'))
            kp     = doc.add_paragraph()
            kp.paragraph_format.space_before = Pt(3)
            kp.paragraph_format.space_after  = Pt(3)
            kp.paragraph_format.left_indent  = Inches(0.12)
            _add_para_left_border(kp, bg, sz='12')
            sr = kp.add_run(f'  {sev}  ')
            sr.bold           = True
            sr.font.size      = Pt(8)
            sr.font.color.rgb = _rgb(bg)
            tr = kp.add_run(f'  {f.title}')
            tr.font.size = Pt(10)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# ENGAGEMENT DETAILS
# ══════════════════════════════════════════════════════════════════════════════

def _build_engagement_details(doc, report):
    _section_heading(doc, '3.  Engagement Details')
    eng = report.engagement
    org = getattr(eng, 'organization', None)
    pm  = getattr(eng, 'project_manager', None)
    lead = getattr(eng, 'lead_pentester', None)

    def _name(u):
        if not u:
            return '—'
        n = f'{getattr(u,"first_name","")} {getattr(u,"last_name","")}'.strip()
        return n or getattr(u, 'email', '—')

    _sub_heading(doc, 'Project Information')
    _info_table(doc, [
        ('Project Name',   getattr(eng, 'name', '—')),
        ('Report Title',   report.title or '—'),
        ('Version',        f'v{report.version}'),
        ('Status',         'DRAFT' if getattr(report, 'is_draft', True) else 'FINAL'),
        ('Start Date',     _fmt_date(getattr(eng, 'start_date', None))),
        ('End Date',       _fmt_date(getattr(eng, 'end_date', None))),
        ('Report Date',    _fmt_date(getattr(report, 'created_at', date.today()))),
        ('Engagement Type', getattr(eng, 'get_engagement_type_display', lambda: getattr(eng, 'engagement_type', '—'))()),
    ])
    _spacer(doc, 8)

    _sub_heading(doc, 'Testing Team')
    _info_table(doc, [
        ('Project Manager', _name(pm)),
        ('Lead Pentester',  _name(lead)),
    ])
    _spacer(doc, 8)

    _sub_heading(doc, 'Client Information')
    _info_table(doc, [
        ('Organisation', getattr(org, 'name', '—') if org else '—'),
        ('Website',      getattr(org, 'website', '—') if org else '—'),
        ('Contact',      _plain_text(getattr(report, 'client_contact', '') or '') or '—'),
    ])
    _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# SCOPE
# ══════════════════════════════════════════════════════════════════════════════

def _build_scope(doc, report):
    _section_heading(doc, '4.  Scope of Assessment')

    scope_text = _plain_text(getattr(report, 'scope', '') or '')
    if scope_text:
        _sub_heading(doc, 'In Scope')
        for line in scope_text.splitlines():
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.2)
            p.add_run(line).font.size = Pt(10)
    else:
        p = doc.add_paragraph('Scope not defined in this report.')
        p.runs[0].font.size = Pt(10)

    oos = _plain_text(getattr(report, 'out_of_scope', '') or '')
    if oos:
        _sub_heading(doc, 'Out of Scope')
        for line in oos.splitlines():
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.2)
            p.add_run(line).font.size = Pt(10)

    limitations = _plain_text(getattr(report, 'limitations', '') or '')
    if limitations:
        _sub_heading(doc, 'Limitations & Exclusions')
        for line in limitations.splitlines():
            line = line.strip()
            if line:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after  = Pt(2)
                p.add_run(line).font.size = Pt(10)

    _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# FINDING SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

def _build_finding_summary(doc, report, findings, sev_counts):
    _section_heading(doc, '5.  Finding Summary')

    total = len(findings)
    hdrs   = ['Severity', 'Count', '% of Total', 'Risk Exposure']
    widths = [int(PAGE_W_DXA * 0.25), int(PAGE_W_DXA * 0.12),
              int(PAGE_W_DXA * 0.15), int(PAGE_W_DXA * 0.48)]

    stbl = doc.add_table(rows=1, cols=4)
    _tbl_width(stbl._tbl, PAGE_W_DXA)
    _tbl_no_borders(stbl._tbl)

    hrow = stbl.rows[0]
    for i, (h, w) in enumerate(zip(hdrs, widths)):
        c = hrow.cells[i]
        _cell_shade(c, GREEN)
        _cell_width(c, w)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.06)
        r = p.add_run(h)
        r.bold           = True
        r.font.size      = Pt(9)
        r.font.color.rgb = _rgb(WHITE)

    risk_desc = {
        'CRITICAL':      'Immediate remediation required — active exploitation risk',
        'HIGH':          'Urgent remediation recommended within 1–2 weeks',
        'MEDIUM':        'Remediate within 30 days as part of normal patching cycle',
        'LOW':           'Address in next scheduled maintenance window',
        'INFORMATIONAL': 'No immediate action required — awareness only',
    }
    for row_i, sev in enumerate(SEV_ORDER):
        cnt     = sev_counts.get(sev, 0)
        pct     = f'{int(cnt / total * 100)}%' if total else '0%'
        bg, fg  = SEV_COLORS[sev]
        shade   = LIGHT_BG if row_i % 2 == 0 else WHITE
        row     = stbl.add_row()
        vals    = [sev, str(cnt), pct, risk_desc.get(sev, '')]
        for i, (val, w) in enumerate(zip(vals, widths)):
            c = row.cells[i]
            _cell_width(c, w)
            _cell_shade(c, bg if i == 0 else shade)
            _set_cell_border(c, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Inches(0.06)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if i == 0:
                r.bold           = True
                r.font.color.rgb = _rgb(fg)

    _spacer(doc, 8)

    bar_png = _sev_bar_chart_png(sev_counts)
    if bar_png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(bar_png, width=Inches(5.5))

    _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# KEY FINDINGS TABLE
# ══════════════════════════════════════════════════════════════════════════════

def _build_key_findings(doc, report, findings, finding_codes):
    _section_heading(doc, '6.  Key Findings')

    hdrs   = ['ID', 'Finding Title', 'Severity', 'Status', 'CVSS']
    widths = [int(PAGE_W_DXA * 0.08), int(PAGE_W_DXA * 0.44),
              int(PAGE_W_DXA * 0.16), int(PAGE_W_DXA * 0.18),
              int(PAGE_W_DXA * 0.14)]

    ktbl = doc.add_table(rows=1, cols=5)
    _tbl_width(ktbl._tbl, PAGE_W_DXA)
    _tbl_no_borders(ktbl._tbl)

    hrow = ktbl.rows[0]
    for i, (h, w) in enumerate(zip(hdrs, widths)):
        c = hrow.cells[i]
        _cell_shade(c, GREEN)
        _cell_width(c, w)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.06)
        r = p.add_run(h)
        r.bold           = True
        r.font.size      = Pt(9)
        r.font.color.rgb = _rgb(WHITE)

    sorted_f = sorted(findings,
                      key=lambda f: SEV_ORDER.index(getattr(f, 'severity', 'INFORMATIONAL'))
                      if getattr(f, 'severity', '') in SEV_ORDER else 99)

    for row_i, f in enumerate(sorted_f):
        sev    = getattr(f, 'severity', 'INFORMATIONAL')
        status = getattr(f, 'status', 'OPEN')
        cvss   = getattr(f, 'cvss_score', None)
        bg, fg = SEV_COLORS.get(sev, ('888888', 'FFFFFF'))
        shade  = LIGHT_BG if row_i % 2 == 0 else WHITE
        row    = ktbl.add_row()
        vals   = [finding_codes.get(f.id, '—'), f.title, sev, status,
                  str(cvss) if cvss else '—']
        for i, (val, w) in enumerate(zip(vals, widths)):
            c = row.cells[i]
            _cell_width(c, w)
            _cell_shade(c, bg if i == 2 else shade)
            _set_cell_border(c, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Inches(0.06)
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if i == 2:
                r.bold           = True
                r.font.color.rgb = _rgb(fg)

    _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════

def _build_methodology(doc, report):
    _section_heading(doc, '7.  Methodology')

    methodology = _plain_text(getattr(report, 'methodology', '') or '')
    phases = methodology.splitlines() if methodology else [
        'Reconnaissance — OSINT gathering, passive enumeration, and attack surface mapping.',
        'Scanning & Enumeration — Active port scanning, service fingerprinting, and asset discovery.',
        'Vulnerability Analysis — Identification of weaknesses in configurations, services, and applications.',
        'Exploitation — Controlled, safe exploitation of confirmed vulnerabilities with written authorisation.',
        'Post-Exploitation — Privilege escalation, lateral movement, and data-access impact assessment.',
        'Reporting — Documentation of all findings with risk ratings, evidence, and remediation guidance.',
    ]

    for phase in phases:
        phase = phase.strip()
        if not phase:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Inches(0.15)
        _add_para_left_border(p, GOLD, sz='8')
        idx = phase.find('—')
        if idx > 0:
            r1 = p.add_run(phase[:idx + 1])
            r1.bold           = True
            r1.font.size      = Pt(10)
            r1.font.color.rgb = _rgb(GREEN)
            r2 = p.add_run(phase[idx + 1:])
            r2.font.size = Pt(10)
        else:
            r = p.add_run(phase)
            r.font.size = Pt(10)

    _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# DETAILED FINDINGS  (the core of the report)
# ══════════════════════════════════════════════════════════════════════════════

def _build_detailed_findings(doc, report, findings, finding_codes):
    _section_heading(doc, '8.  Detailed Findings')

    sorted_f = sorted(findings,
                      key=lambda f: SEV_ORDER.index(getattr(f, 'severity', 'INFORMATIONAL'))
                      if getattr(f, 'severity', '') in SEV_ORDER else 99)

    for f_idx, f in enumerate(sorted_f):
        sev    = getattr(f, 'severity', 'INFORMATIONAL')
        bg, fg = SEV_COLORS.get(sev, ('888888', 'FFFFFF'))
        code   = finding_codes.get(f.id, f'F-{f_idx + 1:02d}')
        status = getattr(f, 'status', 'OPEN')
        cvss   = getattr(f, 'cvss_score', None)

        # ── Finding header banner ──────────────────────────────────────────────
        # Two columns: severity badge (narrow) | code + title (wide)
        hdr_tbl = doc.add_table(rows=1, cols=2)
        _tbl_no_borders(hdr_tbl._tbl)
        _tbl_width(hdr_tbl._tbl, PAGE_W_DXA)
        badge_w = int(PAGE_W_DXA * 0.14)
        title_w = PAGE_W_DXA - badge_w

        badge_c = hdr_tbl.rows[0].cells[0]
        title_c = hdr_tbl.rows[0].cells[1]

        _cell_shade(badge_c, bg)
        _cell_width(badge_c, badge_w)
        _cell_valign(badge_c, 'center')
        bp = badge_c.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(6)
        bp.paragraph_format.space_after  = Pt(6)
        for char in sev:
            br = bp.add_run(char + '\n')
            br.bold           = True
            br.font.size      = Pt(7)
            br.font.color.rgb = _rgb(fg)

        _cell_shade(title_c, bg)
        _cell_width(title_c, title_w)
        _cell_valign(title_c, 'center')
        _set_cell_border(title_c, left={'val': 'single', 'sz': '6', 'color': fg})
        tp = title_c.paragraphs[0]
        tp.paragraph_format.space_before = Pt(7)
        tp.paragraph_format.space_after  = Pt(7)
        tp.paragraph_format.left_indent  = Inches(0.1)
        code_r = tp.add_run(f'{code}   ')
        code_r.bold           = True
        code_r.font.size      = Pt(8)
        code_r.font.color.rgb = _rgb(fg)
        title_r = tp.add_run(f.title)
        title_r.bold           = True
        title_r.font.size      = Pt(12)
        title_r.font.color.rgb = _rgb(fg)

        # ── Metadata strip ─────────────────────────────────────────────────────
        meta_items = [
            ('Status',       status.replace('_', ' ')),
            ('CVSS Score',   str(cvss) if cvss else '—'),
            ('Affected Asset', (_plain_text(getattr(f, 'affected_asset', '') or '') or '—')[:40]),
            ('Type',         (getattr(f, 'pentest_type', '') or '—').replace('_', ' ')),
        ]
        meta_w = PAGE_W_DXA // 4
        meta_tbl = doc.add_table(rows=1, cols=4)
        _tbl_no_borders(meta_tbl._tbl)
        _tbl_width(meta_tbl._tbl, PAGE_W_DXA)

        for i, (mlabel, mval) in enumerate(meta_items):
            mc = meta_tbl.rows[0].cells[i]
            _cell_shade(mc, MID_BG)
            _cell_width(mc, meta_w)
            _set_cell_border(mc,
                bottom={'val': 'single', 'sz': '6', 'color': GREEN},
                right ={'val': 'single', 'sz': '2', 'color': BORDER} if i < 3 else {})
            mp = mc.paragraphs[0]
            mp.paragraph_format.space_before = Pt(4)
            mp.paragraph_format.space_after  = Pt(4)
            mp.paragraph_format.left_indent  = Inches(0.07)
            ml = mp.add_run(mlabel + '\n')
            ml.font.size      = Pt(7)
            ml.font.color.rgb = _rgb(MID_TEXT)
            mv = mp.add_run(mval)
            mv.bold           = True
            mv.font.size      = Pt(9)
            mv.font.color.rgb = _rgb(DARK_TEXT)

        # ── CWE / CVE row ──────────────────────────────────────────────────────
        cwe = getattr(f, 'cwe_id', '') or ''
        cve = getattr(f, 'cve_id', '') or ''
        if cwe or cve:
            ref_tbl = doc.add_table(rows=1, cols=2)
            _tbl_no_borders(ref_tbl._tbl)
            _tbl_width(ref_tbl._tbl, PAGE_W_DXA)
            half = PAGE_W_DXA // 2
            for i, (rlabel, rval) in enumerate([('CWE', cwe or '—'), ('CVE', cve or '—')]):
                rc = ref_tbl.rows[0].cells[i]
                _cell_shade(rc, LIGHT_BG)
                _cell_width(rc, half)
                _set_cell_border(rc, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
                rp = rc.paragraphs[0]
                rp.paragraph_format.space_before = Pt(3)
                rp.paragraph_format.space_after  = Pt(3)
                rp.paragraph_format.left_indent  = Inches(0.07)
                rl = rp.add_run(rlabel + ':  ')
                rl.bold           = True
                rl.font.size      = Pt(8)
                rl.font.color.rgb = _rgb(GREEN)
                rv = rp.add_run(rval)
                rv.font.size = Pt(8)

        # ── Content sections ───────────────────────────────────────────────────
        _content_section(doc,
                         'Description',
                         _plain_text(getattr(f, 'description', '') or ''),
                         border_color=GREEN, label_color=GREEN)

        _content_section(doc,
                         'Technical Details',
                         _plain_text(getattr(f, 'details', '') or ''),
                         border_color=GREEN, label_color=GREEN)

        _content_section(doc,
                         'Impact',
                         _plain_text(getattr(f, 'impact', '') or ''),
                         border_color='B71C1C', label_color='B71C1C')

        _content_section(doc,
                         'Likelihood',
                         _plain_text(getattr(f, 'likelihood', '') or ''),
                         border_color='E65100', label_color='E65100')

        _content_section(doc,
                         'Recommendations',
                         _plain_text(getattr(f, 'recommendations', '') or ''),
                         border_color='2E7D32', label_color='2E7D32')

        _content_section(doc,
                         'Supporting Evidence',
                         _plain_text(getattr(f, 'supporting_evidence', '') or ''),
                         border_color=GOLD, label_color='8B6914')

        # ── CVSS breakdown ─────────────────────────────────────────────────────
        cvss_fields = ['av', 'ac', 'pr', 'ui', 's', 'c', 'i', 'a']
        cvss_vals   = {k: getattr(f, k, '') for k in cvss_fields}
        has_cvss    = any(cvss_vals.values())
        if has_cvss or cvss:
            cvss_label = doc.add_paragraph()
            cvss_label.paragraph_format.space_before = Pt(5)
            cvss_label.paragraph_format.space_after  = Pt(2)
            cvss_label.paragraph_format.left_indent  = Inches(0.12)
            _add_para_left_border(cvss_label, GREEN, sz='18', space='8')
            cl = cvss_label.add_run('CVSS 3.1 BREAKDOWN')
            cl.bold           = True
            cl.font.size      = Pt(8)
            cl.font.color.rgb = _rgb(GREEN)
            if cvss:
                cs = cvss_label.add_run(f'   Score: {cvss}')
                cs.bold           = True
                cs.font.size      = Pt(8)
                cs.font.color.rgb = _rgb(GREEN)
            if getattr(f, 'cvss_vector', ''):
                cv = cvss_label.add_run(f'   |   {f.cvss_vector}')
                cv.font.size      = Pt(7)
                cv.font.color.rgb = _rgb(MID_TEXT)

            if has_cvss:
                cvss_tbl = doc.add_table(rows=0, cols=4)
                _tbl_no_borders(cvss_tbl._tbl)
                _tbl_width(cvss_tbl._tbl, PAGE_W_DXA)
                metric_w = PAGE_W_DXA // 4

                pairs = list(zip(cvss_fields[::2], cvss_fields[1::2]))
                for pair in pairs:
                    row = cvss_tbl.add_row()
                    for col_i, key in enumerate(pair):
                        val    = cvss_vals.get(key, '')
                        label  = CVSS_METRIC_NAMES.get(key, key)
                        full   = CVSS_LABELS.get(key, {}).get(val, val)
                        lc     = row.cells[col_i * 2]
                        vc     = row.cells[col_i * 2 + 1]
                        shade  = LIGHT_BG if (pairs.index(pair) % 2 == 0) else WHITE
                        _cell_shade(lc, MID_BG)
                        _cell_shade(vc, shade)
                        _cell_width(lc, int(metric_w * 0.7))
                        _cell_width(vc, int(metric_w * 1.3))
                        _set_cell_border(lc, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
                        _set_cell_border(vc, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
                        lp = lc.paragraphs[0]
                        lp.paragraph_format.space_before = Pt(2)
                        lp.paragraph_format.space_after  = Pt(2)
                        lp.paragraph_format.left_indent  = Inches(0.06)
                        lr = lp.add_run(label)
                        lr.bold           = True
                        lr.font.size      = Pt(8)
                        lr.font.color.rgb = _rgb(GREEN)
                        vp = vc.paragraphs[0]
                        vp.paragraph_format.space_before = Pt(2)
                        vp.paragraph_format.space_after  = Pt(2)
                        vp.paragraph_format.left_indent  = Inches(0.06)
                        vr = vp.add_run(full or '—')
                        vr.font.size = Pt(8)

        # ── References ─────────────────────────────────────────────────────────
        refs = _plain_text(getattr(f, 'references', '') or '')
        if refs:
            rp_label = doc.add_paragraph()
            rp_label.paragraph_format.space_before = Pt(5)
            rp_label.paragraph_format.space_after  = Pt(2)
            rp_label.paragraph_format.left_indent  = Inches(0.12)
            _add_para_left_border(rp_label, '888888', sz='12', space='8')
            rl = rp_label.add_run('REFERENCES')
            rl.bold           = True
            rl.font.size      = Pt(8)
            rl.font.color.rgb = _rgb('888888')
            for ref in refs.splitlines():
                ref = ref.strip()
                if ref:
                    rp = doc.add_paragraph(ref)
                    rp.paragraph_format.space_before = Pt(1)
                    rp.paragraph_format.space_after  = Pt(1)
                    rp.paragraph_format.left_indent  = Inches(0.18)
                    _add_para_left_border(rp, '888888', sz='4', space='8')
                    for r in rp.runs:
                        r.font.size = Pt(8)

        # ── Separator between findings ──────────────────────────────────────────
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(10)
        sep.paragraph_format.space_after  = Pt(4)
        _add_para_left_border(sep, BORDER, sz='0')
        pPr  = sep._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), BORDER)
        pBdr.append(bot)
        pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# REMEDIATION TRACKER
# ══════════════════════════════════════════════════════════════════════════════

def _build_remediation_tracker(doc, report, findings, finding_codes):
    _section_heading(doc, '9.  Remediation Tracker')

    sorted_f   = sorted(findings,
                        key=lambda f: SEV_ORDER.index(getattr(f, 'severity', 'INFORMATIONAL'))
                        if getattr(f, 'severity', '') in SEV_ORDER else 99)
    total      = len(sorted_f)
    open_c     = sum(1 for f in sorted_f if getattr(f, 'status', '') in ('OPEN', 'DRAFT', 'IN_REVIEW'))
    closed_c   = sum(1 for f in sorted_f if getattr(f, 'status', '') in ('REMEDIATED', 'FALSE_POSITIVE', 'ACCEPTED_RISK'))
    in_prog    = max(0, total - open_c - closed_c)

    # Status tiles
    stats_tbl = doc.add_table(rows=1, cols=4)
    _tbl_no_borders(stats_tbl._tbl)
    _tbl_width(stats_tbl._tbl, PAGE_W_DXA)
    tile_w = PAGE_W_DXA // 4

    for i, (label, count, bg) in enumerate([
        ('TOTAL',       str(total),    GREEN),
        ('OPEN',        str(open_c),   'B71C1C'),
        ('IN PROGRESS', str(in_prog),  'E65100'),
        ('RESOLVED',    str(closed_c), '2E7D32'),
    ]):
        c = stats_tbl.rows[0].cells[i]
        _cell_shade(c, bg)
        _cell_width(c, tile_w)
        _cell_valign(c, 'center')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(8)
        nr = p.add_run(count + '\n')
        nr.bold           = True
        nr.font.size      = Pt(18)
        nr.font.color.rgb = _rgb(WHITE)
        lr = p.add_run(label)
        lr.bold           = True
        lr.font.size      = Pt(7)
        lr.font.color.rgb = _rgb(WHITE)

    _spacer(doc, 8)

    pie_png = _status_donut_chart_png(open_c, closed_c, in_prog)
    if pie_png:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.add_run().add_picture(pie_png, width=Inches(3.0))
        _spacer(doc, 6)

    hdrs   = ['ID', 'Finding Title', 'Severity', 'Status', 'CVSS', 'Remediated Date']
    widths = [int(PAGE_W_DXA * 0.08), int(PAGE_W_DXA * 0.36),
              int(PAGE_W_DXA * 0.14), int(PAGE_W_DXA * 0.16),
              int(PAGE_W_DXA * 0.10), int(PAGE_W_DXA * 0.16)]

    tbl = doc.add_table(rows=1, cols=6)
    _tbl_no_borders(tbl._tbl)
    _tbl_width(tbl._tbl, PAGE_W_DXA)

    hrow = tbl.rows[0]
    for i, (h, w) in enumerate(zip(hdrs, widths)):
        c = hrow.cells[i]
        _cell_shade(c, GREEN)
        _cell_width(c, w)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.05)
        r = p.add_run(h)
        r.bold           = True
        r.font.size      = Pt(8)
        r.font.color.rgb = _rgb(WHITE)

    for row_i, f in enumerate(sorted_f):
        sev    = getattr(f, 'severity', 'INFORMATIONAL')
        status = getattr(f, 'status', 'OPEN').replace('_', ' ')
        cvss   = getattr(f, 'cvss_score', None)
        rem_at = _fmt_date(getattr(f, 'remediated_at', None))
        bg, fg = SEV_COLORS.get(sev, ('888888', 'FFFFFF'))
        shade  = LIGHT_BG if row_i % 2 == 0 else WHITE
        row    = tbl.add_row()
        vals   = [finding_codes.get(f.id, '—'), f.title, sev, status,
                  str(cvss) if cvss else '—', rem_at]
        for i, (val, w) in enumerate(zip(vals, widths)):
            c = row.cells[i]
            _cell_width(c, w)
            _cell_shade(c, bg if i == 2 else shade)
            _set_cell_border(c, bottom={'val': 'single', 'sz': '2', 'color': BORDER})
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.05)
            r = p.add_run(str(val))
            r.font.size = Pt(8)
            if i == 2:
                r.bold           = True
                r.font.color.rgb = _rgb(fg)

    _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# ATTACK CHAIN
# ══════════════════════════════════════════════════════════════════════════════

def _build_attack_chain(doc, report, attack_chain_entries):
    if not attack_chain_entries:
        return

    _section_heading(doc, '10.  Attack Chain')

    p = doc.add_paragraph(
        'The following maps the observed attack path to MITRE ATT&CK phases, '
        'illustrating how vulnerabilities were chained during the engagement.'
    )
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.font.size = Pt(10)

    phases = {}
    for entry in attack_chain_entries:
        ph = getattr(entry, 'phase', 'Unknown')
        phases.setdefault(ph, []).append(entry)

    for phase_name, entries in phases.items():
        # Phase header
        ph_tbl = doc.add_table(rows=1, cols=1)
        _tbl_no_borders(ph_tbl._tbl)
        _tbl_width(ph_tbl._tbl, PAGE_W_DXA)
        ph_c = ph_tbl.rows[0].cells[0]
        _cell_shade(ph_c, MID_BG)
        _cell_width(ph_c, PAGE_W_DXA)
        _set_cell_border(ph_c, left={'val': 'single', 'sz': '18', 'color': GOLD})
        pp = ph_c.paragraphs[0]
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after  = Pt(4)
        pp.paragraph_format.left_indent  = Inches(0.1)
        pr = pp.add_run(phase_name.upper())
        pr.bold           = True
        pr.font.size      = Pt(9)
        pr.font.color.rgb = _rgb(GREEN)

        for entry in entries:
            f      = getattr(entry, 'finding', None)
            sev    = getattr(f, 'severity', 'INFORMATIONAL') if f else 'INFORMATIONAL'
            bg, fg = SEV_COLORS.get(sev, ('888888', 'FFFFFF'))
            ep     = doc.add_paragraph()
            ep.paragraph_format.space_before = Pt(3)
            ep.paragraph_format.space_after  = Pt(3)
            ep.paragraph_format.left_indent  = Inches(0.25)
            _add_para_left_border(ep, bg, sz='8')
            sr = ep.add_run(f'[{sev}]  ')
            sr.bold           = True
            sr.font.size      = Pt(9)
            sr.font.color.rgb = _rgb(bg)
            tr = ep.add_run(f.title if f else '—')
            tr.font.size = Pt(10)
            notes = getattr(entry, 'notes', '') or ''
            if notes:
                nr = ep.add_run(f'  —  {notes}')
                nr.italic    = True
                nr.font.size = Pt(9)

        _spacer(doc, 4)


# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

def _build_conclusion(doc, report):
    _section_heading(doc, '11.  Conclusion')

    conclusion = _plain_text(getattr(report, 'conclusion', '') or '')
    if not conclusion:
        conclusion = (
            'OziCyber has completed the penetration test engagement as scoped and authorised. '
            'The findings documented in this report represent a point-in-time assessment of the '
            'security posture of the target environment. Remediation of identified vulnerabilities '
            'is strongly recommended in priority order as outlined in the Remediation Tracker. '
            'OziCyber is available to provide re-test services upon request to validate that '
            'remediation efforts have been effective.'
        )
    p = doc.add_paragraph(conclusion)
    p.paragraph_format.space_after = Pt(10)
    for r in p.runs:
        r.font.size = Pt(10)

    notes = _plain_text(getattr(report, 'client_notes', '') or '')
    if notes:
        _sub_heading(doc, 'Client Notes')
        np = doc.add_paragraph(notes)
        np.paragraph_format.left_indent = Inches(0.15)
        _add_para_left_border(np, GOLD, sz='12')
        for r in np.runs:
            r.font.size = Pt(10)

    _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# HEADER & FOOTER
# ══════════════════════════════════════════════════════════════════════════════

def _add_header_footer(doc, report):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # ── Header ────────────────────────────────────────────────────────────────
    header = section.header
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    htbl = header.add_table(1, 3, Inches(6.5))
    _tbl_no_borders(htbl._tbl)
    _tbl_width(htbl._tbl, PAGE_W_DXA)
    lw = int(PAGE_W_DXA * 0.50)
    mw = int(PAGE_W_DXA * 0.25)
    rw = PAGE_W_DXA - lw - mw
    _cell_width(htbl.rows[0].cells[0], lw)
    _cell_width(htbl.rows[0].cells[1], mw)
    _cell_width(htbl.rows[0].cells[2], rw)

    for c in htbl.rows[0].cells:
        _set_cell_border(c, bottom={'val': 'single', 'sz': '6', 'color': GREEN})

    lp = htbl.rows[0].cells[0].paragraphs[0]
    lp.paragraph_format.space_after = Pt(3)
    lr = lp.add_run(report.title or 'Penetration Test Report')
    lr.bold           = True
    lr.font.size      = Pt(8)
    lr.font.color.rgb = _rgb(GREEN)

    mp = htbl.rows[0].cells[1].paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_after = Pt(3)
    mr = mp.add_run('OziCyber Security')
    mr.font.size      = Pt(8)
    mr.font.color.rgb = _rgb(MID_TEXT)

    rp = htbl.rows[0].cells[2].paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(3)
    rr = rp.add_run('CONFIDENTIAL')
    rr.bold           = True
    rr.font.size      = Pt(8)
    rr.font.color.rgb = _rgb(GOLD)

    # ── Footer ────────────────────────────────────────────────────────────────
    footer = section.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    ftbl = footer.add_table(1, 2, Inches(6.5))
    _tbl_no_borders(ftbl._tbl)
    _tbl_width(ftbl._tbl, PAGE_W_DXA)
    _cell_width(ftbl.rows[0].cells[0], lw + mw)
    _cell_width(ftbl.rows[0].cells[1], rw)

    for c in ftbl.rows[0].cells:
        _set_cell_border(c, top={'val': 'single', 'sz': '6', 'color': GREEN})

    flp = ftbl.rows[0].cells[0].paragraphs[0]
    flp.paragraph_format.space_before = Pt(3)
    flr = flp.add_run(f'{getattr(report.engagement, "name", "Engagement")}  —  v{report.version}')
    flr.font.size      = Pt(8)
    flr.font.color.rgb = _rgb(MID_TEXT)

    frp = ftbl.rows[0].cells[1].paragraphs[0]
    frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    frp.paragraph_format.space_before = Pt(3)
    frr = frp.add_run('Page ')
    frr.font.size      = Pt(8)
    frr.font.color.rgb = _rgb(MID_TEXT)
    _add_page_number_field(frp)
    frp.add_run(' of ').font.size = Pt(8)
    _add_num_pages_field(frp)


# ══════════════════════════════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def _build_finding_codes(findings):
    codes    = {}
    counters = {}
    for f in findings:
        ptype  = getattr(f, 'pentest_type', '') or ''
        prefix = ptype[:3].upper() if ptype else 'FND'
        counters[prefix] = counters.get(prefix, 0) + 1
        codes[f.id] = f'{prefix}-{counters[prefix]:02d}'
    return codes


def generate_pentest_docx_report(report, findings_qs, attack_chain_entries=None):
    """
    Build a fully styled DOCX report from scratch.
    Returns a BytesIO buffer containing the .docx bytes.
    """
    findings   = list(findings_qs)
    ac_entries = list(attack_chain_entries) if attack_chain_entries else []

    sev_counts    = {s: sum(1 for f in findings if getattr(f, 'severity', '') == s) for s in SEV_ORDER}
    finding_codes = _build_finding_codes(findings)

    doc = Document()

    # A4, 2cm margins
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    style           = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    _build_cover_page(doc, report)
    _build_doc_control(doc, report)
    _build_exec_summary(doc, report, findings, sev_counts)
    _build_engagement_details(doc, report)
    _build_scope(doc, report)
    _build_finding_summary(doc, report, findings, sev_counts)
    _build_key_findings(doc, report, findings, finding_codes)
    _build_methodology(doc, report)
    _build_detailed_findings(doc, report, findings, finding_codes)
    _build_remediation_tracker(doc, report, findings, finding_codes)
    _build_attack_chain(doc, report, ac_entries)
    _build_conclusion(doc, report)
    _add_header_footer(doc, report)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
