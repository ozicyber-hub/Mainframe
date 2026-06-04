"""
Django management command: build_report_template

Transforms the annotated OziCyber placeholder DOCX template into a
generator.py-compatible template with proper <<marker>> syntax.

Usage (inside Docker):
    python manage.py build_report_template \\
        --input  /path/to/annotated_template.docx \\
        --output /path/to/generator_template.docx

What it does:
  1. Normalises <<MARKER>> to <<marker>> (strip + lower + spaces→underscores)
  2. TABLE[finding list]   → one template row with <<fl.xxx>> markers
  3. TABLE[duplicate list] → DELETED (including preceding duplicate heading)
  4. TABLE[APP-1]          → wrapped in <<f:start>> / <<f:end>>, cells replaced
  5. TABLE[APP-2..5]       → DELETED
  6. TABLE[remediation]    → one template row with <<r.xxx>> markers
  7. TABLE[version]        → one template row with <<v.xxx>> markers
  8. TABLE[scope / IP]     → one template row with <<s.item>>
  9. TABLE[scoped targets] → one template row with <<s.item>>
 10. Dashboard table       → chart/key-findings markers injected
 11. Developer-note paragraphs → DELETED or replaced with count markers
"""

import re
import sys
from copy import deepcopy

from django.core.management.base import BaseCommand, CommandError

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise CommandError("python-docx is not installed — run: pip install python-docx")


# ── XML helpers ───────────────────────────────────────────────────────────────

_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_W_T  = f'{{{_W_NS}}}t'
_W_P  = f'{{{_W_NS}}}p'
_W_TR = f'{{{_W_NS}}}tr'
_W_TC = f'{{{_W_NS}}}tc'


def _elem_text(elem):
    return ''.join(t.text or '' for t in elem.iter(_W_T))


def _first_unique_cell(row, col_idx):
    """Return the first distinct Cell at column col_idx (skips vMerge duplicates)."""
    seen = set()
    for i, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        if i == col_idx:
            return cell
    # Fallback: just return cells[col_idx]
    return row.cells[col_idx]


def _set_cell_marker(cell, marker_text):
    """
    Replace ALL text inside a cell with marker_text, preserving the first
    run's character properties so Word doesn't break the table styling.
    """
    tc = cell._tc
    all_t = list(tc.iter(_W_T))
    if not all_t:
        # Build a minimal paragraph with a run
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = marker_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        return

    first_t = all_t[0]
    first_t.text = marker_text
    first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    for t in all_t[1:]:
        t.text = ''


def _make_marker_para(text):
    """Create a bare paragraph element containing just marker_text."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def _remove_body_elem(elem):
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)


def _normalize_marker_text(s):
    return re.sub(
        r'<<([^>]+)>>',
        lambda m: '<<' + m.group(1).strip().lower().replace(' ', '_') + '>>',
        s,
    )


def _normalize_all_markers(doc):
    """
    Strip + lowercase + spaces→underscores for every <<MARKER>> in the doc.

    Pass 1: normalise within individual <w:t> elements.
    Pass 2: handle markers split across multiple runs in a paragraph by
            concatenating, normalising, writing back to the first run.
    """
    # Pass 1 — per text-node
    for t in doc.element.iter(_W_T):
        if t.text and '<<' in t.text:
            t.text = _normalize_marker_text(t.text)

    # Pass 2 — per paragraph (catches split-run markers)
    def _fix_paras(paras):
        for para in paras:
            if '<<' not in (para.text or ''):
                continue
            runs = para.runs
            full = ''.join(r.text or '' for r in runs)
            if '<<' not in full:
                continue
            normed = _normalize_marker_text(full)
            if normed != full and runs:
                runs[0].text = normed
                for r in runs[1:]:
                    r.text = ''

    _fix_paras(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _fix_paras(cell.paragraphs)


def _strip_dev_annotations(doc):
    """
    Remove any remaining <<...>> markers that are purely developer annotations
    (long comments, ALL CAPS instructions, anything with uppercase letters that
    didn't normalise to a known marker).  These are leftover dev notes.
    We detect them: after normalisation every annotation marker still contains
    words like 'color', 'match', 'background', 'comment', 'here', 'put', etc.
    We blank those text nodes rather than removing them so we don't break layout.
    """
    annotation_words = {
        'color', 'match', 'background', 'severity', 'here', 'put', 'tester',
        'amounts', 'cricitcal', 'go', 'goes', 'screenshots', 'poc', 'figure',
        'stop', 'ffinding', 'finding_1', 'finding_2', 'finding_3',
    }
    for t in doc.element.iter(_W_T):
        if not t.text or '<<' not in t.text:
            continue
        # Replace each marker that looks like an annotation
        def _clean(m):
            inner = m.group(1).strip()
            if any(w in inner for w in annotation_words):
                return ''
            return m.group(0)
        t.text = re.sub(r'<<([^>]*)>>', _clean, t.text)


# ── Table finders ────────────────────────────────────────────────────────────

def _find_table_by_header(doc, *header_keywords, min_cols=1):
    """
    Return the first table whose first row contains ALL given keywords
    (case-insensitive substring match across all cells).
    """
    for tbl in doc.tables:
        if len(tbl.columns) < min_cols:
            continue
        if not tbl.rows:
            continue
        header_text = ' '.join(c.text for c in tbl.rows[0].cells).lower()
        if all(kw.lower() in header_text for kw in header_keywords):
            return tbl
    return None


def _find_table_by_first_cell(doc, prefix, col_idx=0):
    """Return the first table where cell(0, col_idx) text starts with prefix."""
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        cells = tbl.rows[0].cells
        if len(cells) > col_idx and cells[col_idx].text.strip().startswith(prefix):
            return tbl
    return None


# ── Paragraph finders ────────────────────────────────────────────────────────

def _find_para_containing(doc, substring):
    """Return the first body paragraph whose text contains substring."""
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def _para_elem(para):
    return para._p


# ── Main transformation functions ─────────────────────────────────────────────

def _transform_finding_list_table(tbl, log):
    """
    TABLE: Risk / Reference / Weakness  (summary list)
    → keep header row, replace all data rows with ONE template row using <<fl.xxx>>.
    """
    if len(tbl.rows) < 2:
        log('  [skip] finding list table has < 2 rows')
        return

    # Delete rows from the end, keep only header row
    for row in list(tbl.rows)[1:]:
        tbl._tbl.remove(row._tr)

    # Append one template row cloned from original row 1 structure
    tmpl_tr = OxmlElement('w:tr')
    for marker in ['<<fl.severity_display>>', '<<fl.code>>', '<<fl.title>>']:
        tc = OxmlElement('w:tc')
        p  = OxmlElement('w:p')
        r  = OxmlElement('w:r')
        t  = OxmlElement('w:t')
        t.text = marker
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        tmpl_tr.append(tc)
    tbl._tbl.append(tmpl_tr)
    log('  [ok]   transformed finding list table → <<fl.xxx>> template row')


def _delete_duplicate_finding_list(doc, log):
    """
    Remove the annotated duplicate finding list table (TABLE[3] in the original)
    and the preceding duplicate section heading paragraph.
    Detection: a table whose row 0 has both a severity marker and <<severity>> or
    <<reference>> dev-annotation markers.
    """
    body = doc.element.body

    # Find the paragraph with <<type_of_test>> Findings (duplicate heading).
    # After _normalize_all_markers (2-pass), the full paragraph text will be
    # normalised even if the marker was originally split across runs.
    dup_heading = None
    for p in doc.paragraphs:
        pt = p.text
        if 'type_of_test' in pt and 'Findings' in pt and \
                'Detailed' not in pt and 'Assessment' not in pt:
            dup_heading = p
            break

    # Find a table whose header row contains 'severity' or 'reference' annotation markers
    dup_table = None
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        r0_text = _elem_text(tbl.rows[0]._tr).lower()
        if 'risk' in r0_text and 'reference' in r0_text and 'weakness' in r0_text:
            # Check if this looks like an annotation/dev table by looking for
            # finding list table — prefer one that has dev markers in its header row
            # We want to REMOVE the one that still has annotation text in header
            raw_header = ' '.join(c.text for c in tbl.rows[0].cells)
            if '<<' in raw_header:
                dup_table = tbl
                break

    if dup_heading:
        _remove_body_elem(_para_elem(dup_heading))
        log('  [ok]   removed duplicate section heading paragraph')

    if dup_table:
        _remove_body_elem(dup_table._tbl)
        log('  [ok]   removed annotated duplicate finding list table')
    else:
        log('  [skip] annotated duplicate finding list table not found')


def _transform_finding_detail_tables(doc, log):
    """
    TABLE[4] (APP-1): transform cells to <<f.xxx>> markers.
    TABLE[5..8] (APP-2..5): DELETE.
    Wrap TABLE[4] in <<f:start>> / <<f:end>> paragraphs.
    """
    app_tables = []
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        first_cell_text = tbl.rows[0].cells[0].text.strip()
        if re.match(r'^APP-\d+$', first_cell_text):
            app_tables.append(tbl)

    if not app_tables:
        log('  [skip] no APP-N finding detail tables found')
        return

    log(f'  [info] found {len(app_tables)} APP-N finding tables')

    tmpl_tbl = app_tables[0]

    # ── Transform APP-1 template table ───────────────────────────────────────
    rows = tmpl_tbl.rows
    if len(rows) >= 1:
        # Row 0: code | title | severity (merged cols 2+3)
        unique_cells_r0 = list(dict.fromkeys(c._tc for c in rows[0].cells))
        if len(unique_cells_r0) >= 3:
            # code
            all_t = list(unique_cells_r0[0].iter(_W_T))
            if all_t:
                all_t[0].text = '<<f.code>>'
                for t in all_t[1:]: t.text = ''
            # title
            all_t = list(unique_cells_r0[1].iter(_W_T))
            if all_t:
                all_t[0].text = '<<f.title>>'
                for t in all_t[1:]: t.text = ''
            # severity (may be two merged cells — set both)
            for tc in unique_cells_r0[2:]:
                all_t = list(tc.iter(_W_T))
                if all_t:
                    all_t[0].text = 'RISK RATING: <<f.severity_display>>'
                    for t in all_t[1:]: t.text = ''

    if len(rows) >= 4:
        # Rows 1-3: affected_asset (merged vertically, cols 0+1), label, value
        # Affected asset — first unique cell in row 1
        unique_cells_r1 = list(dict.fromkeys(c._tc for c in rows[1].cells))
        if len(unique_cells_r1) >= 1:
            all_t = list(unique_cells_r1[0].iter(_W_T))
            if all_t:
                all_t[0].text = '<<f.affected_asset>>'
                for t in all_t[1:]: t.text = ''

        # Value cells: last unique tc per row
        for row_i, key in zip([1, 2, 3], ['level_of_access', 'consequence', 'likelihood_label']):
            unique = list(dict.fromkeys(c._tc for c in rows[row_i].cells))
            if len(unique) >= 2:
                last_tc = unique[-1]
                all_t = list(last_tc.iter(_W_T))
                if all_t:
                    all_t[0].text = f'<<f.{key}>>'
                    for t in all_t[1:]: t.text = ''

    content_row_map = {
        4: ('Description', 'description'),
        5: ('Consequence',  'impact'),
        6: ('Proof of Concept', 'supporting_evidence'),
        7: ('Remediation',  'recommendations'),
    }
    for row_i, (label, key) in content_row_map.items():
        if row_i >= len(rows):
            continue
        unique = list(dict.fromkeys(c._tc for c in rows[row_i].cells))
        if unique:
            tc = unique[0]
            all_t = list(tc.iter(_W_T))
            if all_t:
                all_t[0].text = f'{label}\n<<f.{key}>>'
                for t in all_t[1:]: t.text = ''

    log('  [ok]   transformed APP-1 table to <<f.xxx>> markers')

    # ── Add <<f:start>> before APP-1 table ───────────────────────────────────
    tbl_elem = tmpl_tbl._tbl
    tbl_elem.addprevious(_make_marker_para('<<f:start>>'))
    log('  [ok]   inserted <<f:start>> before APP-1 table')

    # ── Delete APP-2 through APP-5 tables ────────────────────────────────────
    for tbl in app_tables[1:]:
        _remove_body_elem(tbl._tbl)
    log(f'  [ok]   deleted {len(app_tables)-1} duplicate APP-N tables')

    # ── Add <<f:end>> immediately after the template table ───────────────────
    tbl_elem.addnext(_make_marker_para('<<f:end>>'))
    log('  [ok]   inserted <<f:end>> after APP-1 table')


def _delete_static_detail_headings(doc, log):
    """
    Remove the static '<<TYPE OF TEST>> Assessment Findings' paragraph and
    the 'The following security issues…' intro paragraph from the Detailed
    Findings section. The generator inserts its own type headings and intro text.
    """
    removed = 0
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if 'assessment_findings' in t or 'type_of_test' in t and 'assessment' in t.lower():
            _remove_body_elem(_para_elem(p))
            removed += 1
        elif t.startswith('The following') and 'assessment' in t.lower() and 'security' in t.lower():
            _remove_body_elem(_para_elem(p))
            removed += 1
    if removed:
        log(f'  [ok]   removed {removed} static detail section heading/intro paragraph(s)')
    else:
        log('  [skip] no static detail headings found to remove')


def _transform_remediation_table(tbl, log):
    """
    Keep header row. Replace all data rows with one template row using <<r.xxx>>.
    """
    if len(tbl.rows) < 2:
        log('  [skip] remediation table has < 2 rows')
        return

    for row in list(tbl.rows)[1:]:
        tbl._tbl.remove(row._tr)

    markers = ['<<r.code>>', '<<r.title>>', '<<r.remediated>>', '<<r.date>>', '<<r.status>>']
    tmpl_tr = OxmlElement('w:tr')
    for marker in markers:
        tc = OxmlElement('w:tc')
        p  = OxmlElement('w:p')
        r  = OxmlElement('w:r')
        t  = OxmlElement('w:t')
        t.text = marker
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        tmpl_tr.append(tc)
    tbl._tbl.append(tmpl_tr)
    log('  [ok]   transformed remediation table → <<r.xxx>> template row')


def _transform_version_table(tbl, log):
    """
    Keep header row. Replace all data rows with one template row using <<v.xxx>>.
    """
    if len(tbl.rows) < 2:
        log('  [skip] version table has < 2 rows')
        return

    for row in list(tbl.rows)[1:]:
        tbl._tbl.remove(row._tr)

    markers = ['<<v.version>>', '<<v.date>>', '<<v.author>>', '<<v.comment>>']
    tmpl_tr = OxmlElement('w:tr')
    for marker in markers:
        tc = OxmlElement('w:tc')
        p  = OxmlElement('w:p')
        r  = OxmlElement('w:r')
        t  = OxmlElement('w:t')
        t.text = marker
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        tmpl_tr.append(tc)
    tbl._tbl.append(tmpl_tr)
    log('  [ok]   transformed version table → <<v.xxx>> template row')


def _transform_scope_table(tbl, log, label='scope'):
    """
    Keep header row. Replace all data rows with one template row using <<s.item>>.
    """
    if len(tbl.rows) < 2:
        log(f'  [skip] {label} table has < 2 rows')
        return

    for row in list(tbl.rows)[1:]:
        tbl._tbl.remove(row._tr)

    # One template row — number of cells = same as header
    n_cols = len(tbl.rows[0].cells)
    tmpl_tr = OxmlElement('w:tr')
    for ci in range(n_cols):
        tc = OxmlElement('w:tc')
        p  = OxmlElement('w:p')
        r  = OxmlElement('w:r')
        t  = OxmlElement('w:t')
        t.text = '<<s.item>>' if ci == 0 else ''
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        tmpl_tr.append(tc)
    tbl._tbl.append(tmpl_tr)
    log(f'  [ok]   transformed {label} table → <<s.item>> template row')


def _transform_dashboard_table(tbl, log):
    """
    Transform the 2×2 dashboard table:
      - Left column: leave as-is (template's own embedded charts stay)
      - Right column: replace <<FINDING X TITLE>> dev-note markers with
        a <<kf.title>> paragraph (one per line, processed by _process_cell_lists)
    """
    if len(tbl.rows) < 1:
        log('  [skip] dashboard table has < 1 rows')
        return

    rows = tbl.rows

    # Right column (merged cell spanning both rows in the original template)
    unique_r0 = list(dict.fromkeys(c._tc for c in rows[0].cells))
    if len(unique_r0) >= 2:
        right_tc = unique_r0[-1]
        # Clear ALL existing text / paragraphs from the right cell
        for p_elem in list(right_tc.findall(_W_P)):
            right_tc.remove(p_elem)
        # Add "KEY FINDINGS" label paragraph
        def _make_text_para(text, bold=False):
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            if bold:
                rpr = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                rpr.append(b)
                r.append(rpr)
            t = OxmlElement('w:t')
            t.text = text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            p.append(r)
            return p
        right_tc.append(_make_text_para('KEY FINDINGS', bold=True))
        right_tc.append(_make_text_para('<<kf.title>>'))

    log('  [ok]   dashboard right cell → KEY FINDINGS + <<kf.title>> marker')


def _fix_cover_page(doc, log):
    """
    Replace any occurrence of the literal word 'REDACTED' with <<client_name>>.
    Searches the full document XML so text boxes and frames on the cover page
    are found, not just body paragraphs.
    """
    replaced = 0
    for t in doc.element.iter(_W_T):
        if t.text and 'REDACTED' in t.text:
            t.text = t.text.replace('REDACTED', '<<client_name>>')
            replaced += 1
    if replaced:
        log(f'  [ok]   replaced REDACTED with <<client_name>> in {replaced} text node(s)')
    else:
        log('  [skip] REDACTED cover text not found')


def _fix_lorem_ipsum_disclaimer(doc, log):
    """
    Replace lorem ipsum placeholder text in the Disclaimer section with
    <<disclaimer>> so the generator can fill it from the replacements dict.
    Searches all text nodes so text-box content on the cover page is caught.
    """
    found = False
    # Collect all <w:t> nodes that are part of a lorem ipsum paragraph
    # Strategy: find any <w:t> whose text starts with 'Lorem ipsum' and blank
    # all sibling <w:t> nodes in the same run/paragraph, then set the first to marker.
    _W_R = f'{{{_W_NS}}}r'
    for t in list(doc.element.iter(_W_T)):
        if t.text and t.text.strip().lower().startswith('lorem ipsum'):
            # Walk up to the parent paragraph and blank all its text nodes
            parent_r = t.getparent()
            parent_p = parent_r.getparent() if parent_r is not None else None
            if parent_p is not None:
                all_t_in_para = list(parent_p.iter(_W_T))
                for sibling_t in all_t_in_para:
                    sibling_t.text = ''
            t.text = '<<disclaimer>>'
            found = True
            log('  [ok]   lorem ipsum → <<disclaimer>> marker')
    if not found:
        log('  [skip] lorem ipsum paragraph not found')


def _replace_dev_comment_para(doc, log):
    """
    Replace the dev-note paragraph that says <<PUT A LINE COMMENT…>> with a
    human-readable finding count summary using scalar markers.
    """
    for p in list(doc.paragraphs):
        if 'put_a_line_comment' in p.text.lower() or 'put a line comment' in p.text.lower():
            # Clear all runs and set new text
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = (
                    '<<critical_count>> Critical  /  <<high_count>> High  /  '
                    '<<medium_count>> Medium  /  <<low_count>> Low  /  '
                    '<<info_count>> Informational findings identified.'
                )
            else:
                p.add_run(
                    '<<critical_count>> Critical  /  <<high_count>> High  /  '
                    '<<medium_count>> Medium  /  <<low_count>> Low  /  '
                    '<<info_count>> Informational findings identified.'
                )
            log('  [ok]   replaced dev-comment paragraph with finding count markers')
            return
    log('  [skip] dev-comment paragraph not found')


# ── Command ───────────────────────────────────────────────────────────────────

class Command(BaseCommand):
    help = (
        'Transform an annotated OziCyber DOCX template into a generator.py-compatible '
        'template with proper <<marker>> syntax for upload to the template repo.'
    )

    def add_arguments(self, parser):
        parser.add_argument('--input',  required=True, help='Path to annotated input template .docx')
        parser.add_argument('--output', required=True, help='Path to save the generator-ready template .docx')

    def handle(self, *args, **options):
        inp  = options['input']
        outp = options['output']

        def log(msg):
            self.stdout.write(msg)

        log(f'\nLoading template: {inp}')
        try:
            doc = Document(inp)
        except Exception as e:
            raise CommandError(f'Could not open template: {e}')

        log('\n── Step 1: Normalise markers ─────────────────────────')
        _normalize_all_markers(doc)
        log('  [ok]   all <<MARKERS>> normalised to <<lowercase_underscores>>')

        log('\n── Step 2: Fix cover page ───────────────────────────')
        _fix_cover_page(doc, log)

        log('\n── Step 3: Fix disclaimer (lorem ipsum) ─────────────')
        _fix_lorem_ipsum_disclaimer(doc, log)

        log('\n── Step 4: Replace dev-note paragraph ───────────────')
        _replace_dev_comment_para(doc, log)

        log('\n── Step 5: Remove annotated duplicate finding list ──')
        _delete_duplicate_finding_list(doc, log)

        log('\n── Step 6: Transform finding summary list table ─────')
        all_risk_tbls = [
            t for t in doc.tables
            if t.rows and all(
                kw in ' '.join(c.text for c in t.rows[0].cells).lower()
                for kw in ('risk', 'reference', 'weakness')
            )
        ]
        if all_risk_tbls:
            _transform_finding_list_table(all_risk_tbls[0], log)
            for extra in all_risk_tbls[1:]:
                _remove_body_elem(extra._tbl)
                log(f'  [ok]   deleted extra duplicate finding list table')
        else:
            log('  [skip] Risk/Reference/Weakness table not found')

        log('\n── Step 7: Transform dashboard table ────────────────')
        dash_tbl = _find_table_by_header(doc, 'Overall', 'Vulnerability')
        if dash_tbl:
            _transform_dashboard_table(dash_tbl, log)
        else:
            log('  [skip] dashboard table not found')

        log('\n── Step 8: Delete static detailed-findings headings ─')
        _delete_static_detail_headings(doc, log)

        log('\n── Step 9: Transform & wrap finding detail tables ───')
        _transform_finding_detail_tables(doc, log)

        log('\n── Step 10: Transform remediation table ─────────────')
        rem_tbl = _find_table_by_header(doc, 'Finding Reference', 'Remediated')
        if rem_tbl:
            _transform_remediation_table(rem_tbl, log)
        else:
            log('  [skip] remediation table not found')

        log('\n── Step 11: Transform scope / IP table ──────────────')
        scope_tbl = _find_table_by_header(doc, 'IP Address')
        if scope_tbl:
            _transform_scope_table(scope_tbl, log, label='scope/IP')
        else:
            log('  [skip] IP Address table not found')

        log('\n── Step 12: Transform scoped targets table ──────────')
        targets_tbl = _find_table_by_header(doc, 'Domain')
        if targets_tbl:
            _transform_scope_table(targets_tbl, log, label='scoped targets')
        else:
            log('  [skip] Domain table not found')

        log('\n── Step 13: Transform version control table ─────────')
        ver_tbl = _find_table_by_header(doc, 'Version', 'Revision')
        if ver_tbl:
            _transform_version_table(ver_tbl, log)
        else:
            log('  [skip] Version/Revision table not found')

        log('\n── Step 14: Strip leftover dev annotations ──────────')
        _strip_dev_annotations(doc)
        log('  [ok]   dev annotation markers cleared')

        log(f'\n── Saving → {outp}')
        try:
            doc.save(outp)
        except Exception as e:
            raise CommandError(f'Could not save output: {e}')

        log('\n✓ Template built successfully.\n')
        log('Next steps:')
        log('  1. Copy the output file into the Docker container (or mount a volume)')
        log('  2. Upload it via the Reports → Templates section in the UI')
        log('  3. Set it as the default template for your organisation')
        log('\nKnown placeholders WITHOUT a matching app field (shown as empty):')
        log('  <<level_of_access>>  — add to Finding model if needed')
        log('  <<concessions>>      — add to Report model if needed')
        log('  <<constraints>>      — add to Report model if needed')
