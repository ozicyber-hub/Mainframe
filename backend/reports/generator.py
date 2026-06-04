"""
Report generator — populates a styled Word template using python-docx directly.

No Jinja2 / docxtpl. Template conventions:

  Scalar placeholders  →  <<placeholder_name>>  anywhere in paragraph text
                          e.g. <<report_title>>, <<client_name>>, <<today>>

  Loop tables          →  Put ONE template row in the table with cell text like:
                            <<f.code>>  <<f.title>>  <<f.severity_display>>  …  (findings)
                            <<r.code>>  <<r.title>>  <<r.severity>>  …           (remediation)
                            <<v.version>>  <<v.date>>  <<v.author>>  <<v.comment>> (version log)
                          The generator finds the template row automatically, clones it
                          for every item, fills it in, then removes the original.

  Chart placeholder    →  A paragraph containing only <<SEVERITY_CHART>> is replaced
                          with an inline severity bar chart (requires matplotlib).

See backend/TEMPLATE_VARIABLES.md for the complete placeholder reference.
"""

import io
import os
from copy import deepcopy
from datetime import date

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE_PATH = os.path.normpath(
    os.path.join(BASE_DIR, '..', 'report_fixtures', 'template.docx')
)

CONSEQUENCE_MAP = {'HIGH': 'Major', 'MEDIUM': 'Moderate', 'LOW': 'Minor'}
LIKELIHOOD_MAP  = {'HIGH': 'Likely', 'MEDIUM': 'Possible', 'LOW': 'Unlikely'}
SEVERITIES = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFORMATIONAL']

TYPE_PREFIXES = {
    'WEB_APP':    'WEB',  'INTERNAL':  'INT',  'EXTERNAL':  'EXT',
    'MOBILE':     'MOB',  'API':       'API',  'CLOUD':     'CLD',
    'SOCIAL_ENG': 'SOC',  'PHYSICAL':  'PHY',  'RED_TEAM':  'RED',
    'WIRELESS':   'WLS',  'OTHER':     'OTH',
}

SEV_DISPLAY = {
    'CRITICAL': 'Critical', 'HIGH': 'High', 'MEDIUM': 'Medium',
    'LOW': 'Low', 'INFORMATIONAL': 'Informational',
}

SEV_COLORS = {
    'CRITICAL':     'FF0000',
    'HIGH':         'E17468',
    'MEDIUM':       'FFF1AA',
    'LOW':          '49A58B',
    'INFORMATIONAL':'00467A',
}

STATUS_DISPLAY = {
    'DRAFT': 'Draft', 'OPEN': 'Open', 'IN_REVIEW': 'In Review',
    'PUBLISHED': 'Published', 'REMEDIATED': 'Remediated',
    'FALSE_POSITIVE': 'False Positive', 'ACCEPTED_RISK': 'Risk Accepted',
}

CVSS_METRIC_DISPLAY = {
    'av': {'N': 'Network', 'A': 'Adjacent', 'L': 'Local', 'P': 'Physical'},
    'ac': {'L': 'Low', 'H': 'High'},
    'pr': {'N': 'None', 'L': 'Low', 'H': 'High'},
    'ui': {'N': 'None', 'R': 'Required'},
    's':  {'U': 'Unchanged', 'C': 'Changed'},
    'c':  {'N': 'None', 'L': 'Low', 'H': 'High'},
    'i':  {'N': 'None', 'L': 'Low', 'H': 'High'},
    'a':  {'N': 'None', 'L': 'Low', 'H': 'High'},
}


# ── Text / image helpers ──────────────────────────────────────────────────────

def html_to_plain(html):
    """Strip Quill HTML to clean plain text, preserving list bullets."""
    if not html:
        return ''
    soup = BeautifulSoup(html, 'html.parser')
    lines = []
    for elem in soup.descendants:
        if elem.name == 'li':
            lines.append('• ' + elem.get_text(strip=True))
        elif elem.name in ('p', 'div', 'br'):
            text = elem.get_text(strip=True)
            if text:
                lines.append(text)
    if not lines:
        lines = [soup.get_text(separator='\n', strip=True)]
    return '\n'.join(lines)


def _extract_images_from_html(html):
    """
    Extract embedded images from Quill HTML.
    Returns a list of bytes objects (PNG/JPEG raw data).
    Quill stores images as <img src="data:image/png;base64,..."> data URLs.
    """
    if not html:
        return []
    import base64
    soup = BeautifulSoup(html, 'html.parser')
    images = []
    for img in soup.find_all('img'):
        src = img.get('src', '')
        if src.startswith('data:') and ',' in src:
            try:
                images.append(base64.b64decode(src.split(',', 1)[1]))
            except Exception:
                pass
    return images


def _fmt_date(d, fmt='%d %B %Y'):
    return d.strftime(fmt) if d else ''


def _cvss_display(metric_key, value):
    return CVSS_METRIC_DISPLAY.get(metric_key, {}).get(value, value or '')


def _short_mitigation(html, max_sentences=2):
    """Return the first max_sentences sentences of the plain-text recommendations."""
    text = html_to_plain(html or '')
    if not text:
        return ''
    import re as _re2
    sentences = _re2.split(r'(?<=[.!?])\s+', text.strip())
    return ' '.join(sentences[:max_sentences])


def build_finding_codes(findings):
    """Return {finding.id: 'WEB-01'} stable codes, severity-sorted within type group."""
    groups = {}
    for f in findings:
        key = f.pentest_type or '__untagged__'
        groups.setdefault(key, []).append(f)
    codes = {}
    for key, group in groups.items():
        prefix = TYPE_PREFIXES.get(key, 'FIN')
        sorted_group = sorted(
            group,
            key=lambda f: SEVERITIES.index(f.severity) if f.severity in SEVERITIES else 99
        )
        for idx, f in enumerate(sorted_group, 1):
            codes[f.id] = f'{prefix}-{idx:02d}'
    return codes


def _overall_risk(counts):
    for sev in SEVERITIES:
        if counts.get(sev, 0) > 0:
            return SEV_DISPLAY.get(sev, sev)
    return 'Informational'


# ── Marker normalisation ─────────────────────────────────────────────────────

import re as _re

def _normalize_markers(doc):
    """
    Normalise every <<...>> marker in every text node of the document:
    strip whitespace, lowercase, and replace internal spaces with underscores.

    This makes marker matching case-insensitive and space-insensitive so that
    <<CLIENT NAME>>, <<client_name>>, and <<Client Name>> all resolve to the
    same replacement key.
    Call once at the very start of generate_report_docx, before anything else.
    """
    _W_T = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    for t in doc.element.iter(_W_T):
        if t.text and '<<' in t.text:
            t.text = _re.sub(
                r'<<([^>]+)>>',
                lambda m: '<<' + m.group(1).strip().lower().replace(' ', '_') + '>>',
                t.text,
            )


# ── Template replacement helpers ──────────────────────────────────────────────

def _replace_para(para, replacements):
    """
    Replace <<key>> markers in a paragraph.

    Reads full text across all runs first (avoids split-run issues),
    normalises any remaining un-normalised markers (handles cases where
    <<CLIENT NAME>> is split across runs and survived _normalize_markers),
    applies all replacements, then writes result back into the first run
    and clears the others — preserving the first run's character formatting.
    """
    full = ''.join(r.text or '' for r in para.runs)
    if '<<' not in full:
        return
    # Second-chance normalisation for markers that were split across runs
    full = _re.sub(
        r'<<([^>]+)>>',
        lambda m: '<<' + m.group(1).strip().lower().replace(' ', '_') + '>>',
        full,
    )
    for key, value in replacements.items():
        full = full.replace(f'<<{key}>>', str(value) if value is not None else '')
    if para.runs:
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(full)


def _iter_all_paragraphs(doc):
    """Yield every paragraph in the document including tables, headers, footers."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for container in (section.header, section.footer):
            try:
                yield from container.paragraphs
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs
            except Exception:
                pass


def _replace_all_scalars(doc, replacements):
    for para in _iter_all_paragraphs(doc):
        _replace_para(para, replacements)


# ── Loop table helpers ────────────────────────────────────────────────────────

def _set_cell_shading(tc_elem, hex_color):
    """Set the background fill of a <w:tc> element to hex_color."""
    from docx.oxml import OxmlElement
    tcPr = tc_elem.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc_elem.insert(0, tcPr)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.insert(0, shd)


def _fill_element_markers(elem, item, prefix):
    """Replace <<prefix.key>> in every text node inside an XML element.
    Also colors table cells that contained a severity marker."""
    _W_T  = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    _W_TC = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'
    marker = f'<<{prefix}.'
    sev_trigger = {f'<<{prefix}.severity>>', f'<<{prefix}.severity_display>>'}

    # Color severity cells BEFORE text replacement so we can still detect the marker
    sev = item.get('severity', '')
    color = SEV_COLORS.get(sev)
    if color:
        for tc in elem.iter(_W_TC):
            cell_text = ''.join(t.text or '' for t in tc.iter(_W_T))
            if any(m in cell_text for m in sev_trigger):
                _set_cell_shading(tc, color)

    # Replace text markers
    for t in elem.iter(_W_T):
        if not t.text or marker not in t.text:
            continue
        for key, value in item.items():
            t.text = t.text.replace(f'<<{prefix}.{key}>>', str(value) if value is not None else '')


def _process_loop_tables(doc, prefix, items):
    """
    For every table in the document that contains <<prefix. markers:
      - Find the first row with a marker (template start)
      - Find the last *consecutive* row from that point that also has markers
      - For <<f.>> findings: clone the whole template table once per pentest type,
        prepend a Heading-2 sub-title, and remove the original template table
      - Other prefixes: clone rows in-place as before
    """
    _W_TR = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'
    marker = f'<<{prefix}.'

    def _is_pure_template_row(row):
        """
        A row qualifies as a loop-template row only if the cell containing
        the marker has NO other non-marker text paragraphs.  This prevents
        layout table rows (e.g. charts + key-findings text in the same cell)
        from being mis-identified as template rows.
        For the 'f' prefix, skip cells that also contain '<<fl.' markers —
        those belong to the flat-list table and must not be grabbed here.
        """
        for cell in row.cells:
            if marker not in cell.text:
                continue
            # '<<f.' must not match cells that belong to the '<<fl.' flat-list table
            if marker == '<<f.' and '<<fl.' in cell.text:
                continue
            non_empty = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
            if non_empty and all(marker in pt for pt in non_empty):
                return True
        return False

    for table in doc.tables:
        first_idx = None
        for i, row in enumerate(table.rows):
            if _is_pure_template_row(row):
                first_idx = i
                break
        if first_idx is None:
            continue

        last_idx = first_idx
        for i in range(first_idx + 1, len(table.rows)):
            if _is_pure_template_row(table.rows[i]):
                last_idx = i
            else:
                break

        tmpl_trs = [deepcopy(row._tr) for row in table.rows[first_idx:last_idx + 1]]

        if prefix == 'f':
            # Group items by pentest type (order from items list is already type+severity sorted)
            groups: dict = {}
            for item in items:
                groups.setdefault(item.get('pentest_type_code', ''), []).append(item)

            tbl_elem = table._tbl
            body = tbl_elem.getparent()

            for group_items in groups.values():
                type_display = group_items[0].get('pentest_type') or 'Other'

                # Spacer + sub-heading before the table
                spacer = doc.add_paragraph()
                spacer._element.getparent().remove(spacer._element)
                tbl_elem.addprevious(spacer._element)

                for style_name in ('Heading 2', 'Heading 3', 'Heading 1'):
                    try:
                        h = doc.add_paragraph(type_display + ' Findings', style=style_name)
                        h._element.getparent().remove(h._element)
                        tbl_elem.addprevious(h._element)
                        break
                    except Exception:
                        continue

                # Deep-copy the whole table, populate with this group's items
                tbl_clone = deepcopy(tbl_elem)
                for item in group_items:
                    for tmpl_tr in tmpl_trs:
                        clone_tr = deepcopy(tmpl_tr)
                        _fill_element_markers(clone_tr, item, prefix)
                        tbl_clone.append(clone_tr)

                # Remove template rows from the clone
                clone_trs = [c for c in tbl_clone if c.tag == _W_TR]
                for tr in clone_trs[first_idx:last_idx + 1]:
                    tbl_clone.remove(tr)

                tbl_elem.addprevious(tbl_clone)

            body.remove(tbl_elem)
            continue   # original template table is gone; skip the normal row-removal below

        # Safety: if nothing to fill, don't touch the table at all
        if not items:
            continue

        # Non-finding prefixes: clone rows and insert them IN-PLACE
        # (immediately after the template rows, before any subsequent rows)
        # so that two separate loop sections in the same table stay in order.
        clones = []
        for item in items:
            for tmpl_tr in tmpl_trs:
                clone = deepcopy(tmpl_tr)
                _fill_element_markers(clone, item, prefix)
                clones.append(clone)

        all_trs = list(table.rows)
        if last_idx + 1 < len(all_trs):
            after_tr = all_trs[last_idx + 1]._tr
            for clone in reversed(clones):
                after_tr.addprevious(clone)
        else:
            for clone in clones:
                table._tbl.append(clone)

        for row in list(table.rows)[first_idx:last_idx + 1]:
            table._tbl.remove(row._tr)


# ── Cell-embedded list helpers ────────────────────────────────────────────────

def _process_cell_lists(doc, prefix, items):
    """
    Handle <<prefix.field>> markers that live as PARAGRAPHS inside a table cell
    alongside other text (e.g. 'KEY FINDINGS\\n<<kf.title>>\\nKEY MITIGATIONS…').

    For each such paragraph: clone it once per item (inserting clones immediately
    before it), fill markers, then remove the original template paragraph.
    If items is empty the template paragraph is simply removed silently.
    """
    import re as _re
    _W_T = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    _W_P = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    marker_pat = _re.compile(r'^<<' + _re.escape(prefix) + r'\.[^>]+>>$')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                for p in list(tc):
                    if p.tag != _W_P:
                        continue
                    p_text = ''.join(t.text or '' for t in p.iter(_W_T)).strip()
                    if not marker_pat.match(p_text):
                        continue
                    tmpl = deepcopy(p)
                    for item in reversed(items):
                        clone = deepcopy(tmpl)
                        _fill_element_markers(clone, item, prefix)
                        p.addprevious(clone)
                    tc.remove(p)


# ── Body-block loop helpers ───────────────────────────────────────────────────
#
# For finding sections that mix tables AND paragraphs, wrap the entire block
# in your template with marker paragraphs:
#
#   <<f:start>>          ← paragraph containing only this text
#   ... tables, paragraphs with <<f.xxx>> markers ...
#   <<f:end>>            ← paragraph containing only this text
#
# Similarly <<r:start>> / <<r:end>> for remediation, <<v:start>> / <<v:end>>
# for version rows, <<s:start>> / <<s:end>> for scope items.

_W_NS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_W_T   = f'{{{_W_NS}}}t'
_W_P   = f'{{{_W_NS}}}p'


def _elem_text(elem):
    return ''.join(t.text or '' for t in elem.iter(_W_T))


def _remove_marker_everywhere(doc, marker):
    """
    Remove any paragraph (in body, headers, footers, text boxes) whose
    full text contains the given marker string.  Catches cases where the
    user accidentally put a marker in a header or floating text box.
    """
    _W_T   = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    _W_P   = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    _W_TXB = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'

    def _clear_if_contains(root):
        for p in list(root.iter(_W_P)):
            text = ''.join(t.text or '' for t in p.iter(_W_T))
            if marker in text:
                parent = p.getparent()
                if parent is not None:
                    parent.remove(p)

    _clear_if_contains(doc.element)          # body + all nested elements
    for section in doc.sections:
        for container in (section.header, section.footer,
                          section.even_page_header, section.even_page_footer,
                          section.first_page_header, section.first_page_footer):
            try:
                _clear_if_contains(container._element)
            except Exception:
                pass


def _insert_type_heading(doc, anchor_elem, heading_text, type_display):
    """Insert a Heading-1 paragraph + intro line before anchor_elem."""
    # Heading
    for style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
        try:
            h = doc.add_paragraph(heading_text, style=style_name)
            h._element.getparent().remove(h._element)
            anchor_elem.addprevious(h._element)
            break
        except Exception:
            continue
    else:
        h = doc.add_paragraph()
        h.add_run(heading_text).bold = True
        h._element.getparent().remove(h._element)
        anchor_elem.addprevious(h._element)

    # Intro line immediately after heading (inserted before anchor → appears after heading)
    intro = doc.add_paragraph(
        f'The following {type_display} security issues were identified during the assessment.'
    )
    intro._element.getparent().remove(intro._element)
    anchor_elem.addprevious(intro._element)


def _process_body_block(doc, prefix, items):
    """
    Find <<prefix:start>> … <<prefix:end>> blocks in the document body,
    clone everything between them for each item, fill <<prefix.key>> markers,
    then remove the original block including the marker paragraphs.

    Robustness features:
    - Markers are detected by substring match (not exact), so they can share
      a paragraph with other text.
    - If <<prefix:end>> is not found in the body (e.g. accidentally placed in a
      header/text-box), the block auto-extends to the last body element that
      contains <<prefix. markers.
    - Stray marker paragraphs in headers / text-boxes are cleaned up afterwards.
    """
    start_tag = f'<<{prefix}:start>>'
    end_tag   = f'<<{prefix}:end>>'
    field_pfx = f'<<{prefix}.'
    body      = doc.element.body

    while True:
        children  = list(body)
        start_idx = end_idx = None

        for i, child in enumerate(children):
            text = _elem_text(child).strip()
            if start_tag in text:
                start_idx = i
            elif end_tag in text and start_idx is not None:
                end_idx = i
                break

        if start_idx is None:
            break

        # If end marker not found in body, auto-extend to last element with <<prefix. markers
        if end_idx is None:
            for i in range(start_idx + 1, len(children)):
                if field_pfx in _elem_text(children[i]):
                    end_idx = i
            if end_idx is None:
                break  # Nothing to clone

        # Deep-copy the template elements (everything between the markers)
        tmpl_elems = [deepcopy(e) for e in children[start_idx + 1:end_idx]]
        anchor     = children[end_idx]

        # For findings ('f' prefix): group by pentest type and insert section headings
        if prefix == 'f':
            current_type = object()  # sentinel
            for item_idx, item in enumerate(items):
                item_type = item.get('pentest_type_code', '')
                if item_type != current_type:
                    current_type = item_type
                    type_display = item.get('pentest_type') or 'Other'
                    label = type_display + ' Assessment Findings'
                    _insert_type_heading(doc, anchor, label, type_display)
                for tmpl in tmpl_elems:
                    clone = deepcopy(tmpl)
                    _fill_element_markers(clone, item, prefix)
                    anchor.addprevious(clone)
                # Page break after every finding so the next section always starts fresh
                pb_p = OxmlElement('w:p')
                pb_r = OxmlElement('w:r')
                pb_br = OxmlElement('w:br')
                pb_br.set(qn('w:type'), 'page')
                pb_r.append(pb_br)
                pb_p.append(pb_r)
                anchor.addprevious(pb_p)
        else:
            for item in items:
                for tmpl in tmpl_elems:
                    clone = deepcopy(tmpl)
                    _fill_element_markers(clone, item, prefix)
                    anchor.addprevious(clone)

        # Remove the marker paragraphs and the original template elements
        for elem in children[start_idx:end_idx + 1]:
            body.remove(elem)

    # Clean up any stray start/end marker paragraphs left in headers, text-boxes, etc.
    _remove_marker_everywhere(doc, start_tag)
    _remove_marker_everywhere(doc, end_tag)


# ── Image marker replacement ─────────────────────────────────────────────────

def _replace_image_markers(doc, prefix, items):
    """
    Replace <<prefix.supporting_evidence_img>> paragraphs with actual images.

    Because _fill_element_markers only replaces keys that exist in the item
    dict, and 'supporting_evidence_img' is NOT a key, the marker survives
    cloning.  We then walk the document in order and pair the Nth marker with
    the Nth item's image list.

    Place <<f.supporting_evidence_img>> as a standalone paragraph in your
    template wherever you want the embedded screenshots to appear.
    """
    marker   = f'<<{prefix}.supporting_evidence_img>>'
    item_idx = 0

    for para in _iter_all_paragraphs(doc):
        if marker not in para.text:
            continue
        if item_idx >= len(items):
            para.clear()
            break

        images = items[item_idx].get('_images', [])
        item_idx += 1

        para.clear()
        if not images:
            continue

        for img_bytes in images:
            run = para.add_run()
            try:
                run.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
            except Exception:
                pass
            # Each image on its own line within the paragraph
            para.add_run('\n')


# ── Executive Summary Block ───────────────────────────────────────────────────

_PAGE_W = 9360   # 6.5" content width in twips

# Light tint backgrounds for severity tiles
_SEV_TILE_TINT = {
    'CRITICAL': 'FFEDED', 'HIGH': 'FAEAE8', 'MEDIUM': 'FFFCE0',
    'LOW': 'E8F5F0', 'INFORMATIONAL': 'EAF0F7',
}
# Border + text color for tiles (slightly darkened for readability on light bg)
_SEV_TILE_COLOR = {
    'CRITICAL': 'CC0000', 'HIGH': 'C94A3E', 'MEDIUM': 'B8860B',
    'LOW': '2E8B6B', 'INFORMATIONAL': '00467A',
}
# Solid fill for severity badge pills
_SEV_BADGE_BG = {
    'CRITICAL': 'CC0000', 'HIGH': 'C94A3E', 'MEDIUM': 'C49000',
    'LOW': '3A9B7A', 'INFORMATIONAL': '00467A',
}


def _rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _get_or_add_tcPr(tc):
    p = tc.find(qn('w:tcPr'))
    if p is None:
        p = OxmlElement('w:tcPr')
        tc.insert(0, p)
    return p


def _es_tc_width(tc, w):
    pr = _get_or_add_tcPr(tc)
    for old in pr.findall(qn('w:tcW')):
        pr.remove(old)
    el = OxmlElement('w:tcW')
    el.set(qn('w:w'), str(w))
    el.set(qn('w:type'), 'dxa')
    pr.insert(0, el)


def _es_tc_margins(tc, top=80, right=120, bottom=80, left=120):
    pr = _get_or_add_tcPr(tc)
    for old in pr.findall(qn('w:tcMar')):
        pr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    pr.append(mar)


def _es_tc_valign(tc, val='center'):
    pr = _get_or_add_tcPr(tc)
    for old in pr.findall(qn('w:vAlign')):
        pr.remove(old)
    v = OxmlElement('w:vAlign')
    v.set(qn('w:val'), val)
    pr.append(v)


def _es_tc_borders(tc, top=None, right=None, bottom=None, left=None):
    """Each side: ('single', 'RRGGBB', size_eighths) or None to suppress."""
    pr = _get_or_add_tcPr(tc)
    for old in pr.findall(qn('w:tcBorders')):
        pr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for side, spec in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
        b = OxmlElement(f'w:{side}')
        if spec:
            b.set(qn('w:val'), spec[0])
            b.set(qn('w:color'), spec[1])
            b.set(qn('w:sz'), str(spec[2]))
            b.set(qn('w:space'), '0')
        else:
            b.set(qn('w:val'), 'none')
        borders.append(b)
    pr.append(borders)


def _es_tbl_props(tbl, total_w, cell_spacing=0):
    tp = tbl.find(qn('w:tblPr'))
    if tp is None:
        tp = OxmlElement('w:tblPr')
        tbl.insert(0, tp)
    for tag in ('w:tblW', 'w:tblLayout', 'w:tblBorders', 'w:tblCellSpacing', 'w:tblCellMar'):
        for old in tp.findall(qn(tag)):
            tp.remove(old)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), str(total_w))
    tw.set(qn('w:type'), 'dxa')
    tp.append(tw)
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    tp.append(lay)
    bdr = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        bdr.append(b)
    tp.append(bdr)
    if cell_spacing:
        cs = OxmlElement('w:tblCellSpacing')
        cs.set(qn('w:w'), str(cell_spacing))
        cs.set(qn('w:type'), 'dxa')
        tp.append(cs)
    # Zero default cell margins
    cm = OxmlElement('w:tblCellMar')
    for side in ('top', 'right', 'bottom', 'left'):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        cm.append(m)
    tp.append(cm)


def _es_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    sp.set(qn('w:line'), '240')
    sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)


def _es_clear_list_style(para):
    """Remove any inherited list/bullet numbering from a paragraph."""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:numPr')):
        pPr.remove(old)
    for old in pPr.findall(qn('w:pStyle')):
        pPr.remove(old)



_badge_cache = {}

def _build_badge_image(sev):
    """Render a small rounded severity badge as a PNG (cached per severity)."""
    if sev in _badge_cache:
        return _badge_cache[sev]
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches

        label     = SEV_DISPLAY.get(sev, sev).upper()
        badge_col = _SEV_BADGE_BG.get(sev, '666666')

        def h2f(h):
            return (int(h[0:2],16)/255, int(h[2:4],16)/255, int(h[4:6],16)/255)

        # Fixed figure, no auto-cropping — axes fill the whole canvas
        fig = plt.figure(figsize=(1.1, 0.30))
        fig.patch.set_alpha(0)
        ax = fig.add_axes([0.0, 0.0, 1.0, 1.0])
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis('off')

        ax.add_patch(mpatches.FancyBboxPatch(
            (0.03, 0.10), 0.94, 0.80,
            boxstyle='round,pad=0.0,rounding_size=0.28',
            facecolor=h2f(badge_col),
            edgecolor='none',
            zorder=2,
        ))
        ax.text(0.5, 0.52, label,
                ha='center', va='center',
                fontsize=8, fontweight='bold',
                color='white', zorder=3)

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=180, transparent=True)
        plt.close(fig)
        buf.seek(0)
        result = buf.read()
        _badge_cache[sev] = result
        return result
    except ImportError:
        return None


def _build_exec_summary_block(doc, counts, key_findings):
    """
    Find <<exec_summary_block>> and replace with:
      1. Summary sentence ("A total of X findings…")
      2. Severity count tiles (light tint fill + colored border + colored text)
      3. Key findings list (compact severity tag | bold title + gray description)
    """
    import re as _re2

    anchor = None
    for para in _iter_all_paragraphs(doc):
        if '<<exec_summary_block>>' in para.text:
            anchor = para
            break
    if anchor is None:
        return

    anchor_p = anchor._p

    def insert_before(elem):
        p = elem.getparent()
        if p is not None:
            p.remove(elem)
        anchor_p.addprevious(elem)

    def bare_para(before=0, after=120):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spc = OxmlElement('w:spacing')
        spc.set(qn('w:before'), str(before))
        spc.set(qn('w:after'), str(after))
        spc.set(qn('w:line'), '276')
        spc.set(qn('w:lineRule'), 'auto')
        pPr.append(spc)
        p.append(pPr)
        return p

    def add_run(p, text, bold=False, size_pt=10, color='1A1A1A'):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(szCs)
        cl = OxmlElement('w:color')
        cl.set(qn('w:val'), color)
        rPr.append(cl)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)

    # ── Summary sentence ──────────────────────────────────────────────────────
    total = sum(counts.values())
    summary_p = bare_para(before=0, after=140)
    add_run(summary_p, 'A total of ', size_pt=10, color='333333')
    add_run(summary_p, f'{total} finding{"s" if total != 1 else ""}',
            bold=True, size_pt=10, color='00467A')
    add_run(summary_p, ' were identified during the assessment. '
            'The distribution of findings by severity is shown below.',
            size_pt=10, color='333333')
    anchor_p.addprevious(summary_p)

    # ── Severity tiles (Word table) ───────────────────────────────────────────
    SEV_ORDER  = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFORMATIONAL']
    SEV_LABELS = ['Critical', 'High', 'Medium', 'Low', 'Info']
    GAP    = 108
    tile_w = (_PAGE_W - 4 * GAP) // 5

    sev_tbl = doc.add_table(rows=1, cols=5)
    _es_tbl_props(sev_tbl._tbl, _PAGE_W, cell_spacing=GAP)

    for i, (sev, label) in enumerate(zip(SEV_ORDER, SEV_LABELS)):
        cell  = sev_tbl.rows[0].cells[i]
        tc    = cell._tc
        col   = _SEV_TILE_COLOR[sev]
        bspec = ('single', col, 12)

        _set_cell_shading(tc, _SEV_TILE_TINT[sev])
        _es_tc_width(tc, tile_w)
        _es_tc_margins(tc, top=160, right=60, bottom=130, left=60)
        _es_tc_valign(tc, 'center')
        _es_tc_borders(tc, top=bspec, right=bspec, bottom=bspec, left=bspec)

        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _es_para_spacing(p, 0, 0)
        _es_clear_list_style(p)

        r_n = p.add_run(str(counts.get(sev, 0)))
        r_n.bold = True
        r_n.font.size = Pt(22)
        r_n.font.color.rgb = _rgb(col)
        p.add_run('\n')
        r_l = p.add_run(label.upper())
        r_l.font.size = Pt(6.5)
        r_l.font.bold = False
        r_l.font.color.rgb = _rgb(col)

    insert_before(sev_tbl._tbl)

    # ── Key findings ──────────────────────────────────────────────────────────
    if not key_findings:
        anchor_p.getparent().remove(anchor_p)
        return

    anchor_p.addprevious(bare_para(before=200, after=0))

    # Sort by pentest type then severity
    sev_rank = {s: i for i, s in enumerate(
        ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFORMATIONAL'])}
    key_findings = sorted(
        key_findings,
        key=lambda f: (f.pentest_type or '', sev_rank.get(f.severity, 99))
    )

    BADGE_W   = 720
    CONTENT_W = _PAGE_W - BADGE_W
    SEP = ('single', 'DEDEDE', 4)

    current_type = None
    kf_tbl = None

    for idx, finding in enumerate(key_findings):
        sev       = finding.severity or 'INFORMATIONAL'
        badge_col = _SEV_BADGE_BG.get(sev, '666666')
        ptype     = finding.pentest_type or ''
        is_last   = idx == len(key_findings) - 1

        # ── New pentest-type group header ─────────────────────────────────────
        if ptype != current_type:
            current_type = ptype
            # Flush previous table before new group
            if kf_tbl is not None:
                insert_before(kf_tbl._tbl)

            # Group label paragraph
            type_label = finding.get_pentest_type_display() if ptype else 'Other'
            lp = bare_para(before=(160 if idx > 0 else 80), after=40)
            add_run(lp, type_label.upper(), bold=True, size_pt=8, color='00467A')
            anchor_p.addprevious(lp)

            # Start a fresh table for this group
            kf_tbl = doc.add_table(rows=0, cols=2)
            _es_tbl_props(kf_tbl._tbl, _PAGE_W)

        raw = html_to_plain(finding.description or '').strip()
        sentences = _re2.split(r'(?<=[.!?])\s+', raw)
        short = sentences[0] if sentences else ''
        if len(short) > 200:
            short = short[:197] + '...'

        # ── Row 1: badge | title ──────────────────────────────────────────────
        r1 = kf_tbl.add_row()

        b1tc = r1.cells[0]._tc
        _set_cell_shading(b1tc, 'FFFFFF')
        _es_tc_width(b1tc, BADGE_W)
        _es_tc_margins(b1tc, top=90, right=60, bottom=60, left=0)
        _es_tc_valign(b1tc, 'center')
        _es_tc_borders(b1tc)  # all none

        b1p = r1.cells[0].paragraphs[0]
        b1p.clear()
        b1p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _es_para_spacing(b1p, 0, 0)
        _es_clear_list_style(b1p)
        badge_png = _build_badge_image(sev)
        if badge_png:
            b1p.add_run().add_picture(io.BytesIO(badge_png), width=Inches(0.58))

        t1tc = r1.cells[1]._tc
        _set_cell_shading(t1tc, 'FFFFFF')
        _es_tc_width(t1tc, CONTENT_W)
        _es_tc_margins(t1tc, top=100, right=0, bottom=short and 0 or 120, left=160)
        _es_tc_valign(t1tc, 'center')
        _es_tc_borders(t1tc)  # all none — separator goes on desc row only

        t1p = r1.cells[1].paragraphs[0]
        t1p.clear()
        _es_para_spacing(t1p, 0, 0)
        _es_clear_list_style(t1p)
        rt = t1p.add_run(finding.title or '')
        rt.bold = True
        rt.font.size = Pt(10)
        rt.font.color.rgb = _rgb('1A1A1A')

        if not short:
            # No description — separator goes on the title row itself
            if not is_last:
                _es_tc_borders(b1tc, bottom=SEP)
                _es_tc_borders(t1tc, bottom=SEP)
            continue

        # ── Row 2: empty | description ───────────────────────────────────────
        r2 = kf_tbl.add_row()

        b2tc = r2.cells[0]._tc
        _set_cell_shading(b2tc, 'FFFFFF')
        _es_tc_width(b2tc, BADGE_W)
        _es_tc_margins(b2tc, 0, 0, 0, 0)
        _es_tc_borders(b2tc, bottom=(None if is_last else SEP))
        b2p = r2.cells[0].paragraphs[0]
        b2p.clear()
        _es_para_spacing(b2p, 0, 0)
        _es_clear_list_style(b2p)

        d2tc = r2.cells[1]._tc
        _set_cell_shading(d2tc, 'FFFFFF')
        _es_tc_width(d2tc, CONTENT_W)
        _es_tc_margins(d2tc, top=0, right=0, bottom=120, left=160)
        _es_tc_borders(d2tc, bottom=(None if is_last else SEP))

        d2p = r2.cells[1].paragraphs[0]
        d2p.clear()
        _es_para_spacing(d2p, 0, 0)
        _es_clear_list_style(d2p)
        rd = d2p.add_run(short)
        rd.font.size = Pt(8.5)
        rd.font.color.rgb = _rgb('666666')

    if kf_tbl is not None:
        insert_before(kf_tbl._tbl)
    anchor_p.getparent().remove(anchor_p)


# ── Chart generation ──────────────────────────────────────────────────────────

def _build_severity_chart(counts):
    """Return PNG bytes of a severity distribution bar chart, or None if matplotlib missing."""
    try:
        import matplotlib
        import matplotlib.ticker as ticker
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        # Informational → Critical (left to right)
        labels = ['Info', 'Low', 'Medium', 'High', 'Critical']
        values = [
            counts['INFORMATIONAL'], counts['LOW'], counts['MEDIUM'],
            counts['HIGH'], counts['CRITICAL'],
        ]
        colors = ['#00467A', '#49A58B', '#FFF1AA', '#E17468', '#FF0000']

        fig, ax = plt.subplots(figsize=(3.2, 1.9))
        bars = ax.bar(labels, values, color=colors, edgecolor='white', linewidth=0.8, width=0.55)
        ax.set_ylabel('Findings', fontsize=7)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.tick_params(axis='both', labelsize=7)
        max_val = max(values) if max(values) > 0 else 1
        ax.set_ylim(0, max_val + max(1, int(max_val * 0.25)))
        ax.yaxis.set_major_locator(ticker.MaxNLocator(integer=True))
        for bar, val in zip(bars, values):
            if val > 0:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + 0.03,
                    str(val),
                    ha='center', va='bottom', fontsize=7, fontweight='bold',
                )
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except ImportError:
        return None


def _build_remediation_chart(findings):
    """Return PNG bytes of a remediation priority pie chart, or None if matplotlib missing."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        immediate  = sum(1 for f in findings if f.severity in ('CRITICAL', 'HIGH'))
        short_term = sum(1 for f in findings if f.severity == 'MEDIUM')
        long_term  = sum(1 for f in findings if f.severity in ('LOW', 'INFORMATIONAL'))

        all_labels = ['Immediate (Critical / High)', 'Short-term (Medium)', 'Long-term (Low / Info)']
        all_values = [immediate, short_term, long_term]
        all_colors = ['#E17468', '#FFF1AA', '#49A58B']

        # Pie only draws non-zero slices, but legend always shows all 3
        pie_data = [(l, v, c) for l, v, c in zip(all_labels, all_values, all_colors) if v > 0]
        if not pie_data:
            return None
        pie_labels, pie_values, pie_colors = zip(*pie_data)

        fig, ax = plt.subplots(figsize=(3.2, 2.2))
        wedges, _, autotexts = ax.pie(
            pie_values,
            labels=None,
            colors=pie_colors,
            autopct=lambda p: f'{int(round(p * sum(pie_values) / 100))}' if p > 0 else '',
            startangle=90,
            wedgeprops={'edgecolor': 'white', 'linewidth': 1.5},
        )
        for at in autotexts:
            at.set_fontsize(7)
            at.set_fontweight('bold')
        # Always show full legend with all 3 tiers using patch proxies
        from matplotlib.patches import Patch
        legend_handles = [Patch(facecolor=c, edgecolor='white') for c in all_colors]
        ax.legend(
            legend_handles, all_labels,
            loc='lower center',
            bbox_to_anchor=(0.5, -0.30),
            ncol=1,
            fontsize=6.5,
            frameon=False,
        )
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except ImportError:
        return None


def _insert_chart(doc, marker, chart_png, width=Inches(3.0)):
    """Replace a paragraph containing marker with the chart image."""
    if not chart_png:
        return
    for para in _iter_all_paragraphs(doc):
        if marker in para.text:
            para.clear()
            run = para.add_run()
            run.add_picture(io.BytesIO(chart_png), width=width)
            return


# ── Table layout helpers ──────────────────────────────────────────────────────

def _fix_table_layouts(doc):
    """
    Enforce fixed column-width layout on every table in the document.
    This prevents long unbreakable text in one cell from expanding that column
    and squeezing neighbouring columns (e.g. the chart cell).
    Fixed layout honours the column widths already set in the template.
    """
    W_TBLLAYOUT = qn('w:tblLayout')
    W_TBLPR     = qn('w:tblPr')

    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(W_TBLPR)
        if tblPr is None:
            from docx.oxml import OxmlElement
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        layout = tblPr.find(W_TBLLAYOUT)
        if layout is None:
            from docx.oxml import OxmlElement
            layout = OxmlElement('w:tblLayout')
            tblPr.append(layout)
        layout.set(qn('w:type'), 'fixed')


# ── Table of Contents reset ───────────────────────────────────────────────────

def _reset_toc(doc):
    """
    Clear stale TOC entries by finding the exact field boundaries.

    Phase 1 — find the body-element index that contains the TOC field's
              <w:fldChar type="separate"/> (comes right after the instrText).
    Phase 2 — scan forward from there, tracking nested begin/end depth.
              The first 'end' at depth 0 is the TOC field's closing marker.
    Remove every body element strictly between separate and end so the field
    instruction stays intact and Word can right-click → Update Field.
    """
    W      = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _W_FLD  = f'{{{W}}}fldChar'
    _W_INST = f'{{{W}}}instrText'

    body     = doc.element.body
    children = list(body)

    # Phase 1: find the body element containing the TOC 'separate' marker
    sep_idx  = None
    toc_seen = False
    for i, elem in enumerate(children):
        for node in elem.iter():
            if node.tag == _W_INST and 'TOC' in (node.text or ''):
                toc_seen = True
            if toc_seen and node.tag == _W_FLD:
                if node.get(f'{{{W}}}fldCharType', '') == 'separate':
                    sep_idx = i
                    break
        if sep_idx is not None:
            break

    if sep_idx is None:
        return  # no TOC field found

    # Phase 2: from the element after 'separate', find the matching 'end'
    # tracking nested field depth so we don't stop at a hyperlink's end.
    nest = 0
    for i in range(sep_idx + 1, len(children)):
        for node in children[i].iter():
            if node.tag != _W_FLD:
                continue
            ftype = node.get(f'{{{W}}}fldCharType', '')
            if ftype == 'begin':
                nest += 1
            elif ftype == 'end':
                if nest == 0:
                    for e in children[sep_idx + 1 : i]:
                        body.remove(e)
                    return
                nest -= 1


# ── Attack Chain ─────────────────────────────────────────────────────────────

def _build_attack_chain_image(entries, codes):
    """
    Render a horizontal attack-chain flow diagram using matplotlib.
    Returns PNG bytes, or None if matplotlib is unavailable.
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches

        # Group by phase, preserving insertion order
        phases = {}
        for e in entries:
            phases.setdefault(e.phase, []).append(e)

        n = len(phases)
        if n == 0:
            return None

        BOX_W, BOX_H = 1.9, 1.2
        GAP = 0.55
        fig_w = n * BOX_W + (n - 1) * GAP + 1.2
        fig_h = max(3.5, BOX_H + 2.2)

        fig, ax = plt.subplots(figsize=(fig_w, fig_h))
        ax.set_xlim(0, fig_w - 1.0)
        ax.set_ylim(-0.6, fig_h - 0.2)
        ax.axis('off')
        fig.patch.set_facecolor('#FAFAFA')

        palette = [
            '#c0392b', '#e67e22', '#27ae60', '#2980b9',
            '#8e44ad', '#16a085', '#d35400', '#2c3e50',
        ]

        x = 0.2
        for idx, (phase, phase_entries) in enumerate(phases.items()):
            col = palette[idx % len(palette)]

            # Header box
            ax.add_patch(mpatches.FancyBboxPatch(
                (x, BOX_H + 0.6), BOX_W, 0.52,
                boxstyle='round,pad=0.05',
                facecolor=col, edgecolor='white', linewidth=0,
            ))
            ax.text(x + BOX_W / 2, BOX_H + 0.86, phase,
                    ha='center', va='center', fontsize=7.5,
                    fontweight='bold', color='white', clip_on=True)

            # Body box
            ax.add_patch(mpatches.FancyBboxPatch(
                (x, 0.55), BOX_W, BOX_H,
                boxstyle='round,pad=0.05',
                facecolor='#FFFFFF', edgecolor=col, linewidth=1.5,
            ))

            # Finding codes + truncated title
            lines = []
            for e in phase_entries[:4]:
                code = codes.get(e.finding_id, '?') if e.finding_id else '?'
                title = (e.finding.title or '') if e.finding else ''
                label = f'{code}  {title[:28]}' if title else code
                lines.append(label)
            if len(phase_entries) > 4:
                lines.append(f'(+{len(phase_entries)-4} more)')

            body_text = '\n'.join(lines)
            ax.text(x + BOX_W / 2, 0.55 + BOX_H / 2, body_text,
                    ha='center', va='center', fontsize=6.0, color='#333333',
                    linespacing=1.5)

            # Arrow to next phase
            if idx < n - 1:
                ax.annotate(
                    '', xy=(x + BOX_W + GAP, BOX_H + 0.86),
                    xytext=(x + BOX_W + 0.04, BOX_H + 0.86),
                    arrowprops=dict(arrowstyle='->', color=col, lw=2),
                )

            x += BOX_W + GAP

        fig.tight_layout(pad=0.3)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=180, bbox_inches='tight',
                    facecolor=fig.get_facecolor())
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except ImportError:
        return None


def _generate_phase_blurb(phase_name, finding_titles):
    """
    Call the configured AI provider (Ollama / Gemini) to produce a 2-sentence
    description of what the attacker achieved during this attack-chain phase.
    Returns plain text.  Falls back to a generic sentence on any error.
    """
    try:
        from django.conf import settings
        from findings.ai_views import call_ollama, call_gemini

        titles_str = '\n'.join(f'- {t}' for t in finding_titles)
        prompt = (
            f"You are writing an attack chain narrative for a professional penetration test report. "
            f"In exactly 2 concise sentences, describe what the attacker achieved during the "
            f"'{phase_name}' phase of this engagement, based on the following findings:\n"
            f"{titles_str}\n\n"
            f"Be specific about the technique and its significance. Write in past tense. "
            f"Do not use bullet points, headings, or labels — just the 2 sentences."
        )

        provider = getattr(settings, 'AI_PROVIDER', 'ollama')
        blurb = call_gemini(prompt) if provider == 'gemini' else call_ollama(prompt)
        return blurb.strip() if blurb else ''
    except Exception:
        return ''


def _build_attack_chain(doc, report_obj, finding_codes):
    """
    Replace <<attack_chain_cards>> with a native Word table attack-chain.

    Layout: one column per phase separated by narrow arrow columns.
      Row 0 — phase name header (dark blue, white bold text)
      Row 1 — findings list (light grey, code + truncated title per line)
    """
    anchor = None
    for para in _iter_all_paragraphs(doc):
        if '<<attack_chain_cards>>' in para.text:
            anchor = para
            break
    if anchor is None:
        return

    entries = []
    try:
        from .models import AttackChainEntry
        entries = list(
            AttackChainEntry.objects.filter(report=report_obj)
            .select_related('finding')
            .order_by('phase', 'position')
        )
    except Exception:
        pass

    if not entries:
        anchor.clear()
        anchor.add_run('No attack chain data has been defined for this report.')
        return

    # Group by phase, preserving insertion order
    phases = {}
    for e in entries:
        phases.setdefault(e.phase, []).append(e)

    n          = len(phases)
    ARROW_W    = 320    # twips — narrow arrow separator column
    TOTAL_W    = 9360   # twips — 6.5" content width
    phase_w    = (TOTAL_W - (n - 1) * ARROW_W) // n if n > 1 else TOTAL_W
    n_cols     = 2 * n - 1 if n > 1 else 1

    col_widths = []
    for i in range(n):
        col_widths.append(phase_w)
        if i < n - 1:
            col_widths.append(ARROW_W)

    # Phase header colour palette (dark blues / teals matching OziCyber palette)
    HEADER_COLORS = [
        '00467A', '1F4E79', '2E6DA4', '1A5276',
        '0D47A1', '154360', '1B4F72', '375A7F',
    ]

    # Pre-generate AI blurbs for all phases (one AI call per phase)
    phase_blurbs = {}
    for phase_name, phase_entries in phases.items():
        titles = [
            (e.finding.title or '') for e in phase_entries if e.finding
        ]
        blurb = _generate_phase_blurb(phase_name, titles)
        if not blurb:
            blurb = (
                f'During the {phase_name} phase, the assessment identified '
                f'{len(phase_entries)} finding{"s" if len(phase_entries) != 1 else ""} '
                f'that contributed to this attack path.'
            )
        phase_blurbs[phase_name] = blurb

    # 3 rows: 0=phase header, 1=AI blurb, 2=finding codes
    tbl = doc.add_table(rows=3, cols=n_cols)

    def _tc_no_borders(tc):
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        bdr = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            bdr.append(b)
        tcPr.append(bdr)

    for i, (phase_name, phase_entries) in enumerate(phases.items()):
        col_i   = i * 2
        hdr_col = HEADER_COLORS[i % len(HEADER_COLORS)]

        # ── Row 0: phase header ───────────────────────────────────────────────
        hdr_tc = tbl.rows[0].cells[col_i]._tc
        _set_cell_shading(hdr_tc, hdr_col)
        _es_tc_width(hdr_tc, phase_w)
        _es_tc_margins(hdr_tc, top=100, right=80, bottom=100, left=80)
        _tc_no_borders(hdr_tc)

        hdr_p = tbl.rows[0].cells[col_i].paragraphs[0]
        hdr_p.clear()
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _es_para_spacing(hdr_p, 0, 0)
        r = hdr_p.add_run(phase_name.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # ── Row 1: AI blurb ──────────────────────────────────────────────────
        blurb_tc = tbl.rows[1].cells[col_i]._tc
        _set_cell_shading(blurb_tc, 'EAF0F7')
        _es_tc_width(blurb_tc, phase_w)
        _es_tc_margins(blurb_tc, top=90, right=90, bottom=90, left=90)
        _tc_no_borders(blurb_tc)

        blurb_p = tbl.rows[1].cells[col_i].paragraphs[0]
        blurb_p.clear()
        _es_para_spacing(blurb_p, 0, 0)
        rb = blurb_p.add_run(phase_blurbs.get(phase_name, ''))
        rb.italic = True
        rb.font.size = Pt(7.5)
        rb.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # ── Row 2: finding codes ─────────────────────────────────────────────
        cnt_tc = tbl.rows[2].cells[col_i]._tc
        _set_cell_shading(cnt_tc, 'F2F2F2')
        _es_tc_width(cnt_tc, phase_w)
        _es_tc_margins(cnt_tc, top=80, right=80, bottom=80, left=80)
        _tc_no_borders(cnt_tc)

        cnt_p = tbl.rows[2].cells[col_i].paragraphs[0]
        cnt_p.clear()
        _es_para_spacing(cnt_p, 0, 0)

        for j, entry in enumerate(phase_entries):
            code  = finding_codes.get(entry.finding_id, '') if entry.finding_id else ''
            title = (entry.finding.title or '') if entry.finding else ''
            line  = f'{code}  {title}' if title else code
            if len(line) > 55:
                line = line[:52] + '…'
            if j > 0:
                cnt_p.add_run('\n')
            run = cnt_p.add_run(f'• {line}')
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        # ── Arrow separator columns ───────────────────────────────────────────
        if i < n - 1:
            for row_idx in range(3):
                arr_tc = tbl.rows[row_idx].cells[col_i + 1]._tc
                _set_cell_shading(arr_tc, 'FFFFFF')
                _es_tc_width(arr_tc, ARROW_W)
                _es_tc_margins(arr_tc, top=0, right=0, bottom=0, left=0)
                _tc_no_borders(arr_tc)

                arr_p = tbl.rows[row_idx].cells[col_i + 1].paragraphs[0]
                arr_p.clear()
                arr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _es_para_spacing(arr_p, 0, 0)
                if row_idx == 0:  # arrow only in header row
                    ar = arr_p.add_run('→')
                    ar.font.size = Pt(14)
                    ar.bold = True
                    ar.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Move table from doc body to position of the anchor paragraph
    tbl._tbl.getparent().remove(tbl._tbl)
    anchor._p.addprevious(tbl._tbl)
    anchor._p.getparent().remove(anchor._p)


def _process_flat_finding_list(doc, finding_rows):
    """
    Find the single <<fl.>> template table and replace it with one cloned table
    per pentest type, each containing only that type's findings, preceded by a heading.

    Also finds and removes the static sub-heading that precedes the template table
    in the source template (e.g. "2.1. Web Application Penetration Test Findings"),
    clones its XML to preserve the outline numPr so Word auto-numbers the generated
    headings as 2.1, 2.2, etc.
    """
    _W_TR = f'{{{_W_NS}}}tr'
    _W_R  = f'{{{_W_NS}}}r'
    marker = '<<fl.'

    # Locate the template table and its template rows
    tmpl_table = None
    first_idx = last_idx = None
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            if any(marker in cell.text for cell in row.cells):
                if tmpl_table is None:
                    tmpl_table = table
                    first_idx = i
                last_idx = i
        if tmpl_table is not None:
            break

    if tmpl_table is None:
        return

    tmpl_trs = [deepcopy(row._tr) for row in tmpl_table.rows[first_idx:last_idx + 1]]
    tbl_elem = tmpl_table._tbl
    body = tbl_elem.getparent()

    # Walk backwards from the template table through body children.
    # Find the static sub-heading (e.g. "2.1. Web Application…"), clone its XML
    # to preserve numPr/style, then remove it.  Skip blank paragraphs; stop at
    # any Heading 1 (section boundary) or non-paragraph element.
    static_heading_tmpl = None
    children = list(body)
    tbl_body_idx = next(i for i, c in enumerate(children) if c is tbl_elem)

    for elem in reversed(children[:tbl_body_idx]):
        if elem.tag != _W_P:
            break
        pPr    = elem.find(f'{{{_W_NS}}}pPr')
        pStyle = pPr.find(f'{{{_W_NS}}}pStyle') if pPr is not None else None
        val    = (pStyle.get(f'{{{_W_NS}}}val', '') if pStyle is not None else '').lower()
        if val in ('heading1', '1'):
            break  # section boundary — stop
        if 'heading' in val:
            static_heading_tmpl = deepcopy(elem)
            body.remove(elem)
            break
        # blank or body paragraph — keep looking

    # Group findings by pentest type (preserve sorted order)
    groups: dict = {}
    for item in finding_rows:
        groups.setdefault(item.get('pentest_type_code', ''), []).append(item)

    for idx, group_items in enumerate(groups.values()):
        type_display = group_items[0].get('pentest_type') or 'Other'
        heading_text = f'{type_display} Findings'

        # Spacer paragraph between tables and the next sub-heading
        if idx > 0:
            spacer = OxmlElement('w:p')
            sp_pPr = OxmlElement('w:pPr')
            sp_spc = OxmlElement('w:spacing')
            sp_spc.set(qn('w:before'), '160')
            sp_spc.set(qn('w:after'), '0')
            sp_pPr.append(sp_spc)
            spacer.append(sp_pPr)
            tbl_elem.addprevious(spacer)

        if static_heading_tmpl is not None:
            # Clone the static heading — preserves numPr so Word auto-numbers (2.1, 2.2…)
            h_elem = deepcopy(static_heading_tmpl)
            # Replace all direct run children with a single clean run containing new text
            for r in [c for c in list(h_elem) if c.tag == _W_R]:
                h_elem.remove(r)
            r_new = OxmlElement('w:r')
            t_new = OxmlElement('w:t')
            t_new.text = heading_text
            r_new.append(t_new)
            h_elem.append(r_new)
            tbl_elem.addprevious(h_elem)
        else:
            for style_name in ('Heading 2', 'Heading 3', 'Heading 1'):
                try:
                    h = doc.add_paragraph(heading_text, style=style_name)
                    h._element.getparent().remove(h._element)
                    tbl_elem.addprevious(h._element)
                    break
                except Exception:
                    continue

        # Clone the template table and fill with this group's findings only
        tbl_clone = deepcopy(tbl_elem)
        for item in group_items:
            for tmpl_tr in tmpl_trs:
                clone_tr = deepcopy(tmpl_tr)
                _fill_element_markers(clone_tr, item, 'fl')
                _fill_element_markers(clone_tr, item, 'f')  # handle mixed <<f.xxx>> markers (e.g. severity cell)
                tbl_clone.append(clone_tr)

        # Remove the template rows from the clone (keep header row)
        clone_trs = [c for c in tbl_clone if c.tag == _W_TR]
        for tr in clone_trs[first_idx:last_idx + 1]:
            tbl_clone.remove(tr)

        tbl_elem.addprevious(tbl_clone)

    body.remove(tbl_elem)


def _remove_detailed_findings_heading(doc):
    """Remove the static 'Detailed Findings' heading that duplicates the generator-inserted type headings."""
    for para in list(doc.paragraphs):
        if para.text.strip().lower() in ('detailed findings', 'detailed finding'):
            style_name = (para.style.name or '').lower()
            if 'heading' in style_name:
                para._element.getparent().remove(para._element)


# ── Public API ────────────────────────────────────────────────────────────────

def generate_report_docx(report_obj, findings_qs, template_path=None):
    """
    Render a python-docx Word template with report data.
    Returns a BytesIO containing the populated .docx.
    """
    path = template_path or DEFAULT_TEMPLATE_PATH
    if not os.path.exists(path):
        raise FileNotFoundError(f'Template not found: {path}')

    doc = Document(path)

    # Normalise all <<MARKER>> text to lowercase so templates are case-insensitive.
    # <<F.DESCRIPTION>>, <<f.description>>, <<F.Description>> all work identically.
    _normalize_markers(doc)

    engagement = report_obj.engagement
    org        = engagement.organization if engagement else None
    findings   = sorted(
        list(findings_qs),
        key=lambda f: (
            f.pentest_type or '',
            SEVERITIES.index(f.severity) if f.severity in SEVERITIES else 99,
        )
    )
    codes      = build_finding_codes(findings)

    counts = {s: 0 for s in SEVERITIES}
    for f in findings:
        if f.severity in counts:
            counts[f.severity] += 1

    scope_raw = (engagement.scope or '') if engagement else ''
    scope_items = [s.strip() for s in scope_raw.split('\n') if s.strip()] or ['—']

    out_of_scope_raw = (engagement.out_of_scope or '') if engagement else ''
    out_of_scope_items = [s.strip() for s in out_of_scope_raw.split('\n') if s.strip()] or []

    lead = engagement.lead_pentester if engagement else None
    pm   = engagement.project_manager if engagement else None
    today_str   = date.today().strftime('%d %B %Y')
    author_name = lead.get_full_name() if lead else (
        report_obj.generated_by.get_full_name() if report_obj.generated_by else 'OziCyber'
    )

    # ── Scalar replacements ───────────────────────────────────────────────────
    replacements = {
        # Report metadata
        'report_title':     report_obj.title or 'Penetration Test Report',
        'report_version':   report_obj.version or '1.0',
        'version':          report_obj.version or '1.0',
        'is_draft':         'DRAFT' if report_obj.is_draft else 'FINAL',
        'today':            today_str,
        'generated_by':     report_obj.generated_by.get_full_name() if report_obj.generated_by else '',

        # Engagement
        'project_id':           f'ENG-{engagement.id:04d}' if engagement else '',
        'engagement_name':      engagement.name if engagement else '',
        'engagement_type':      engagement.get_engagement_type_display() if engagement else '',
        'engagement_status':    engagement.get_status_display() if engagement else '',
        'engagement_desc':      engagement.description or '' if engagement else '',
        'engagement_start':     _fmt_date(engagement.start_date) if engagement else '',
        'engagement_end':       _fmt_date(engagement.end_date) if engagement else '',
        'report_due_date':      _fmt_date(engagement.report_due_date) if engagement else '',
        'objectives':           engagement.objectives or '' if engagement else '',

        # Client / org
        'client_name':          org.name if org else '',
        'org_name':             org.name if org else '',
        'org_website':          org.website or '' if org else '',
        'org_phone':            org.phone or '' if org else '',
        'org_address':          org.address or '' if org else '',
        'client_contact_name':  engagement.client_name or '' if engagement else '',
        'client_contact_email': engagement.client_email or '' if engagement else '',
        'client_contact_phone': engagement.client_phone or '' if engagement else '',

        # Team
        'author_name':           author_name,
        'lead_tester':           lead.get_full_name() if lead else '',
        'lead_tester_email':     lead.email if lead else '',
        'project_manager':       pm.get_full_name() if pm else '',
        'project_manager_email': pm.email if pm else '',

        # Report content
        'executive_summary': html_to_plain(report_obj.executive_summary or ''),
        'methodology':       html_to_plain(report_obj.methodology or ''),
        'conclusion':        html_to_plain(report_obj.conclusion or ''),
        'client_notes':      html_to_plain(report_obj.client_notes or ''),

        # Scope (joined for inline display; loop table uses <<s.item>>)
        'scope_raw':            scope_raw,
        'scope_list':           '\n'.join(f'• {s}' for s in scope_items),
        'out_of_scope_list':    '\n'.join(f'• {s}' for s in out_of_scope_items),

        # Finding counts
        'finding_count':   str(len(findings)),
        'critical_count':  str(counts['CRITICAL']),
        'high_count':      str(counts['HIGH']),
        'medium_count':    str(counts['MEDIUM']),
        'low_count':       str(counts['LOW']),
        'info_count':      str(counts['INFORMATIONAL']),
        'overall_risk':    _overall_risk(counts),

        # Aliases for templates that use variant/legacy names
        'client':             org.name if org else '',
        'type_of_test':       engagement.get_engagement_type_display() if engagement else '',
        'web_application':    engagement.get_engagement_type_display() if engagement else '',
        'date':               today_str,
        'projectid':          f'ENG-{engagement.id:04d}' if engagement else '',
        'author':             author_name,
        'exec_summary':       html_to_plain(report_obj.executive_summary or ''),
        'concessions':        '',   # placeholder — no matching model field yet
        'constraints':        '',   # placeholder — no matching model field yet
        'disclaimer': (
            'This report has been prepared by OziCyber and is intended solely for the use of the '
            'organisation named above. The contents of this report are confidential and must not be '
            'disclosed to any third party without prior written consent. The findings and '
            'recommendations are based on information available at the time of the assessment and '
            'reflect the security posture at that point in time only.'
        ),
    }

    # ── Loop data ─────────────────────────────────────────────────────────────
    finding_rows = []
    for f in findings:
        finding_rows.append({
            'code':               codes.get(f.id, '—'),
            'title':              f.title or '',
            'severity':           f.severity or '',
            'severity_display':   SEV_DISPLAY.get(f.severity, f.severity or ''),
            'impact_rating':      f.impact_rating or '',
            'likelihood_rating':  f.likelihood_rating or '',
            'consequence':        CONSEQUENCE_MAP.get(f.impact_rating or '', '—'),
            'likelihood_label':   LIKELIHOOD_MAP.get(f.likelihood_rating or '', '—'),
            'pentest_type':       f.get_pentest_type_display() if f.pentest_type else '',
            'pentest_type_code':  f.pentest_type or '',
            'status':             f.status or '',
            'status_display':     STATUS_DISPLAY.get(f.status, f.status or ''),
            'affected_asset':     f.affected_asset or '',
            'description':        html_to_plain(f.description or ''),
            'details':            html_to_plain(f.details or ''),
            'impact':             html_to_plain(f.impact or ''),
            'likelihood':         html_to_plain(f.likelihood or ''),
            'recommendations':    html_to_plain(f.recommendations or ''),
            'supporting_evidence': html_to_plain(f.supporting_evidence or ''),
            '_images':             _extract_images_from_html(f.supporting_evidence or ''),
            'references':         f.references or '',
            'cwe_id':             f.cwe_id or '',
            'cve_id':             f.cve_id or '',
            'level_of_access':     f.get_level_of_access_display() if f.level_of_access else '',
            'cvss_score':         str(f.cvss_score) if f.cvss_score else '',
            'cvss_vector':        f.cvss_vector or '',
            'cvss_av':            _cvss_display('av', f.av),
            'cvss_ac':            _cvss_display('ac', f.ac),
            'cvss_pr':            _cvss_display('pr', f.pr),
            'cvss_ui':            _cvss_display('ui', f.ui),
            'cvss_s':             _cvss_display('s',  f.s),
            'cvss_c':             _cvss_display('c',  f.c),
            'cvss_i':             _cvss_display('i',  f.i),
            'cvss_a':             _cvss_display('a',  f.a),
            'created_at':         _fmt_date(f.created_at.date() if f.created_at else None),
            'remediated_at':      _fmt_date(f.remediated_at.date() if f.remediated_at else None),
        })

    remediation_rows = [
        {
            'code':       codes.get(f.id, '—'),
            'title':      f.title or '',
            'severity':   SEV_DISPLAY.get(f.severity, f.severity or ''),
            'remediated': 'Yes' if f.status == 'REMEDIATED' else 'No',
            'date':       _fmt_date(f.remediated_at.date() if f.remediated_at else None, '%d/%m/%Y'),
            'status':     STATUS_DISPLAY.get(f.status, f.status or ''),
        }
        for f in findings
    ]

    version_rows = [{
        'version': report_obj.version or '1.0',
        'date':    today_str,
        'author':  author_name,
        'comment': 'Generated by OziReport',
    }]

    scope_rows = [{'item': s} for s in scope_items]

    key_findings = [f for f in findings if getattr(f, 'is_key_finding', False)]
    key_finding_rows    = [{'title': f.title or ''} for f in key_findings]
    key_mitigation_rows = [{'title': _short_mitigation(f.recommendations)} for f in key_findings]

    # ── Apply to document ─────────────────────────────────────────────────────
    _replace_all_scalars(doc, replacements)

    # Remove static 'Detailed Findings' heading before inserting generated type headings
    _remove_detailed_findings_heading(doc)

    # Flat finding list — one table per pentest type in the List of Vulnerabilities section.
    _process_flat_finding_list(doc, finding_rows)

    # Body-block loops first (<<f:start>> / <<f:end>> etc.) — handles sections
    # where a finding block mixes tables and paragraphs together.
    _process_body_block(doc, 'f', finding_rows)
    _process_body_block(doc, 'r', remediation_rows)
    _process_body_block(doc, 'v', version_rows)
    _process_body_block(doc, 's', scope_rows)
    _process_body_block(doc, 'kf', key_finding_rows)
    _process_body_block(doc, 'km', key_mitigation_rows)

    # Table-row loops — handles simple single/multi-row table repetition.
    _process_loop_tables(doc, 'f', finding_rows)
    _process_loop_tables(doc, 'r', remediation_rows)
    _process_loop_tables(doc, 'v', version_rows)
    _process_loop_tables(doc, 's', scope_rows)
    _process_cell_lists(doc, 'kf', key_finding_rows)
    _process_cell_lists(doc, 'km', key_mitigation_rows)

    # Image markers — replace <<f.supporting_evidence_img>> paragraphs with screenshots.
    _replace_image_markers(doc, 'f', finding_rows)

    _build_exec_summary_block(doc, counts, key_findings)
    _fix_table_layouts(doc)
    _insert_chart(doc, '<<severity_chart>>', _build_severity_chart(counts))
    _insert_chart(doc, '<<remediation_chart>>', _build_remediation_chart(findings))

    _reset_toc(doc)
    _build_attack_chain(doc, report_obj, codes)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
