"""
GRC Gap Analysis Report — DOCX Generator
==========================================
Generates a professional Word (.docx) gap analysis report for a GRC project.
Mirrors the structure and narrative content of report_generator.py (HTML report).

Usage:
    from grc.docx_generator import generate_gap_analysis_docx
    docx_bytes = generate_gap_analysis_docx(project, control_statuses, family_stats)

Author: OziCyber Security Platform
"""

import io
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Import shared constants and helpers from report_generator
# ---------------------------------------------------------------------------
from .report_generator import (
    _generate_finding_description,
    _generate_recommendation,
    _generate_gap_analysis_text,
    _gather_policy_evidence,
    _get_evidence_list,
    _get_findings,
    _compute_overall_stats,
    _get_accent,
    _fmt_date,
    _pct_color,
    _risk_label_from_pct,
    _truncate,
    _first_sentence,
    EXPECTED_POLICIES,
    STATUS_LABEL,
    STATUS_COLOR,
    RISK_FOR_STATUS,
    RISK_COLOR,
    FRAMEWORK_ACCENT,
    GLOSSARY_TERMS,
    FRAMEWORK_DESCRIPTION,
)

# ---------------------------------------------------------------------------
# Color constants
# ---------------------------------------------------------------------------
NAVY = '0d2137'
WHITE = 'FFFFFF'
LIGHT_GREY = 'F5F5F5'
ALT_ROW = 'F0F4F8'
BODY_TEXT = '1a1a2e'
DARK_GREY = '333333'
MID_GREY = '555555'
LIGHT_BORDER = 'CCCCCC'

# A4 (21 cm) minus 2 × 2.5 cm margins = 16 cm content width
CONTENT_WIDTH = Cm(16)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_color(cell, hex_color: str):
    """Set the background fill colour of a table cell."""
    hex_color = hex_color.lstrip('#')
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(existing)
    tc_pr.append(shd)


def _set_cell_vertical_align(cell, align='center'):
    """Set vertical alignment of a table cell (top/center/bottom)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement('w:vAlign')
    v_align.set(qn('w:val'), align)
    for existing in tc_pr.findall(qn('w:vAlign')):
        tc_pr.remove(existing)
    tc_pr.append(v_align)


def _set_table_borders(table, color='CCCCCC', size=4):
    """Apply uniform borders to a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    # Remove existing tblBorders
    for existing in tbl_pr.findall(qn('w:tblBorders')):
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color.lstrip('#'))
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def _remove_table_borders(table):
    """Remove all borders from a table (invisible/borderless table)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    for existing in tbl_pr.findall(qn('w:tblBorders')):
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def _set_cell_borders(cell, color='CCCCCC', size=4, sides=None):
    """Set borders on individual sides of a cell."""
    if sides is None:
        sides = ('top', 'left', 'bottom', 'right')
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in sides:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color.lstrip('#'))
        tc_borders.append(border)
    for existing in tc_pr.findall(qn('w:tcBorders')):
        tc_pr.remove(existing)
    tc_pr.append(tc_borders)


def _set_table_width(table, width_dxa=9072):
    """Set table to a fixed width in twentieths-of-a-point (dxa)."""
    tbl = table._tbl
    tbl_pr = tbl.get_or_add_tblPr()
    for existing in tbl_pr.findall(qn('w:tblW')):
        tbl_pr.remove(existing)
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), str(width_dxa))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_w)


def _set_col_width(cell, width_dxa):
    """Set a cell's column width in dxa."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')
    for existing in tc_pr.findall(qn('w:tcW')):
        tc_pr.remove(existing)
    tc_pr.append(tc_w)


def _set_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _add_colored_run(paragraph, text, color_hex, bold=False, italic=False, size=None):
    """Add a run with a specific colour."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color_hex.lstrip('#'))
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return run


def _add_rule(doc, color=NAVY, thickness=6):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color.lstrip('#'))
    pb.append(bottom)
    p_pr.append(pb)
    _set_spacing(p, before=0, after=0)
    return p


def _add_page_number(paragraph):
    """Insert a PAGE field into the paragraph."""
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def _set_paragraph_font(paragraph, name='Calibri', size=11, color=BODY_TEXT,
                        bold=False, italic=False, line_spacing=1.15):
    """Apply default body font settings to a paragraph."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))
        run.bold = bold
        run.italic = italic


def _style_run(run, name='Calibri', size=11, color=BODY_TEXT, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))
    run.bold = bold
    run.italic = italic


def _make_table_full_width(table):
    """Set table to 100% page width (A4 with 2.5cm margins ≈ 9072 dxa for full content area)."""
    _set_table_width(table, 9072)


def _add_heading1(doc, text, accent=None):
    """Add a Heading 1 paragraph: Calibri 18pt bold navy."""
    p = doc.add_paragraph()
    _set_spacing(p, before=24, after=12)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run.bold = True
    return p


def _add_heading2(doc, text, accent=NAVY):
    """Add a Heading 2 paragraph: Calibri 14pt bold accent colour."""
    p = doc.add_paragraph()
    _set_spacing(p, before=18, after=8)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(accent.lstrip('#'))
    run.bold = True
    return p


def _add_heading3(doc, text):
    """Add a Heading 3 paragraph: Calibri 12pt bold dark grey."""
    p = doc.add_paragraph()
    _set_spacing(p, before=12, after=6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(DARK_GREY)
    run.bold = True
    return p


def _add_body_paragraph(doc, text='', bold=False, italic=False, color=BODY_TEXT,
                         size=11, before=0, after=6):
    """Add a body paragraph with Calibri 11pt."""
    p = doc.add_paragraph()
    _set_spacing(p, before=before, after=after, line=1.15)
    if text:
        run = p.add_run(text)
        _style_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def _add_bullet(doc, text, indent_level=0, color=BODY_TEXT, size=10):
    """Add a bullet list paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
    _set_spacing(p, before=2, after=2)
    run = p.add_run(text)
    _style_run(run, size=size, color=color)
    return p


def _add_numbered(doc, text, color=BODY_TEXT, size=10):
    """Add a numbered list paragraph."""
    p = doc.add_paragraph(style='List Number')
    _set_spacing(p, before=2, after=4)
    run = p.add_run(text)
    _style_run(run, size=size, color=color)
    return p


def _header_row(table, headers, col_widths=None):
    """Style the first row of a table as a navy header with white bold text."""
    row = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(row.cells, headers)):
        _set_cell_color(cell, NAVY)
        _set_cell_vertical_align(cell, 'center')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(p, before=3, after=3)
        run = p.add_run(hdr)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        if col_widths and i < len(col_widths):
            _set_col_width(cell, col_widths[i])


def _cell_text(cell, text, bold=False, italic=False, color=BODY_TEXT,
               size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Set text content of a cell with styling."""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = align
    _set_spacing(p, before=2, after=2)
    run = p.add_run(str(text) if text is not None else '')
    _style_run(run, size=size, color=color, bold=bold, italic=italic)
    return run


def _alt_row_shade(table, start_row=1, even_color=ALT_ROW, odd_color=WHITE):
    """Apply alternating row shading starting from start_row."""
    for i, row in enumerate(table.rows[start_row:], start=0):
        color = even_color if i % 2 == 0 else odd_color
        for cell in row.cells:
            _set_cell_color(cell, color)


def _pct_bar_text(pct, width=20):
    """Return a simple filled/empty block text for progress bar."""
    filled = round(pct / 100 * width)
    empty = width - filled
    return '█' * filled + '░' * empty


def _section_title_table(doc, section_num, title, subtitle, accent):
    """
    Render a section header as a full-width navy table row (mimicking HTML section headers).
    """
    table = doc.add_table(rows=1, cols=2, width=CONTENT_WIDTH)
    _make_table_full_width(table)
    _remove_table_borders(table)
    row = table.rows[0]

    # Left cell: section number
    left = row.cells[0]
    _set_col_width(left, 900)
    _set_cell_color(left, NAVY)
    _set_cell_vertical_align(left, 'center')
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=8, after=8)
    run = p.add_run(str(section_num))
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(accent.lstrip('#'))
    run.bold = True

    # Right cell: title + subtitle
    right = row.cells[1]
    _set_cell_color(right, NAVY)
    _set_cell_vertical_align(right, 'center')
    right.paragraphs[0].clear()
    p1 = right.paragraphs[0]
    _set_spacing(p1, before=6, after=2)
    r1 = p1.add_run(title)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(255, 255, 255)
    r1.bold = True

    p2 = right.add_paragraph()
    _set_spacing(p2, before=0, after=6)
    r2 = p2.add_run(subtitle)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(180, 180, 180)
    r2.italic = True

    # Bottom accent border on right cell
    _set_cell_borders(right, color=accent.lstrip('#'), size=8, sides=('bottom',))

    return table


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def _setup_document() -> Document:
    """Create a new Document with A4 page size and margins."""
    doc = Document()
    section = doc.sections[0]

    # A4 page size
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # Margins: 2.5cm all around
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Remove any default header/footer distance weirdness
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    return doc


def _setup_header_footer(doc, project_title: str, generated_date: str):
    """Add header and footer to the default (non-first-page) section."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True  # cover page gets no header

    # Header
    header = section.header
    header.is_linked_to_previous = False
    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()

    # Build a 3-column header table
    tbl = header.add_table(rows=1, cols=3, width=CONTENT_WIDTH)
    _make_table_full_width(tbl)
    _remove_table_borders(tbl)
    row = tbl.rows[0]

    # Left: CONFIDENTIAL
    left = row.cells[0]
    _set_col_width(left, 2000)
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run('CONFIDENTIAL')
    lr.font.name = 'Calibri'
    lr.font.size = Pt(8)
    lr.font.color.rgb = RGBColor(231, 76, 60)
    lr.bold = True

    # Centre: page number
    centre = row.cells[1]
    cp = centre.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run('Page ')
    cr.font.name = 'Calibri'
    cr.font.size = Pt(8)
    cr.font.color.rgb = RGBColor(150, 150, 150)
    _add_page_number(cp)

    # Right: project title
    right = row.cells[2]
    _set_col_width(right, 3000)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(project_title[:60])
    rr.font.name = 'Calibri'
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor(100, 100, 100)

    # Header bottom border
    for cell in row.cells:
        _set_cell_borders(cell, color='CCCCCC', size=4, sides=('bottom',))

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f'© OziCyber | Information Security Gap Analysis Report | Generated {generated_date}'
    )
    fr.font.name = 'Calibri'
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(130, 130, 130)
    fr.italic = True


# ---------------------------------------------------------------------------
# Cover Page
# ---------------------------------------------------------------------------

def _build_cover_page(doc, project, stats: dict, accent: str, generated_date: str):
    """Build the cover page (first page, no header/footer)."""
    framework_name = getattr(project.framework, 'name', 'N/A')
    framework_key = getattr(project.framework, 'key', '')
    assessor = getattr(project, 'assessor_name', None) or 'OziCyber Security Team'
    proj_status = getattr(project, 'status', 'ACTIVE') or 'ACTIVE'
    target_date_str = _fmt_date(getattr(project, 'target_date', None)) or 'Not specified'
    pct = stats['pct']
    risk_label = _risk_label_from_pct(pct)

    # Top bar: navy table with OziCyber | CONFIDENTIAL
    top_table = doc.add_table(rows=1, cols=2, width=CONTENT_WIDTH)
    _make_table_full_width(top_table)
    _remove_table_borders(top_table)
    top_row = top_table.rows[0]

    left = top_row.cells[0]
    _set_cell_color(left, NAVY)
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(lp, before=8, after=8)
    lr = lp.add_run('OziCyber')
    lr.font.name = 'Calibri'
    lr.font.size = Pt(24)
    lr.font.color.rgb = RGBColor(255, 255, 255)
    lr.bold = True

    right = top_row.cells[1]
    _set_cell_color(right, NAVY)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_spacing(rp, before=8, after=8)
    rr = rp.add_run('CONFIDENTIAL')
    rr.font.name = 'Calibri'
    rr.font.size = Pt(12)
    rr.font.color.rgb = RGBColor(231, 76, 60)
    rr.bold = True

    # Spacer
    for _ in range(3):
        sp = doc.add_paragraph()
        _set_spacing(sp, before=0, after=4)

    # Framework chip
    fw_p = doc.add_paragraph()
    fw_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(fw_p, before=0, after=4)
    fw_r = fw_p.add_run(framework_name)
    fw_r.font.name = 'Calibri'
    fw_r.font.size = Pt(13)
    fw_r.font.color.rgb = RGBColor.from_string(accent.lstrip('#'))
    fw_r.bold = True

    # "INFORMATION SECURITY ASSESSMENT" label
    label_p = doc.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(label_p, before=4, after=4)
    label_r = label_p.add_run('INFORMATION SECURITY ASSESSMENT')
    label_r.font.name = 'Calibri'
    label_r.font.size = Pt(10)
    label_r.font.color.rgb = RGBColor(150, 150, 150)
    label_r.bold = True

    # Main title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(title_p, before=4, after=2)
    tr = title_p.add_run('Gap Analysis Report')
    tr.font.name = 'Calibri'
    tr.font.size = Pt(32)
    tr.font.color.rgb = RGBColor.from_string(NAVY)
    tr.bold = True

    # Subtitle: project title
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(sub_p, before=2, after=12)
    sr = sub_p.add_run(project.title or 'Untitled Project')
    sr.font.name = 'Calibri'
    sr.font.size = Pt(18)
    sr.font.color.rgb = RGBColor.from_string(NAVY)

    # Accent rule
    rule_table = doc.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
    _make_table_full_width(rule_table)
    _remove_table_borders(rule_table)
    rule_cell = rule_table.rows[0].cells[0]
    _set_cell_color(rule_cell, accent.lstrip('#'))
    rule_p = rule_cell.paragraphs[0]
    _set_spacing(rule_p, before=3, after=3)
    rule_p.add_run('')

    # Spacer
    sp2 = doc.add_paragraph()
    _set_spacing(sp2, before=0, after=8)

    # Project details table (borderless, 2 cols)
    details_table = doc.add_table(rows=7, cols=2, width=CONTENT_WIDTH)
    _make_table_full_width(details_table)
    _remove_table_borders(details_table)

    details_data = [
        ('Organisation:', project.title or 'N/A'),
        ('Assessment Framework:', framework_name),
        ('Assessment Status:', proj_status),
        ('Lead Assessor:', assessor),
        ('Report Date:', generated_date),
        ('Target Completion:', target_date_str),
        ('Classification:', 'CONFIDENTIAL'),
    ]

    for i, (label, value) in enumerate(details_data):
        row = details_table.rows[i]
        label_cell = row.cells[0]
        _set_col_width(label_cell, 2000)
        lp = label_cell.paragraphs[0]
        _set_spacing(lp, before=3, after=3)
        lr = lp.add_run(label)
        lr.font.name = 'Calibri'
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(100, 100, 100)
        lr.bold = True

        val_cell = row.cells[1]
        vp = val_cell.paragraphs[0]
        _set_spacing(vp, before=3, after=3)
        # Special colouring for classification
        if label == 'Classification:':
            vr = vp.add_run(value)
            vr.font.name = 'Calibri'
            vr.font.size = Pt(10)
            vr.font.color.rgb = RGBColor(231, 76, 60)
            vr.bold = True
        else:
            vr = vp.add_run(value)
            vr.font.name = 'Calibri'
            vr.font.size = Pt(10)
            vr.font.color.rgb = RGBColor.from_string(NAVY)

    # Compliance posture line
    risk_color_hex = _pct_color(pct).lstrip('#')
    sp3 = doc.add_paragraph()
    _set_spacing(sp3, before=12, after=4)
    sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_colored_run(sp3, f'{pct}% — {risk_label}', risk_color_hex, bold=True, size=14)

    # Page break (end of cover)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 1: Executive Summary
# ---------------------------------------------------------------------------

def _build_executive_summary(doc, project, stats: dict, control_statuses, accent: str):
    framework_name = getattr(project.framework, 'name', 'N/A')
    pct = stats['pct']
    risk_label = _risk_label_from_pct(pct)
    risk_color = _pct_color(pct).lstrip('#')
    total = stats['total']
    implemented = stats['implemented']
    partial = stats['partial']
    in_progress = stats['in_progress']
    planned = stats['planned']
    not_applicable = stats['not_applicable']
    not_started = stats['not_started']
    evidence_total = stats['evidence_total']
    effective = stats['effective']

    def _pct_of(n):
        return f'{round(n / total * 100, 1)}%' if total > 0 else '0%'

    # Section header
    _section_title_table(doc, '1', 'Executive Summary',
                         'High-level overview of assessment findings and compliance posture',
                         accent)
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Purpose & Scope
    _add_heading2(doc, 'Assessment Purpose & Scope', accent)
    proj_desc = (getattr(project, 'description', '') or '').strip()
    purpose = (
        f'OziCyber has conducted a comprehensive gap analysis of {project.title}\'s information '
        f'security controls against the requirements of {framework_name}. '
    )
    if proj_desc:
        purpose += proj_desc + ' '
    purpose += (
        f'The purpose of this assessment is to identify gaps between the organisation\'s current '
        f'control environment and the requirements of the framework, to quantify the overall '
        f'compliance posture, and to provide prioritised, actionable recommendations to address '
        f'identified deficiencies. The assessment evaluated {total} controls across all applicable '
        f'control families.'
    )
    _add_body_paragraph(doc, purpose, before=0, after=8)

    # Overall Posture box
    _add_heading2(doc, 'Overall Compliance Posture', accent)
    posture_bg = {
        'Critical Risk': 'FFF5F5',
        'High Risk':     'FFF8F0',
        'Moderate Risk': 'FFFBF0',
        'Low Risk':      'F0FFF4',
    }.get(risk_label, 'F5F5F5')

    posture_table = doc.add_table(rows=1, cols=2, width=CONTENT_WIDTH)
    _make_table_full_width(posture_table)
    _remove_table_borders(posture_table)
    row = posture_table.rows[0]

    # Left cell: big percentage
    left = row.cells[0]
    _set_col_width(left, 1800)
    _set_cell_color(left, posture_bg)
    _set_cell_vertical_align(left, 'center')
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(lp, before=10, after=4)
    _add_colored_run(lp, f'{pct}%', risk_color, bold=True, size=32)

    lp2 = left.add_paragraph()
    lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(lp2, before=2, after=10)
    _add_colored_run(lp2, risk_label, risk_color, bold=True, size=13)

    # Right cell: narrative
    right = row.cells[1]
    _set_cell_color(right, posture_bg)
    _set_cell_vertical_align(right, 'center')
    rp = right.paragraphs[0]
    _set_spacing(rp, before=10, after=4)
    posture_text = (
        f'{implemented} of {effective} applicable controls are fully implemented. '
        f'{partial} are partially implemented and {not_started} have not yet been commenced. '
    )
    if pct >= 85:
        posture_text += 'This represents a strong security posture.'
    else:
        posture_text += 'Structured remediation is required to improve the organisation\'s security posture.'
    rr = rp.add_run(posture_text)
    rr.font.name = 'Calibri'
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(85, 85, 85)

    # Border on left side of posture table
    for cell in row.cells:
        _set_cell_borders(cell, color=risk_color, size=12, sides=('left',))

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Key Statistics table
    _add_heading2(doc, 'Key Statistics', accent)
    stats_table = doc.add_table(rows=9, cols=3, width=CONTENT_WIDTH)
    _make_table_full_width(stats_table)
    _set_table_borders(stats_table, color='CCCCCC', size=4)

    _header_row(stats_table, ['Metric', 'Count', 'Percentage'],
                col_widths=[4500, 1500, 1500])

    rows_data = [
        ('Total Controls Assessed', str(total), '100%', True, WHITE),
        ('Implemented', str(implemented), _pct_of(implemented), False, WHITE),
        ('Partially Implemented', str(partial), _pct_of(partial), False, ALT_ROW),
        ('In Progress', str(in_progress), _pct_of(in_progress), False, WHITE),
        ('Planned', str(planned), _pct_of(planned), False, ALT_ROW),
        ('Not Started', str(not_started), _pct_of(not_started), False, WHITE),
        ('Not Applicable', str(not_applicable), _pct_of(not_applicable), False, ALT_ROW),
        ('Total Evidence Items Submitted', str(evidence_total), '—', True, 'E8F4FD'),
    ]
    for i, (label, count_val, pct_val, bold, bg) in enumerate(rows_data):
        data_row = stats_table.rows[i + 1]
        _set_cell_color(data_row.cells[0], bg)
        _set_cell_color(data_row.cells[1], bg)
        _set_cell_color(data_row.cells[2], bg)
        _cell_text(data_row.cells[0], label, bold=bold)
        _cell_text(data_row.cells[1], count_val, bold=bold,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(data_row.cells[2], pct_val,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Top Findings
    _add_heading2(doc, 'Top Findings', accent)
    gaps = [cs for cs in control_statuses
            if not getattr(cs.control, 'is_category', False)
            and cs.status in ('NOT_STARTED', 'PARTIALLY_IMPLEMENTED', 'IN_PROGRESS', 'PLANNED')]
    gaps_sorted = sorted(
        gaps,
        key=lambda cs: (
            {'NOT_STARTED': 0, 'PARTIALLY_IMPLEMENTED': 1, 'IN_PROGRESS': 2, 'PLANNED': 3}.get(
                cs.status, 4),
            cs.control.order
        )
    )
    top5 = gaps_sorted[:5]

    if top5:
        top5_table = doc.add_table(rows=len(top5) + 1, cols=3, width=CONTENT_WIDTH)
        _make_table_full_width(top5_table)
        _set_table_borders(top5_table, color='CCCCCC', size=4)
        _header_row(top5_table, ['Control ID', 'Control Title', 'Status / Risk'],
                    col_widths=[1200, 5000, 2200])
        for i, cs in enumerate(top5):
            data_row = top5_table.rows[i + 1]
            bg = WHITE if i % 2 == 0 else ALT_ROW
            for cell in data_row.cells:
                _set_cell_color(cell, bg)
            _cell_text(data_row.cells[0], cs.control.control_id or '', size=9)
            _cell_text(data_row.cells[1], cs.control.title or '', size=9)
            status_label = STATUS_LABEL.get(cs.status, cs.status)
            risk = RISK_FOR_STATUS.get(cs.status, 'Low')
            risk_color_cell = RISK_COLOR.get(risk, '#999').lstrip('#')
            p = data_row.cells[2].paragraphs[0]
            _set_spacing(p, before=2, after=2)
            _add_colored_run(p, status_label, STATUS_COLOR.get(cs.status, '#999').lstrip('#'),
                             bold=True, size=9)
            p.add_run(' | ')
            _add_colored_run(p, risk, risk_color_cell, bold=True, size=9)
    else:
        _add_body_paragraph(doc,
                            'No critical gaps identified — all applicable controls are implemented.',
                            italic=True, color='888888')

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Key Recommendations
    _add_heading2(doc, 'Key Recommendations', accent)
    rec_statuses_order = ['NOT_STARTED', 'PARTIALLY_IMPLEMENTED', 'IN_PROGRESS', 'PLANNED']
    top_recs = [cs for cs in control_statuses
                if not getattr(cs.control, 'is_category', False)
                and cs.status in rec_statuses_order]
    top_recs = sorted(top_recs,
                      key=lambda cs: rec_statuses_order.index(cs.status))[:5]

    if top_recs:
        for cs in top_recs:
            rec_text = _generate_recommendation(cs)
            _add_numbered(doc, rec_text, size=9)
    else:
        _add_numbered(doc, 'Continue to monitor and maintain current control implementations.', size=9)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Assessor Conclusion
    _add_heading2(doc, "Assessor's Conclusion", accent)
    if pct >= 85:
        conclusion = (
            f'The organisation demonstrates a strong information security posture with {pct}% of '
            f'applicable controls implemented. The assessment findings indicate a mature approach to '
            f'security governance. The identified gaps, while limited, should be addressed to '
            f'achieve full compliance and further strengthen the overall security programme. '
            f'OziCyber commends the organisation on its commitment to information security.'
        )
    elif pct >= 70:
        conclusion = (
            f'The organisation demonstrates a moderate information security posture with {pct}% of '
            f'applicable controls implemented. Whilst the foundational controls appear to be in '
            f'place, a number of important gaps have been identified that require structured '
            f'remediation. OziCyber recommends the development of a prioritised remediation '
            f'roadmap to address the identified findings within agreed timeframes.'
        )
    elif pct >= 40:
        conclusion = (
            f'The assessment has identified significant gaps in the organisation\'s information '
            f'security control environment, with only {pct}% of applicable controls implemented. '
            f'Substantial investment in people, process, and technology will be required to '
            f'achieve compliance with the referenced framework. OziCyber strongly recommends '
            f'the immediate establishment of a formal remediation programme with executive '
            f'sponsorship and adequate resourcing.'
        )
    else:
        conclusion = (
            f'The assessment has identified critical deficiencies in the organisation\'s information '
            f'security control environment, with only {pct}% of applicable controls implemented. '
            f'The organisation faces significant exposure to security threats and regulatory risk. '
            f'OziCyber urges the organisation to treat this assessment as a matter of urgent '
            f'priority and to immediately engage executive leadership to allocate the resources '
            f'necessary to commence a comprehensive remediation programme.'
        )

    # Conclusion box (light blue background table)
    conc_table = doc.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
    _make_table_full_width(conc_table)
    _remove_table_borders(conc_table)
    conc_cell = conc_table.rows[0].cells[0]
    _set_cell_color(conc_cell, 'E8F4FD')
    _set_cell_borders(conc_cell, color=accent.lstrip('#'), size=12, sides=('left',))
    cp = conc_cell.paragraphs[0]
    _set_spacing(cp, before=8, after=8)
    cr = cp.add_run(conclusion)
    cr.font.name = 'Calibri'
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 2: Assessment Scope & Methodology
# ---------------------------------------------------------------------------

def _build_scope_methodology(doc, project, accent: str):
    framework_key = getattr(project.framework, 'key', '')
    framework_name = getattr(project.framework, 'name', 'N/A')

    _section_title_table(doc, '2', 'Assessment Scope & Methodology',
                         'Framework overview, assessment approach, and limitations', accent)
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Framework overview
    _add_heading2(doc, f'Framework Overview: {framework_name}', accent)
    fw_desc = FRAMEWORK_DESCRIPTION.get(framework_key, (
        f'This assessment was conducted against {framework_name}, a recognised information security '
        f'framework providing structured guidance for establishing and maintaining effective '
        f'security controls.'
    ))
    _add_body_paragraph(doc, fw_desc, before=0, after=8)

    # Methodology
    _add_heading2(doc, 'Assessment Methodology', accent)
    _add_body_paragraph(
        doc,
        'The gap analysis was conducted using a structured, evidence-based methodology. '
        'Each control within the framework was individually assessed against the following criteria:',
        before=0, after=4
    )

    methodology_points = [
        ('Documentation Review:', 'Examination of policies, procedures, standards, and other '
         'governance documents provided by the organisation.'),
        ('Evidence Collection:', 'Review of technical and procedural evidence submitted by control '
         'owners, including screenshots, configuration exports, logs, and certificates.'),
        ('Stakeholder Input:', 'Consideration of implementation notes and context provided by '
         'control owners and project stakeholders.'),
        ('Control Mapping:', 'Each control was mapped to one of six implementation statuses: '
         'Implemented, Partially Implemented, In Progress, Planned, Not Started, or Not Applicable.'),
        ('Risk Rating:', 'Gap controls were assigned a risk rating (High, Medium, or Low) based '
         'on the implementation status and the criticality of the control.'),
    ]
    for label, detail in methodology_points:
        p = doc.add_paragraph(style='List Bullet')
        _set_spacing(p, before=2, after=2)
        p.paragraph_format.left_indent = Inches(0.25)
        bold_run = p.add_run(label + ' ')
        _style_run(bold_run, size=10, bold=True)
        detail_run = p.add_run(detail)
        _style_run(detail_run, size=10)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=4, after=0)

    # Status Definitions table
    _add_heading2(doc, 'Status Definitions', accent)
    status_def_data = [
        ('IMPLEMENTED', 'Fully implemented, documented, and evidence can be provided to '
         'demonstrate effectiveness.', 'None', '27ae60'),
        ('PARTIALLY_IMPLEMENTED', 'Some elements are in place but the implementation does not '
         'fully meet framework requirements. Gaps remain.', 'Medium', 'f39c12'),
        ('IN_PROGRESS', 'Active work is underway. Implementation is not yet complete.',
         'Low–Medium', '3498db'),
        ('PLANNED', 'Implementation has been planned and scheduled to commence. No active '
         'implementation work has begun.', 'Low–High', '8e44ad'),
        ('NOT_STARTED', 'No implementation activity has commenced. No evidence or planning '
         'documentation exists.', 'High', 'e74c3c'),
        ('NOT_APPLICABLE', 'The organisation has determined this control does not apply to its '
         'environment. Excluded from compliance calculations.', 'N/A', '95a5a6'),
    ]

    status_table = doc.add_table(rows=len(status_def_data) + 1, cols=3, width=CONTENT_WIDTH)
    _make_table_full_width(status_table)
    _set_table_borders(status_table, color='CCCCCC', size=4)
    _header_row(status_table, ['Status', 'Definition', 'Risk Impact'],
                col_widths=[1800, 5500, 1100])

    for i, (status_key, definition, risk_val, color) in enumerate(status_def_data):
        dr = status_table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ALT_ROW
        for cell in dr.cells:
            _set_cell_color(cell, bg)
        sp_p = dr.cells[0].paragraphs[0]
        _set_spacing(sp_p, before=3, after=3)
        _add_colored_run(sp_p, STATUS_LABEL.get(status_key, status_key), color,
                         bold=True, size=9)
        _cell_text(dr.cells[1], definition, size=9)
        rp = dr.cells[2].paragraphs[0]
        _set_spacing(rp, before=3, after=3)
        _add_colored_run(rp, risk_val, color, bold=True, size=9)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Limitations
    _add_heading2(doc, 'Assessment Limitations', accent)
    limitations = [
        'This assessment is a point-in-time evaluation and may not reflect changes made subsequent to the assessment date.',
        'The assessment is based on self-reported information and submitted evidence. OziCyber has not independently verified all claims.',
        'Not all controls could be fully assessed where evidence or implementation notes were not provided.',
        'This report does not constitute formal certification or accreditation against the referenced framework.',
        'The compliance percentage is calculated against applicable controls only and excludes controls assessed as Not Applicable.',
    ]
    for lim in limitations:
        _add_bullet(doc, lim, size=10)

    # Engagement notes
    notes = getattr(project, 'notes', None)
    if notes:
        sp = doc.add_paragraph()
        _set_spacing(sp, before=6, after=0)
        _add_heading2(doc, 'Engagement Notes', accent)
        notes_table = doc.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
        _make_table_full_width(notes_table)
        _remove_table_borders(notes_table)
        nc = notes_table.rows[0].cells[0]
        _set_cell_color(nc, 'FFFBF0')
        np_ = nc.paragraphs[0]
        _set_spacing(np_, before=6, after=6)
        nr = np_.add_run(notes)
        nr.font.name = 'Calibri'
        nr.font.size = Pt(10)
        nr.font.color.rgb = RGBColor(68, 68, 68)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 3: Compliance Dashboard
# ---------------------------------------------------------------------------

def _build_dashboard(doc, stats: dict, family_stats: list, accent: str):
    pct = stats['pct']
    risk_label = _risk_label_from_pct(pct)
    risk_color = _pct_color(pct).lstrip('#')
    total = stats['total']
    effective = stats['effective']
    implemented = stats['implemented']

    _section_title_table(doc, '3', 'Compliance Dashboard',
                         'Visual overview of compliance status and control distribution', accent)
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Overall Compliance
    _add_heading2(doc, 'Overall Compliance', accent)

    overall_table = doc.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
    _make_table_full_width(overall_table)
    _remove_table_borders(overall_table)
    oc = overall_table.rows[0].cells[0]
    _set_cell_color(oc, LIGHT_GREY)

    # Big percentage
    op1 = oc.paragraphs[0]
    op1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(op1, before=10, after=4)
    _add_colored_run(op1, f'{pct}%', risk_color, bold=True, size=36)

    op2 = oc.add_paragraph()
    op2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(op2, before=0, after=4)
    _add_colored_run(op2, risk_label, risk_color, bold=True, size=14)

    op3 = oc.add_paragraph()
    op3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(op3, before=0, after=6)
    r3 = op3.add_run(f'{implemented} of {effective} applicable controls implemented')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(100, 100, 100)

    # Progress bar (20-column table)
    filled_cols = max(0, min(20, round(pct / 100 * 20)))
    bar_table = doc.add_table(rows=1, cols=20, width=CONTENT_WIDTH)
    _make_table_full_width(bar_table)
    _remove_table_borders(bar_table)
    bar_row = bar_table.rows[0]
    for j in range(20):
        bc = bar_row.cells[j]
        color_fill = risk_color if j < filled_cols else 'E0E0E0'
        _set_cell_color(bc, color_fill)
        _set_col_width(bc, 453)  # 9072 / 20
        bp = bc.paragraphs[0]
        _set_spacing(bp, before=4, after=4)
        bp.add_run('')

    sp = doc.add_paragraph()
    _set_spacing(sp, before=4, after=0)

    # Status Breakdown table
    _add_heading2(doc, 'Status Breakdown', accent)
    status_map = {
        'IMPLEMENTED': stats['implemented'],
        'PARTIALLY_IMPLEMENTED': stats['partial'],
        'IN_PROGRESS': stats['in_progress'],
        'PLANNED': stats['planned'],
        'NOT_STARTED': stats['not_started'],
        'NOT_APPLICABLE': stats['not_applicable'],
    }

    sb_table = doc.add_table(rows=len(STATUS_LABEL) + 1, cols=4, width=CONTENT_WIDTH)
    _make_table_full_width(sb_table)
    _set_table_borders(sb_table, color='CCCCCC', size=4)
    _header_row(sb_table, ['Status', 'Count', '% of Total', 'Visual'],
                col_widths=[2200, 1000, 1000, 4000])

    for i, (s_key, s_label) in enumerate(STATUS_LABEL.items()):
        count = status_map.get(s_key, 0)
        s_pct = round(count / total * 100, 1) if total > 0 else 0.0
        color = STATUS_COLOR.get(s_key, '#999').lstrip('#')
        dr = sb_table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ALT_ROW
        for cell in dr.cells:
            _set_cell_color(cell, bg)

        sp_p = dr.cells[0].paragraphs[0]
        _set_spacing(sp_p, before=3, after=3)
        _add_colored_run(sp_p, s_label, color, bold=True, size=9)

        _cell_text(dr.cells[1], str(count), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(dr.cells[2], f'{s_pct}%', align=WD_ALIGN_PARAGRAPH.CENTER)

        bar_p = dr.cells[3].paragraphs[0]
        _set_spacing(bar_p, before=3, after=3)
        bar_r = bar_p.add_run(_pct_bar_text(s_pct))
        bar_r.font.name = 'Courier New'
        bar_r.font.size = Pt(8)
        bar_r.font.color.rgb = RGBColor.from_string(color)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Family Compliance Summary table
    _add_heading2(doc, 'Family Compliance Summary', accent)
    fam_table = doc.add_table(rows=len(family_stats) + 1, cols=8, width=CONTENT_WIDTH)
    _make_table_full_width(fam_table)
    _set_table_borders(fam_table, color='CCCCCC', size=4)
    _header_row(fam_table,
                ['Family ID', 'Family Name', 'Total', 'Impl.', 'Partial',
                 'In Prog.', 'Not Started', 'Compliance'],
                col_widths=[900, 2500, 600, 600, 700, 700, 900, 1300])

    for i, fam in enumerate(family_stats):
        fam_total = fam.get('total', 0)
        fam_impl = fam.get('implemented', 0)
        fam_partial = fam.get('partial', 0)
        fam_inp = fam.get('in_progress', 0)
        fam_ns = max(
            fam_total - fam_impl - fam_partial - fam_inp
            - fam.get('not_applicable', 0) - fam.get('planned', 0), 0)
        fam_pct = fam.get('pct', 0)
        fam_color = _pct_color(fam_pct).lstrip('#')

        dr = fam_table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ALT_ROW
        for cell in dr.cells:
            _set_cell_color(cell, bg)

        _cell_text(dr.cells[0], str(fam.get('identifier', '')), size=9)
        _cell_text(dr.cells[1], str(fam.get('name', '')), size=9)
        _cell_text(dr.cells[2], str(fam_total), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(dr.cells[3], str(fam_impl), align=WD_ALIGN_PARAGRAPH.CENTER,
                   color='27ae60', bold=True)
        _cell_text(dr.cells[4], str(fam_partial), align=WD_ALIGN_PARAGRAPH.CENTER,
                   color='f39c12', bold=True)
        _cell_text(dr.cells[5], str(fam_inp), align=WD_ALIGN_PARAGRAPH.CENTER,
                   color='3498db', bold=True)
        _cell_text(dr.cells[6], str(fam_ns), align=WD_ALIGN_PARAGRAPH.CENTER,
                   color='e74c3c', bold=True)

        # Colour-code compliance cell
        comp_p = dr.cells[7].paragraphs[0]
        comp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(comp_p, before=3, after=3)
        _add_colored_run(comp_p, f'{fam_pct}%', fam_color, bold=True, size=10)

        # Green ≥70%, amber 40-69%, red <40%
        cell_bg = ('E8F5E9' if fam_pct >= 70 else
                   'FFF8E1' if fam_pct >= 40 else
                   'FFEBEE')
        _set_cell_color(dr.cells[7], cell_bg)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Risk Distribution
    _add_heading2(doc, 'Risk Distribution', accent)
    high_count = stats['not_started']
    medium_count = stats['partial']
    low_count = stats['in_progress'] + stats['planned']

    risk_table = doc.add_table(rows=4, cols=3, width=CONTENT_WIDTH)
    _make_table_full_width(risk_table)
    _set_table_borders(risk_table, color='CCCCCC', size=4)
    _header_row(risk_table, ['Risk Rating', 'Finding Count', 'Description'],
                col_widths=[1500, 1200, 5800])

    risk_data = [
        ('High', str(high_count), 'Controls not started; significant exposure; prioritise for immediate remediation', 'e65100'),
        ('Medium', str(medium_count), 'Partially implemented; residual gaps exist; schedule remediation within 90 days', 'f57f17'),
        ('Low', str(low_count), 'In progress or planned; monitor to ensure timely completion', '2e7d32'),
    ]
    for i, (rating, cnt, desc, clr) in enumerate(risk_data):
        dr = risk_table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ALT_ROW
        for cell in dr.cells:
            _set_cell_color(cell, bg)
        rp = dr.cells[0].paragraphs[0]
        _set_spacing(rp, before=3, after=3)
        _add_colored_run(rp, rating, clr, bold=True, size=10)
        _cell_text(dr.cells[1], cnt, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        _cell_text(dr.cells[2], desc, size=9, color='666666')

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 4: Key Findings & Issues
# ---------------------------------------------------------------------------

def _build_findings_section(doc, control_statuses, accent: str):
    findings = _get_findings(control_statuses)

    priority_order = {'NOT_STARTED': 0, 'PARTIALLY_IMPLEMENTED': 1, 'IN_PROGRESS': 2, 'PLANNED': 3}
    findings_sorted = sorted(findings, key=lambda cs: (
        priority_order.get(cs.status, 9),
        cs.control.order
    ))

    _section_title_table(
        doc, '4', 'Key Findings & Issues',
        f'{len(findings_sorted)} gap(s) identified requiring remediation', accent
    )
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    _add_body_paragraph(
        doc,
        'The following findings have been identified based on the gap analysis. Each finding '
        'represents a control that has not been fully implemented. Findings are ordered by risk '
        'priority, with the highest-risk items presented first. Each finding includes a description '
        'of the gap, the current state of implementation, and a specific recommendation for '
        'remediation.',
        before=0, after=8
    )

    if not findings_sorted:
        no_gap_table = doc.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
        _make_table_full_width(no_gap_table)
        _remove_table_borders(no_gap_table)
        nc = no_gap_table.rows[0].cells[0]
        _set_cell_color(nc, 'F0FFF4')
        np_ = nc.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(np_, before=14, after=14)
        _add_colored_run(np_,
                         'No gaps identified. All applicable controls are implemented.',
                         '27ae60', bold=True, size=13)
        doc.add_page_break()
        return

    for idx, cs in enumerate(findings_sorted, 1):
        finding_id = f'F-{idx:03d}'
        status = cs.status
        risk = RISK_FOR_STATUS.get(status, 'Low')
        risk_color_hex = RISK_COLOR.get(risk, '#999').lstrip('#')
        status_color_hex = STATUS_COLOR.get(status, '#999').lstrip('#')
        finding_desc = _generate_finding_description(cs)
        recommendation = _generate_recommendation(cs)

        notes = (cs.implementation_notes or '').strip()
        current_state = notes if notes else 'No implementation evidence or notes have been provided.'

        ev_list = _get_evidence_list(cs)
        ev_count = len(ev_list)

        # Finding header row (colored per risk)
        header_table = doc.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
        _make_table_full_width(header_table)
        _remove_table_borders(header_table)
        hc = header_table.rows[0].cells[0]
        _set_cell_color(hc, NAVY)
        hp = hc.paragraphs[0]
        _set_spacing(hp, before=6, after=6)
        _add_colored_run(hp, f'{finding_id}  ', '64b5f6', bold=True, size=10)
        _add_colored_run(hp, f'Control: {cs.control.control_id}  ', 'AAAAAA', size=10)
        _add_colored_run(hp, f'{cs.control.title}  ', 'FFFFFF', bold=True, size=10)
        _add_colored_run(hp, f'Risk: {risk}', risk_color_hex, bold=True, size=10)

        # Body details table (2 cols: label | content)
        body_table = doc.add_table(rows=5, cols=2, width=CONTENT_WIDTH)
        _make_table_full_width(body_table)
        _set_table_borders(body_table, color='E0E0E0', size=4)

        body_labels = [
            'Current Status',
            'Current State',
            'Finding Description',
            'Evidence',
            'Recommendation',
        ]
        body_contents = [None, None, None, None, None]  # set below
        col_widths_body = [1400, 7200]

        for bi, label in enumerate(body_labels):
            row = body_table.rows[bi]
            label_cell = row.cells[0]
            _set_col_width(label_cell, col_widths_body[0])
            _set_cell_color(label_cell, 'F0F0F0')
            lp = label_cell.paragraphs[0]
            _set_spacing(lp, before=4, after=4)
            lr = lp.add_run(label)
            lr.font.name = 'Calibri'
            lr.font.size = Pt(8)
            lr.font.color.rgb = RGBColor(85, 85, 85)
            lr.bold = True

            content_cell = row.cells[1]
            _set_col_width(content_cell, col_widths_body[1])
            _set_cell_color(content_cell, WHITE)
            cp_ = content_cell.paragraphs[0]
            _set_spacing(cp_, before=4, after=4)

            if bi == 0:  # Current Status
                _add_colored_run(cp_, STATUS_LABEL.get(status, status),
                                 status_color_hex, bold=True, size=9)
            elif bi == 1:  # Current State
                cr_ = cp_.add_run(current_state)
                cr_.font.name = 'Calibri'
                cr_.font.size = Pt(9)
                cr_.font.color.rgb = RGBColor(85, 85, 85)
                cr_.italic = True
            elif bi == 2:  # Finding Description
                cr_ = cp_.add_run(finding_desc)
                cr_.font.name = 'Calibri'
                cr_.font.size = Pt(9)
                cr_.font.color.rgb = RGBColor(51, 51, 51)
            elif bi == 3:  # Evidence
                if ev_count > 0:
                    _add_colored_run(cp_, f'{ev_count} evidence item(s) submitted: ',
                                     NAVY, bold=True, size=9)
                    titles = [getattr(ev, 'title', None) or 'Untitled'
                              for ev in ev_list[:5]]
                    plain = cp_.add_run(', '.join(titles))
                    plain.font.name = 'Calibri'
                    plain.font.size = Pt(9)
                    if ev_count > 5:
                        plain.text += f', ... (+{ev_count - 5} more)'
                else:
                    _add_colored_run(cp_, 'No evidence items attached.', 'e74c3c', size=9)
            elif bi == 4:  # Recommendation
                _set_cell_color(content_cell, 'E8F4FD')
                cr_ = cp_.add_run(recommendation)
                cr_.font.name = 'Calibri'
                cr_.font.size = Pt(9)
                cr_.font.color.rgb = RGBColor(51, 51, 51)

        # Spacer between findings
        sp = doc.add_paragraph()
        _set_spacing(sp, before=0, after=8)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 5: Control-by-Control Analysis
# ---------------------------------------------------------------------------

def _build_control_analysis(doc, control_statuses, family_stats: list, accent: str):
    # Build family_id -> [cs, ...] map
    family_map: dict = {}
    for cs in control_statuses:
        if getattr(cs.control, 'is_category', False):
            continue
        fam_id = cs.control.family_id
        if fam_id not in family_map:
            family_map[fam_id] = []
        family_map[fam_id].append(cs)

    _section_title_table(doc, '5', 'Control-by-Control Analysis',
                         'Detailed assessment of every control, organised by family', accent)
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    _add_body_paragraph(
        doc,
        'This section provides a comprehensive analysis of every control within the framework, '
        'organised by control family. For each control, the current implementation status, '
        'implementation notes, evidence inventory, and gap analysis narrative are presented. '
        'Controls assessed as "Not Applicable" are included for completeness and auditability.',
        before=0, after=10
    )

    for fam in family_stats:
        fam_id = fam.get('id')
        fam_identifier = str(fam.get('identifier', ''))
        fam_name = str(fam.get('name', ''))
        fam_total = fam.get('total', 0)
        fam_impl = fam.get('implemented', 0)
        fam_pct = fam.get('pct', 0)
        fam_color = _pct_color(fam_pct).lstrip('#')

        controls_in_family = sorted(
            family_map.get(fam_id, []),
            key=lambda cs: cs.control.order
        )

        # Family H2 heading with stats
        _add_heading2(doc,
                      f'{fam_identifier} — {fam_name}  '
                      f'({fam_impl} implemented of {fam_total} total — {fam_pct}%)',
                      accent)

        if not controls_in_family:
            _add_body_paragraph(doc, 'No controls assessed for this family.',
                                italic=True, color='AAAAAA')
            continue

        for cs in controls_in_family:
            status = cs.status
            status_color_hex = STATUS_COLOR.get(status, '#999').lstrip('#')
            notes = (cs.implementation_notes or '').strip()
            statement = (cs.control.statement or '').strip()
            gap_text = _generate_gap_analysis_text(cs)
            ev_list = _get_evidence_list(cs)
            owner_name = getattr(cs, 'owner_name', None)
            due_date = getattr(cs, 'due_date', None)
            review_date = getattr(cs, 'review_date', None)

            # H3: control id — title
            _add_heading3(doc,
                          f'{cs.control.control_id} — {cs.control.title}')

            # Status indicator
            status_p = doc.add_paragraph()
            _set_spacing(status_p, before=2, after=4)
            sr1 = status_p.add_run('Status: ')
            sr1.font.name = 'Calibri'
            sr1.font.size = Pt(10)
            sr1.font.color.rgb = RGBColor(85, 85, 85)
            sr1.bold = True
            _add_colored_run(status_p, STATUS_LABEL.get(status, status),
                             status_color_hex, bold=True, size=10)

            # Statement box (light grey single-cell borderless table)
            if statement:
                stmt_table = doc.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
                _make_table_full_width(stmt_table)
                _remove_table_borders(stmt_table)
                sc = stmt_table.rows[0].cells[0]
                _set_cell_color(sc, 'F5F5F5')
                _set_cell_borders(sc, color=accent.lstrip('#'), size=8, sides=('left',))
                stp = sc.paragraphs[0]
                _set_spacing(stp, before=6, after=6)
                str_r = stp.add_run(statement)
                str_r.font.name = 'Calibri'
                str_r.font.size = Pt(9)
                str_r.font.color.rgb = RGBColor(68, 68, 68)
                str_r.italic = True
                sp_s = doc.add_paragraph()
                _set_spacing(sp_s, before=0, after=4)

            # Implementation Notes
            notes_p = doc.add_paragraph()
            _set_spacing(notes_p, before=2, after=2)
            nr1 = notes_p.add_run('Implementation Notes: ')
            nr1.font.name = 'Calibri'
            nr1.font.size = Pt(9)
            nr1.font.color.rgb = RGBColor(85, 85, 85)
            nr1.bold = True
            if notes:
                nr2 = notes_p.add_run(notes)
                nr2.font.name = 'Calibri'
                nr2.font.size = Pt(9)
                nr2.font.color.rgb = RGBColor(68, 68, 68)
            else:
                nr2 = notes_p.add_run('Not yet documented.')
                nr2.font.name = 'Calibri'
                nr2.font.size = Pt(9)
                nr2.font.color.rgb = RGBColor(170, 170, 170)
                nr2.italic = True

            # Evidence
            if ev_list:
                ev_header_p = doc.add_paragraph()
                _set_spacing(ev_header_p, before=4, after=2)
                eh = ev_header_p.add_run('Evidence:')
                eh.font.name = 'Calibri'
                eh.font.size = Pt(9)
                eh.font.color.rgb = RGBColor(85, 85, 85)
                eh.bold = True
                for ev in ev_list:
                    ev_title = getattr(ev, 'title', None) or 'Untitled'
                    ev_desc = getattr(ev, 'description', '') or ''
                    ev_type = ('File' if (getattr(ev, 'file', None) and str(ev.file))
                               else ('URL' if getattr(ev, 'url', None) else 'Note'))
                    uploaded_at = _fmt_date(getattr(ev, 'uploaded_at', None))
                    ev_line = f'{ev_title} [{ev_type}]'
                    if uploaded_at:
                        ev_line += f' — {uploaded_at}'
                    if ev_desc:
                        ev_line += f'\n  {_truncate(ev_desc, 100)}'
                    _add_bullet(doc, ev_line, size=9)

            # Gap Analysis
            gap_table = doc.add_table(rows=1, cols=1, width=CONTENT_WIDTH)
            _make_table_full_width(gap_table)
            _remove_table_borders(gap_table)
            gc = gap_table.rows[0].cells[0]
            _set_cell_color(gc, 'F0F4F8')
            _set_cell_borders(gc, color='90A4AE', size=8, sides=('left',))
            gp1 = gc.paragraphs[0]
            _set_spacing(gp1, before=6, after=2)
            gb1 = gp1.add_run('Gap Analysis: ')
            gb1.font.name = 'Calibri'
            gb1.font.size = Pt(9)
            gb1.font.color.rgb = RGBColor(51, 51, 51)
            gb1.bold = True
            gp2 = gc.add_paragraph()
            _set_spacing(gp2, before=0, after=6)
            gb2 = gp2.add_run(gap_text)
            gb2.font.name = 'Calibri'
            gb2.font.size = Pt(9)
            gb2.font.color.rgb = RGBColor(51, 51, 51)

            # Owner / due date line
            meta_parts = []
            if owner_name:
                meta_parts.append(f'Owner: {owner_name}')
            if due_date:
                meta_parts.append(f'Due: {_fmt_date(due_date)}')
            if review_date:
                meta_parts.append(f'Review: {_fmt_date(review_date)}')
            if meta_parts:
                meta_p = doc.add_paragraph()
                _set_spacing(meta_p, before=4, after=2)
                mr = meta_p.add_run(' | '.join(meta_parts))
                mr.font.name = 'Calibri'
                mr.font.size = Pt(8)
                mr.font.color.rgb = RGBColor(150, 150, 150)

            # Thin separator between controls
            _add_rule(doc, color='CCCCCC', thickness=4)
            sp_r = doc.add_paragraph()
            _set_spacing(sp_r, before=0, after=6)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 6: Policy Gap Analysis
# ---------------------------------------------------------------------------

def _build_policy_analysis(doc, control_statuses, framework_key: str, accent: str):
    expected = EXPECTED_POLICIES.get(framework_key, [])
    observed = _gather_policy_evidence(control_statuses)
    observed_titles = [p['title'].lower() for p in observed]

    def _is_observed(policy_name: str) -> bool:
        pn_lower = policy_name.lower()
        keywords = [w for w in pn_lower.split() if len(w) > 3]
        return any(
            any(kw in obs for kw in keywords)
            for obs in observed_titles
        )

    gap_policies = []
    policy_rows = []  # list of (name, observed_bool, matched_str)
    for pol in expected:
        obs = _is_observed(pol)
        if obs:
            pol_lower = pol.lower()
            kws = [w for w in pol_lower.split() if len(w) > 3]
            matches = [o['title'] for o in observed
                       if any(kw in o['title'].lower() for kw in kws)]
            matched_str = ', '.join(matches[:3]) if matches else 'Inferred from evidence'
        else:
            gap_policies.append(pol)
            matched_str = 'Not found in evidence'
        policy_rows.append((pol, obs, matched_str))

    obs_count = len(observed)
    gap_count = len(gap_policies)
    fw_display = framework_key.replace('_', ' ')

    _section_title_table(doc, '6', 'Policy Gap Analysis',
                         'Assessment of expected policy documentation against observed evidence',
                         accent)
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    _add_body_paragraph(
        doc,
        f'Effective information security governance requires a comprehensive suite of policy '
        f'documents to establish expectations, assign responsibilities, and provide a basis for '
        f'accountability. This section compares the expected policy documents for the '
        f'{fw_display} framework against policy-related evidence items observed in the assessment '
        f'evidence inventory.',
        before=0, after=8
    )

    # Summary stats
    summary_table = doc.add_table(rows=1, cols=3, width=CONTENT_WIDTH)
    _make_table_full_width(summary_table)
    _remove_table_borders(summary_table)
    row = summary_table.rows[0]
    stats_info = [
        (str(len(expected)), 'Expected Policies', '1a1a2e'),
        (str(obs_count), 'Observed in Evidence', '27ae60'),
        (str(gap_count), 'Policy Gaps', 'e74c3c'),
    ]
    for i, (val, label, clr) in enumerate(stats_info):
        cell = row.cells[i]
        _set_cell_color(cell, LIGHT_GREY)
        _set_cell_vertical_align(cell, 'center')
        vp = cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(vp, before=8, after=4)
        _add_colored_run(vp, val, clr, bold=True, size=24)
        lp = cell.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(lp, before=0, after=8)
        lr = lp.add_run(label)
        lr.font.name = 'Calibri'
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(100, 100, 100)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Expected Policy Register
    _add_heading2(doc, 'Expected Policies', accent)
    if policy_rows:
        pol_table = doc.add_table(rows=len(policy_rows) + 1, cols=3, width=CONTENT_WIDTH)
        _make_table_full_width(pol_table)
        _set_table_borders(pol_table, color='CCCCCC', size=4)
        _header_row(pol_table, ['Policy Name', 'Status', 'Evidence Reference'],
                    col_widths=[3500, 1200, 3800])

        for i, (pol_name, obs, matched) in enumerate(policy_rows):
            dr = pol_table.rows[i + 1]
            bg = 'F0FFF4' if obs else 'FFF5F5'
            for cell in dr.cells:
                _set_cell_color(cell, bg)
            _cell_text(dr.cells[0], pol_name, size=9)
            sp_p = dr.cells[1].paragraphs[0]
            _set_spacing(sp_p, before=3, after=3)
            if obs:
                _add_colored_run(sp_p, '✓ Found', '27ae60', bold=True, size=9)
            else:
                _add_colored_run(sp_p, '✗ Gap', 'e74c3c', bold=True, size=9)
            _cell_text(dr.cells[2], matched, size=9,
                       color='555555', italic=not obs)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Policy Gaps
    _add_heading2(doc, 'Policy Gaps', accent)
    if gap_policies:
        for gp in gap_policies:
            _add_bullet(doc,
                        f'{gp}: This policy document was not found in the evidence inventory. '
                        f'Develop and formally approve this document as a priority.',
                        size=9)
    else:
        _add_body_paragraph(doc,
                            '✓ All expected policy documents were observed in the evidence inventory.',
                            color='27ae60', bold=True)

    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    # Recommendations
    _add_heading2(doc, 'Recommendations', accent)
    pol_recs = [
        'Develop a formal Policy Management Framework to govern the lifecycle of all security policies.',
        'Ensure all policies are formally approved by appropriate senior leadership prior to publication.',
        'Implement version control and annual review cycles for all policy documents.',
        'Communicate all policies to relevant staff and maintain acknowledgement records.',
        'Establish a formal process for granting and tracking policy exceptions.',
    ]
    for i, rec in enumerate(pol_recs, 1):
        _add_numbered(doc, rec, size=9)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section 7: Recommendations Register
# ---------------------------------------------------------------------------

def _build_recommendations(doc, control_statuses, family_stats: list, accent: str):
    family_names: dict = {fam.get('id'): fam.get('name', '') for fam in family_stats}

    gap_statuses = {'NOT_STARTED', 'PARTIALLY_IMPLEMENTED', 'IN_PROGRESS', 'PLANNED'}
    priority_map = {
        'NOT_STARTED':           ('High',   'Immediate'),
        'PARTIALLY_IMPLEMENTED': ('Medium', 'Short-term'),
        'IN_PROGRESS':           ('Low',    'Medium-term'),
        'PLANNED':               ('Low',    'Medium-term'),
    }
    effort_map = {
        'NOT_STARTED':           'High',
        'PARTIALLY_IMPLEMENTED': 'Medium',
        'IN_PROGRESS':           'Low',
        'PLANNED':               'Medium',
    }
    priority_order = {'NOT_STARTED': 0, 'PARTIALLY_IMPLEMENTED': 1, 'IN_PROGRESS': 2, 'PLANNED': 3}
    priority_colors = {'High': 'e65100', 'Medium': 'f57f17', 'Low': '2e7d32'}

    recs = [cs for cs in control_statuses
            if not getattr(cs.control, 'is_category', False)
            and cs.status in gap_statuses]
    recs_sorted = sorted(recs,
                         key=lambda cs: (priority_order.get(cs.status, 9), cs.control.order))

    _section_title_table(doc, '7', 'Recommendations Register',
                         'Comprehensive list of remediation actions ordered by priority', accent)
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    _add_body_paragraph(
        doc,
        'The following table presents all recommendations arising from this gap analysis, ordered '
        'by priority. Each recommendation is linked to the relevant control reference and includes '
        'an indicative timeframe and implementation effort rating. Priority: High items should be '
        'treated as immediate remediation actions. Priority: Medium items should be scheduled '
        'within 90 days. Priority: Low items should be tracked and completed within 180 days or '
        'by the agreed target date.',
        before=0, after=8
    )

    if not recs_sorted:
        _add_body_paragraph(
            doc,
            'No gaps identified — all applicable controls are implemented.',
            italic=True, color='888888'
        )
        doc.add_page_break()
        return

    rec_table = doc.add_table(rows=len(recs_sorted) + 1, cols=6, width=CONTENT_WIDTH)
    _make_table_full_width(rec_table)
    _set_table_borders(rec_table, color='CCCCCC', size=4)
    _header_row(rec_table,
                ['#', 'Control Ref', 'Recommendation', 'Priority', 'Effort', 'Family'],
                col_widths=[500, 900, 4500, 800, 800, 1000])

    for i, cs in enumerate(recs_sorted):
        dr = rec_table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ALT_ROW
        for cell in dr.cells:
            _set_cell_color(cell, bg)

        rec_text = _generate_recommendation(cs)
        priority_label, timeframe = priority_map.get(cs.status, ('Medium', 'Short-term'))
        effort = effort_map.get(cs.status, 'Medium')
        fam_name = family_names.get(cs.control.family_id, '')
        p_color = priority_colors.get(priority_label, '333333')

        _cell_text(dr.cells[0], f'R-{i + 1:03d}', size=8,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(dr.cells[1], cs.control.control_id or '', size=8)
        _cell_text(dr.cells[2], rec_text, size=8)

        pp = dr.cells[3].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(pp, before=3, after=3)
        _add_colored_run(pp, priority_label, p_color, bold=True, size=9)

        _cell_text(dr.cells[4], effort, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(dr.cells[5], fam_name, size=8)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Appendix A: Evidence Inventory
# ---------------------------------------------------------------------------

def _build_appendix_evidence(doc, control_statuses, accent: str):
    _section_title_table(doc, 'A', 'Appendix A: Evidence Inventory',
                         'All evidence items across assessed controls', accent)
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    _add_body_paragraph(
        doc,
        'The following table lists all evidence items that have been uploaded to the assessment '
        'portal across all controls. Evidence items are linked to their respective control '
        'references. The existence of an evidence item in this inventory does not constitute '
        'confirmation that the evidence demonstrates full control implementation; all evidence '
        'should be reviewed in the context of the relevant control requirements.',
        before=0, after=8
    )

    # Gather all evidence
    all_evidence = []
    for cs in control_statuses:
        if getattr(cs.control, 'is_category', False):
            continue
        ev_list = _get_evidence_list(cs)
        for ev in ev_list:
            all_evidence.append((cs, ev))

    if not all_evidence:
        _add_body_paragraph(doc,
                            'No evidence has been submitted for this project.',
                            italic=True, color='AAAAAA')
        doc.add_page_break()
        return

    ev_table = doc.add_table(rows=len(all_evidence) + 1, cols=5, width=CONTENT_WIDTH)
    _make_table_full_width(ev_table)
    _set_table_borders(ev_table, color='CCCCCC', size=4)
    _header_row(ev_table,
                ['Control', 'Evidence Title', 'Description', 'Type', 'Date'],
                col_widths=[900, 2200, 3500, 700, 1100])

    for i, (cs, ev) in enumerate(all_evidence):
        dr = ev_table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ALT_ROW
        for cell in dr.cells:
            _set_cell_color(cell, bg)

        ev_title = getattr(ev, 'title', None) or 'Untitled'
        ev_desc = _truncate(getattr(ev, 'description', '') or '', 100)
        ev_type = ('File' if (getattr(ev, 'file', None) and str(ev.file))
                   else ('URL' if getattr(ev, 'url', None) else 'Note'))
        uploaded_at = _fmt_date(getattr(ev, 'uploaded_at', None))

        _cell_text(dr.cells[0], cs.control.control_id or '', size=8)
        _cell_text(dr.cells[1], ev_title, size=9, bold=True)
        _cell_text(dr.cells[2], ev_desc, size=8, color='555555')
        _cell_text(dr.cells[3], ev_type, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(dr.cells[4], uploaded_at, size=8, color='888888')

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Appendix B: Glossary
# ---------------------------------------------------------------------------

def _build_appendix_glossary(doc, framework_key: str, accent: str):
    always_include = {
        'Gap Analysis', 'Control', 'Implementation Status', 'ISMS',
        'Risk Rating', 'Residual Risk', 'Evidence', 'Remediation',
        'Compensating Control', 'Maturity Level',
    }
    framework_extras = {
        'HIPAA':           {'ePHI'},
        'SOC2':            {'TSC'},
        'NIST_800_171_R3': {'CUI', 'SSP'},
        'ISO_27001_2022':  {'SoA', 'ISMS'},
        'NIST_CSF_2':      set(),
    }
    extras = framework_extras.get(framework_key, set())
    include_terms = always_include | extras

    _section_title_table(doc, 'B', 'Appendix B: Glossary',
                         'Definitions of key terms used in this report', accent)
    sp = doc.add_paragraph()
    _set_spacing(sp, before=6, after=0)

    _add_body_paragraph(
        doc,
        'The following terms are used throughout this report. Definitions are provided '
        'in the context of information security assessments and the referenced framework.',
        before=0, after=8
    )

    terms_to_include = [(term, defn) for term, defn in sorted(GLOSSARY_TERMS.items())
                        if term in include_terms]

    gloss_table = doc.add_table(rows=len(terms_to_include), cols=2, width=CONTENT_WIDTH)
    _make_table_full_width(gloss_table)
    _set_table_borders(gloss_table, color='EEEEEE', size=4)

    for i, (term, definition) in enumerate(terms_to_include):
        dr = gloss_table.rows[i]
        bg = WHITE if i % 2 == 0 else ALT_ROW
        for cell in dr.cells:
            _set_cell_color(cell, bg)

        term_cell = dr.cells[0]
        _set_col_width(term_cell, 2000)
        _cell_text(term_cell, term, bold=True, color=NAVY, size=10)

        _cell_text(dr.cells[1], definition, size=9, color='555555')

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_gap_analysis_docx(project, control_statuses, family_stats) -> bytes:
    """
    Generate a complete, professional Word (.docx) gap analysis report.

    Parameters
    ----------
    project : GrcProject
        The GRC project instance.
    control_statuses : list or QuerySet of GrcControlStatus
        All control statuses for the project (include non-category controls).
    family_stats : list of dict
        Family-level statistics with keys: id, identifier, name, total,
        implemented, partial, in_progress, pct.

    Returns
    -------
    bytes
        Raw bytes of the .docx file.
    """
    # Materialise querysets
    control_statuses = list(control_statuses)
    family_stats = list(family_stats)

    framework_key = getattr(project.framework, 'key', 'NIST_CSF_2')
    accent = _get_accent(framework_key)
    generated_date = _fmt_date(date.today())

    # Compute overall statistics once
    stats = _compute_overall_stats(control_statuses)

    # Initialise document
    doc = _setup_document()
    _setup_header_footer(doc, project.title or 'Gap Analysis Report', generated_date)

    # Build all sections
    _build_cover_page(doc, project, stats, accent, generated_date)
    _build_executive_summary(doc, project, stats, control_statuses, accent)
    _build_scope_methodology(doc, project, accent)
    _build_dashboard(doc, stats, family_stats, accent)
    _build_findings_section(doc, control_statuses, accent)
    _build_control_analysis(doc, control_statuses, family_stats, accent)
    _build_policy_analysis(doc, control_statuses, framework_key, accent)
    _build_recommendations(doc, control_statuses, family_stats, accent)
    _build_appendix_evidence(doc, control_statuses, accent)
    _build_appendix_glossary(doc, framework_key, accent)

    # Serialise to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
